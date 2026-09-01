#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generador de Informes de Seguridad — CREM / TrendAI

v4.1  Eduardo Olivares  Todas las mejoras implementadas:
  • Inventario de activos configurable en config.json
  • Filtros por criticidad de activo (Muy Crítico / Crítico / No Crítico)
  • Badges de criticidad en todas las tablas
  • Filas coloreadas por criticidad del activo
  • Sidebar con inventario completo ordenado por criticidad
  • openSecHL: blink al navegar referencias cruzadas
  • Auto-install dependencias
  • --no-input para cron/scheduler
  • Caché de datos procesados (datos/)
  • Config.json por empresa
  • Modo actualización CSV automático
  • Resumen cambios mes a mes + activos reincidentes
  • Sección ejecutiva Word (1 página)
  • Tendencia mensual en HTML
  • Plantilla Word con branding corporativo genérico
  • Deduplicación sys-conf / sec-conf por tipo de evento
  • Tabla CVEs agrupada por activo en sección 2
  • CVE hyperlinks reales en Word
  • Plan de actuación con acciones concretas
  • Modal detalle de activo en HTML
  • Gráficos Chart.js en resumen (donut + bar)
  • Exportar CSV por sección en HTML
  • Validación de CSVs al arrancar
  • Log de ejecución en informe/log_FECHA.txt

Uso:
    python informe_crem.py
    python informe_crem.py --mes "Mayo 2026"
    python informe_crem.py --mes "Mayo 2026" --no-input
    python informe_crem.py --mes "Mayo 2026" --excels
    python informe_crem.py --mes "Mayo 2026" --solo-word   (usa caché de datos/)

Carpetas:
    csv/              CSVs del mes actual
    csv-MES-AÑO/      CSVs del mes anterior (ej: csv-abril-2026)
    datos/            Caché parquet (auto)
    informe/          Word + HTML + log (salida)
    plantilla/        Revisión_CREM_MES_AÑO.docx


ANTES DE TOCAR ESTE FICHERO
───────────────────────────
    python tests/test_regresion.py

Comprueba 80 cifras de un mes real ya cerrado. Aquí los fallos no son
excepciones: el proceso termina con ✓ y el informe sale con los números mal,
y eso llega al cliente sin que nadie lo note. Pásala antes y después de
cualquier cambio.

Tres reglas que evitan la mayoría de los errores que ha tenido este programa:

  1. Un módulo se referencia por su `id`, nunca por su etiqueta en castellano.
     Todo se declara en `MODULOS` (sección 8b). El mismo módulo llega a tener
     dos nombres visibles distintos, y emparejar por texto ya dejó tablas
     enteras en blanco.

  2. Los datos de CUALQUIER mes se leen con `cargar_mes()`. Si el mes anterior
     o el histórico se leen por otro camino, se acaban comparando filas
     agregadas contra filas crudas y el informe publica caídas que no existen.

  3. Un fallo que recorte el informe se anota con `degradado()`, no con
     `warn()`. Así sale en el resumen final, en el log y dentro del propio
     HTML, en vez de perderse.
"""
# ==============================================================================
# 0. AUTO-INSTALL
# ==============================================================================
import subprocess, sys, importlib, importlib.util

_DEPS = {
    "pandas":   "pandas>=2.0",
    "docx":     "python-docx>=1.0",
    "lxml":     "lxml>=4.9",
    "rich":     "rich>=13.0",
    "openpyxl": "openpyxl>=3.1",
}

def _ensure_deps():
    missing = [pkg for mod, pkg in _DEPS.items()
               if not importlib.util.find_spec(mod)]
    if missing:
        print(f"[INFO] Instalando: {', '.join(missing)}")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet"] + missing)
        print("[INFO] Listo.\n")

_ensure_deps()

# ==============================================================================
# 1. IMPORTS
# ==============================================================================
import argparse, contextlib, copy, json, logging, os, re, time, traceback
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Callable, Iterable, NamedTuple, Optional

if sys.platform == "win32":
    try:
        os.system("mode con cols=120 lines=40 >nul 2>&1")
    except Exception:
        pass
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

import pandas as pd
from docx import Document
from lxml import etree
from rich.console import Console
from rich.panel import Panel
from rich.progress import (BarColumn, MofNCompleteColumn, Progress,
                           SpinnerColumn, TaskProgressColumn, TextColumn,
                           TimeElapsedColumn, TimeRemainingColumn)
from rich.rule import Rule
from rich.table import Table as RichTable
from rich.text import Text as RichText
from rich.theme import Theme
import openpyxl
from openpyxl.utils import get_column_letter

# Enriquecimiento de CVEs (NVD + KEV + EPSS). Import opcional: si falta el módulo,
# el generador sigue funcionando sin la sección de soluciones detalladas.
try:
    import cve_enrich
except Exception:
    cve_enrich = None

try:
    from herramientas import security_frameworks
except Exception:
    try:
        import security_frameworks
    except Exception:
        security_frameworks = None

_theme = Theme({
    "info":"cyan","warning":"bold yellow","error":"bold red",
    "success":"bold green","dim":"dim white","new":"bold magenta",
})
console = Console(theme=_theme, highlight=False)

VERSION  = "4.0"
AUTOR    = "Eduardo Olivares"
PROYECTO = "Informe de Seguridad CREM — TrendAI"

# ==============================================================================
# 2. CONFIG & CONSTANTS
# ==============================================================================
# Directorio raíz del proyecto: SIEMPRE junto a este script, nunca el CWD.
# Así el script (y sus datos en CLIENTES/, plantilla/) es portable: cualquier
# usuario, en cualquier equipo, puede ejecutarlo desde cualquier carpeta.
BASE_DIR = Path(__file__).resolve().parent

PLANTILLA    = BASE_DIR / "plantilla" / "Revisión_CREM_MES_AÑO.docx"


# ── Contexto de la ejecución ──────────────────────────────────────────────────
# Esto eran seis variables globales que main() iba mutando con `global`. Dos
# consecuencias reales: no se podía generar más de una empresa por proceso, y el
# orden de llamada era un contrato implícito (una función normalizaba una
# carpeta y otra, más abajo, dependía de que eso ya hubiera pasado).
#
# Ahora el estado se construye entero de una vez y se instala con
# `usar_contexto()`, que además lo restaura al salir. Eso hace posible generar
# varias empresas seguidas y montar el contexto desde una prueba.
@dataclass
class Contexto:
    empresa:     str  = ""                      # nombre de la empresa cliente
    empresa_dir: Path = BASE_DIR                # [EMPRESA]/
    dir_csv:     Path = BASE_DIR / "csv"        # [EMPRESA]/CSV/  (entrada del mes)
    dir_datos:   Path = BASE_DIR / "datos"      # caché .pkl del mes
    dir_informe: Path = BASE_DIR / "informe"    # [EMPRESA]/INFORMES/[MES_AÑO]/

    @property
    def dir_excels(self) -> Path:
        # Derivado, no duplicado: antes era una global aparte que había que
        # acordarse de reasignar cada vez que cambiaba dir_informe.
        return self.dir_informe / "excels"

    @property
    def dir_historico(self) -> Path:
        return self.empresa_dir / "INFORMES" / "CSV"


CTX = Contexto()


def instalar_contexto(ctx: Contexto) -> Contexto:
    """Instala `ctx` como contexto activo. Para el flujo principal, que genera
    un solo informe y termina."""
    global CTX
    CTX = ctx
    return ctx


@contextlib.contextmanager
def usar_contexto(ctx: Contexto):
    """Igual que `instalar_contexto`, pero restaura el anterior al salir.
    Es lo que permite encadenar varias empresas en el mismo proceso y montar
    el contexto desde una prueba sin ensuciar el estado del módulo."""
    previo = CTX
    instalar_contexto(ctx)
    try:
        yield ctx
    finally:
        instalar_contexto(previo)

for _d in [BASE_DIR / "plantilla"]:
    _d.mkdir(exist_ok=True)

MESES_ES = {
    "January":"Enero","February":"Febrero","March":"Marzo","April":"Abril",
    "May":"Mayo","June":"Junio","July":"Julio","August":"Agosto",
    "September":"Septiembre","October":"Octubre","November":"Noviembre","December":"Diciembre",
}
MESES_ES_INV = {v.lower(): k for k, v in MESES_ES.items()}
NIVEL_PESO = {"Critical":5,"High":4,"Medium":3,"Low":1,"":0}

# Word brand colors (marca EMPRESA)
C_RED="D52B1E"; C_WHITE="FFFFFF"; C_MGRAY="E0E0E0"
C_CRIT="FDECEA"; C_HIGH="FFF3E0"; C_MED="FFFDE7"; C_BORDE=C_MGRAY

# Word table indices (plantilla v4 — estilo EMPRESA)
# T0=resumen  T1-T5=infoboxes top5  T6=cambios  T7=infobox reincidentes
# T8=reincidentes  T9=cve_events  T10=cve_assets  T11=sysconf  T12=secconf
# T13=threats  T14=anomaly  T15=cloud  T16=accounts
# T17-T20=infoboxes prioridad  T21=plan
TBL_RESUMEN=0; TBL_CAMBIOS=6; TBL_REINCID=8
TBL_CVE=9; TBL_CVE_ASSET=10; TBL_SYSCONF=11; TBL_SECCONF=12
TBL_THREATS=13; TBL_ANOMALY=14; TBL_CLOUD=15; TBL_ACCOUNTS=16
TBL_PLAN=21
N_TABLAS=22
# Infoboxes top5 acciones ejecutivas (T1-T5) — se rellenan con texto
TBL_ACCION = [1, 2, 3, 4, 5]

_ENCODINGS   = ["utf-8-sig","utf-8","cp1252","latin-1","iso-8859-15"]
CVE_BASE_URL = "https://www.cve.org/CVERecord?id="

# ──────────────────────────────────────────────────────────────────────────────
# Iconos SVG (estilo lineal, currentColor) — sustituyen a los emojis en el HTML.
# Sin llaves {} para poder incrustarlos en f-strings sin escapar.
# ──────────────────────────────────────────────────────────────────────────────
_ICO_PATHS = {
    "search":   '<circle cx="11" cy="11" r="7"/><line x1="21" y1="21" x2="16.65" y2="16.65"/>',
    "cve":      '<rect x="3" y="11" width="18" height="11" rx="2"/><path d="M7 11V7a5 5 0 0 1 9.9-1"/>',
    "ext":      '<path d="M18 13v6a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2V8a2 2 0 0 1 2-2h6"/><polyline points="15 3 21 3 21 9"/><line x1="10" y1="14" x2="21" y2="3"/>',
    "info":     '<circle cx="12" cy="12" r="10"/><line x1="12" y1="16" x2="12" y2="12"/><line x1="12" y1="8" x2="12.01" y2="8"/>',
    "download": '<path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/>',
    "dot":      '<circle cx="12" cy="12" r="7"/>',
    "new":      '<path d="M12 2 15.09 8.26 22 9.27l-5 4.87 1.18 6.88L12 17.77l-6.18 3.25L7 14.14 2 9.27l6.91-1.01L12 2z"/>',
    "check":    '<polyline points="20 6 9 17 4 12"/>',
    "chart":    '<line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/>',
    "device":   '<rect x="2" y="3" width="20" height="14" rx="2"/><line x1="8" y1="21" x2="16" y2="21"/><line x1="12" y1="17" x2="12" y2="21"/>',
    "gear":     '<circle cx="12" cy="12" r="3"/><path d="M19.4 15a1.65 1.65 0 0 0 .33 1.82l.06.06a2 2 0 1 1-2.83 2.83l-.06-.06a1.65 1.65 0 0 0-1.82-.33 1.65 1.65 0 0 0-1 1.51V21a2 2 0 0 1-4 0v-.09A1.65 1.65 0 0 0 9 19.4a1.65 1.65 0 0 0-1.82.33l-.06.06a2 2 0 1 1-2.83-2.83l.06-.06a1.65 1.65 0 0 0 .33-1.82 1.65 1.65 0 0 0-1.51-1H3a2 2 0 0 1 0-4h.09A1.65 1.65 0 0 0 4.6 9a1.65 1.65 0 0 0-.33-1.82l-.06-.06a2 2 0 1 1 2.83-2.83l.06.06a1.65 1.65 0 0 0 1.82.33H9a1.65 1.65 0 0 0 1-1.51V3a2 2 0 0 1 4 0v.09a1.65 1.65 0 0 0 1 1.51 1.65 1.65 0 0 0 1.82-.33l.06-.06a2 2 0 1 1 2.83 2.83l-.06.06a1.65 1.65 0 0 0-.33 1.82V9a1.65 1.65 0 0 0 1.51 1H21a2 2 0 0 1 0 4h-.09a1.65 1.65 0 0 0-1.51 1z"/>',
    "shield":   '<path d="M12 22s8-4 8-10V5l-8-3-8 3v7c0 6 8 10 8 10z"/>',
    "alert":    '<path d="M10.29 3.86 1.82 18a2 2 0 0 0 1.71 3h16.94a2 2 0 0 0 1.71-3L13.71 3.86a2 2 0 0 0-3.42 0z"/><line x1="12" y1="9" x2="12" y2="13"/><line x1="12" y1="17" x2="12.01" y2="17"/>',
    "activity": '<polyline points="22 12 18 12 15 21 9 3 6 12 2 12"/>',
    "cloud":    '<path d="M18 10h-1.26A8 8 0 1 0 9 20h9a5 5 0 0 0 0-10z"/>',
    "user":     '<path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2"/><circle cx="12" cy="7" r="4"/>',
    "target":   '<circle cx="12" cy="12" r="10"/><circle cx="12" cy="12" r="6"/><circle cx="12" cy="12" r="2"/>',
    "recycle":  '<polyline points="23 4 23 10 17 10"/><polyline points="1 20 1 14 7 14"/><path d="M3.51 9a9 9 0 0 1 14.85-3.36L23 10M1 14l4.64 4.36A9 9 0 0 0 20.49 15"/>',
    "trend-down":'<polyline points="23 18 13.5 8.5 8.5 13.5 1 6"/><polyline points="17 18 23 18 23 12"/>',
    "trend-up": '<polyline points="23 6 13.5 15.5 8.5 10.5 1 18"/><polyline points="17 6 23 6 23 12"/>',
    "wrench":   '<path d="M14.7 6.3a4 4 0 0 0-5.66 5.66l-6.35 6.35a1 1 0 0 0 0 1.42l1.58 1.58a1 1 0 0 0 1.42 0l6.35-6.35a4 4 0 0 0 5.66-5.66l-2.83 2.83-2.12-2.12 2.83-2.83z"/>',
    "globe":    '<circle cx="12" cy="12" r="10"/><line x1="2" y1="12" x2="22" y2="12"/><path d="M12 2a15.3 15.3 0 0 1 4 10 15.3 15.3 0 0 1-4 10 15.3 15.3 0 0 1-4-10 15.3 15.3 0 0 1 4-10z"/>',
    "phone":    '<rect x="5" y="2" width="14" height="20" rx="2" ry="2"/><line x1="12" y1="18" x2="12.01" y2="18"/>',
    "pin":      '<path d="M21 10c0 7-9 13-9 13s-9-6-9-13a9 9 0 0 1 18 0z"/><circle cx="12" cy="10" r="3"/>',
    "list":     '<line x1="8" y1="6" x2="21" y2="6"/><line x1="8" y1="12" x2="21" y2="12"/><line x1="8" y1="18" x2="21" y2="18"/><line x1="3" y1="6" x2="3.01" y2="6"/><line x1="3" y1="12" x2="3.01" y2="12"/><line x1="3" y1="18" x2="3.01" y2="18"/>',
    "note":     '<path d="M11 4H4a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h14a2 2 0 0 0 2-2v-7"/><path d="M18.5 2.5a2.12 2.12 0 0 1 3 3L12 15l-4 1 1-4 9.5-9.5z"/>',
}

def _ico(name: str, cls: str = "ico") -> str:
    p = _ICO_PATHS.get(name, "")
    fill = "currentColor" if name == "dot" else "none"
    return (f'<svg class="{cls}" viewBox="0 0 24 24" width="1em" height="1em" fill="{fill}" '
            f'stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" '
            f'aria-hidden="true">{p}</svg>')

_ICONS_JS = {k: _ico(k) for k in _ICO_PATHS}  # para inyectar en JS (const ICO)

# Traducción emoji → icono (para reemplazar sin tocar todos los call-sites)
_EMOJI_ICO = {
    "🔎": "search", "📊": "chart", "🔓": "cve", "💻": "device", "⚙️": "gear",
    "🛡️": "shield", "🚨": "alert", "📡": "activity", "☁️": "cloud", "👤": "user",
    "🔮": "target", "📈": "trend-up", "♻️": "recycle", "📉": "trend-down", "🔧": "wrench",
    "🎯": "target", "🌐": "globe", "📱": "phone", "📍": "pin",
}

def _icoify(emoji: str) -> str:
    """Devuelve el SVG del icono equivalente al emoji, o el emoji si no hay mapeo."""
    key = _EMOJI_ICO.get(str(emoji).strip())
    return _ico(key) if key else emoji

# Punto de severidad (sustituye a los círculos de color 🔴🟠🟡🟢)
_SEV_DOT = {"🔴": "#c62828", "🟠": "#e65100", "🟡": "#f57f17", "🟢": "#2e7d32",
            "💀": "#8B0000", "⬜": "#9ca3af"}

def _dot(color: str) -> str:
    return f'<span class="sev-dot" style="color:{color}">{_ico("dot")}</span>'

def _sev_label(txt: str) -> str:
    """Convierte 'emoji TEXTO' (ej. '🟡 MEDIO') en 'punto-icono TEXTO'."""
    s = str(txt)
    parts = s.split(" ", 1)
    if len(parts) == 2 and parts[0] in _SEV_DOT:
        return _dot(_SEV_DOT[parts[0]]) + _esc_min(parts[1])
    return _esc_min(s)

def _esc_min(s: str) -> str:
    return str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _os_sig(os_str: str) -> str:
    """Firma normalizada del SO: 'Microsoft Windows Server 2022 10.0 (Build …)' → 'windows server 2022'."""
    s = str(os_str).lower().replace("microsoft", "").strip()
    m = re.search(r'windows server\s+\d{4}(\s*r2)?', s)
    if m: return m.group(0).strip()
    m = re.search(r'windows\s+(11|10|8\.1|8|7|vista|xp)', s)
    if m: return m.group(0).strip()
    m = re.search(r'(ubuntu|debian|red hat|centos|suse|macos|mac os|android|ios)\s*[\d.]*', s)
    if m: return m.group(0).strip()
    toks = [t for t in re.split(r'[\s()]+', s) if t]
    return " ".join(toks[:3]).strip()
def agrupar_soluciones(device_sols: list[str]) -> list[str]:
    individual_sols = []
    for s in device_sols:
        s_clean = s.strip()
        if s_clean.startswith("- "):
            s_clean = s_clean[2:]
        parts = [p.strip() for p in s_clean.split(";") if p.strip()]
        individual_sols.extend(parts)
    
    grouped = {}
    other_sols = []
    
    pattern = re.compile(
        r'^(.+?):\s*actualizar a (?:la versión\s+)?([0-9a-zA-Z.]+)\s*o superior(.*)$', 
        re.IGNORECASE
    )
    
    for sol in individual_sols:
        m = pattern.match(sol)
        if m:
            product = m.group(1).strip()
            version_str = m.group(2).strip()
            suffix = m.group(3).strip()
            
            v_tuple = []
            for p in version_str.split('.'):
                num_m = re.match(r'\d+', p)
                if num_m:
                    v_tuple.append(int(num_m.group(0)))
                else:
                    v_tuple.append(0)
            v_tuple = tuple(v_tuple)
            
            if product not in grouped:
                grouped[product] = (v_tuple, version_str, suffix)
            else:
                if v_tuple > grouped[product][0]:
                    grouped[product] = (v_tuple, version_str, suffix)
        else:
            if sol not in other_sols:
                other_sols.append(sol)
                
    reconstructed = []
    for product, (v_tuple, version_str, suffix) in sorted(grouped.items()):
        suf_str = f" {suffix}" if suffix else ""
        reconstructed.append(f"{product}: actualizar a la versión {version_str} o superior{suf_str}")
    
    for o in other_sols:
        reconstructed.append(o)
        
    return reconstructed



def convertir_html_a_pdf(html_path: Path, pdf_path: Path) -> bool:
    try:
        from PyQt6.QtWidgets import QApplication
        from PyQt6.QtWebEngineCore import QWebEnginePage
        from PyQt6.QtCore import QUrl, QEventLoop
        import sys
        
        app = QApplication.instance()
        created_app = False
        if not app:
            app = QApplication(sys.argv)
            created_app = True
            
        page = QWebEnginePage()
        loop = QEventLoop()
        
        page.loadFinished.connect(lambda ok: loop.quit())
        page.load(QUrl.fromLocalFile(str(html_path.resolve())))
        loop.exec()
        
        print_loop = QEventLoop()
        printed_ok = False
        
        def on_finished(path, ok):
            nonlocal printed_ok
            printed_ok = ok
            print_loop.quit()
            
        page.pdfPrintingFinished.connect(on_finished)
        page.printToPdf(str(pdf_path.resolve()))
        print_loop.exec()
        
        page.deleteLater()
        if created_app:
            app.quit()
            
        return printed_ok
    except Exception as ex:
        print(f"Error al generar PDF: {ex}")
        return False

def _strip_sev(s: str) -> str:
    """Elimina los círculos de severidad emoji del texto (para vistas de solo texto)."""
    out = str(s)
    for e in _SEV_DOT:
        out = out.replace(e, "")
    return out.strip()

def _svg_stacked_bars(filas, empty_msg="Sin incidencias registradas este mes.") -> str:
    """Gráfico SVG de barras apiladas horizontales (Crítico/Alto · Medio · Bajo) por
    módulo. Autocontenido (offline), responsive y SIN overflow: las barras se escalan al
    total máximo y los números del total van en un carril reservado a la derecha."""
    datos_g = [f for f in filas if f.get("Total", 0) > 0]
    if not datos_g:
        return (f'<div style="padding:24px;text-align:center;color:var(--gray);font-size:13px">'
                f'{_esc_min(empty_msg)}</div>')
    C_CRIT, C_MED, C_LOW = "#da291c", "#e0870b", "#1f9d55"
    maxt = max((f["Total"] for f in datos_g), default=1) or 1
    label_w, bar_w, pad_r = 150, 440, 54
    row_h, gap, top = 24, 12, 16
    W = label_w + bar_w + pad_r
    H = top * 2 + len(datos_g) * (row_h + gap) - gap + 4
    p = [f'<svg viewBox="0 0 {W} {H}" width="100%" height="{H}" '
         f'preserveAspectRatio="xMidYMid meet" role="img" '
         f'style="font-family:var(--font);max-width:100%" xmlns="http://www.w3.org/2000/svg">']
    # Rejilla y escala
    for t in range(5):
        gx = label_w + bar_w * t / 4
        p.append(f'<line x1="{gx:.1f}" y1="{top-4}" x2="{gx:.1f}" y2="{H-top}" '
                 f'stroke="rgba(0,0,0,.06)" stroke-width="1"/>')
        p.append(f'<text x="{gx:.1f}" y="{H-top+13}" font-size="9" fill="#9ca3af" '
                 f'text-anchor="middle">{round(maxt*t/4)}</text>')
    for i, f in enumerate(datos_g):
        y = top + i * (row_h + gap)
        cy = y + row_h / 2
        crit, med, low, tot = f["Alto / Crítico"], f["Medio"], f["Bajo"], f["Total"]
        lbl = _esc_min(str(f["Módulo de Seguridad"]))
        if len(lbl) > 20:
            lbl = lbl[:19] + "…"
        p.append(f'<text x="{label_w-10}" y="{cy+4:.1f}" font-size="11.5" '
                 f'fill="#374151" text-anchor="end" font-weight="600">{lbl}</text>')
        p.append(f'<rect x="{label_w}" y="{y}" width="{bar_w}" height="{row_h}" rx="6" fill="#eef1f6"/>')
        x = float(label_w)
        for val, color in ((crit, C_CRIT), (med, C_MED), (low, C_LOW)):
            if val <= 0:
                continue
            w = bar_w * val / maxt
            p.append(f'<rect x="{x:.1f}" y="{y}" width="{w:.1f}" height="{row_h}" rx="4" '
                     f'fill="{color}"><title>{lbl} · {val}</title></rect>')
            if w >= 20:
                p.append(f'<text x="{x+w/2:.1f}" y="{cy+4:.1f}" font-size="10.5" fill="#fff" '
                         f'text-anchor="middle" font-weight="700">{val}</text>')
            x += w
        p.append(f'<text x="{label_w+bar_w+8}" y="{cy+4:.1f}" font-size="11.5" '
                 f'fill="#374151" font-weight="800">{tot}</text>')
    p.append('</svg>')
    legend = ('<div class="chart-legend">'
              '<span><i style="background:#da291c"></i>Crítico / Alto</span>'
              '<span><i style="background:#e0870b"></i>Medio</span>'
              '<span><i style="background:#1f9d55"></i>Bajo</span></div>')
    return "".join(p) + legend

def _svg_donut(alto: int, medio: int, bajo: int) -> str:
    """Donut SVG de distribución de severidad, con total centrado y leyenda HTML
    legible (sin overflow). Autocontenido y offline."""
    segs = [("Alto / Crítico", int(alto or 0), "#c62828"),
            ("Medio",          int(medio or 0), "#f57f17"),
            ("Bajo",           int(bajo or 0),  "#2e7d32")]
    total = sum(v for _, v, _ in segs)
    size, sw = 168, 30
    r  = (size - sw) / 2
    cx = cy = size / 2
    C  = 2 * 3.141592653589793 * r
    p = [f'<svg viewBox="0 0 {size} {size}" width="100%" height="{size}" '
         f'preserveAspectRatio="xMidYMid meet" style="max-width:{size}px;font-family:var(--font)" '
         f'role="img" xmlns="http://www.w3.org/2000/svg">']
    # Aro de fondo
    p.append(f'<circle cx="{cx}" cy="{cy}" r="{r:.2f}" fill="none" stroke="#eef1f6" stroke-width="{sw}"/>')
    cum = 0.0
    for label, val, color in segs:
        if val <= 0 or total <= 0:
            continue
        frac = val / total
        ang  = -90 + cum * 360
        p.append(f'<circle cx="{cx}" cy="{cy}" r="{r:.2f}" fill="none" stroke="{color}" '
                 f'stroke-width="{sw}" stroke-linecap="butt" '
                 f'stroke-dasharray="{frac*C:.2f} {C:.2f}" '
                 f'transform="rotate({ang:.2f} {cx} {cy})">'
                 f'<title>{_esc_min(label)}: {val} ({frac*100:.0f}%)</title></circle>')
        cum += frac
    # Total centrado
    p.append(f'<text x="{cx}" y="{cy-1}" text-anchor="middle" font-size="30" font-weight="800" '
             f'fill="#1a1d27">{total}</text>')
    p.append(f'<text x="{cx}" y="{cy+17}" text-anchor="middle" font-size="10" font-weight="700" '
             f'fill="#9ca3af" letter-spacing="1.5">TOTAL</text>')
    p.append('</svg>')
    def _pct(v):
        if not total:
            return ""
        pc = v / total * 100
        if 0 < pc < 1:
            return " · &lt;1%"
        return f" · {pc:.0f}%"
    legend = '<div class="chart-legend donut-legend">' + "".join(
        f'<span><i style="background:{color}"></i>{_esc_min(label)} '
        f'<b>{val}</b><small>{_pct(val)}</small></span>' for label, val, color in segs) + '</div>'
    return f'<div class="donut-wrap">{"".join(p)}{legend}</div>'


def _svg_risk_gauge(score: float, color: str, nivel: str) -> str:
    """Gauge semicircular (0–100) autocontenido para el Riesgo CREM ejecutivo.
    Zonas de color de fondo (Bajo/Medio/Alto/Crítico) + arco de valor en el color
    del nivel + aguja marcadora. Offline y responsive."""
    import math
    try:
        s = max(0.0, min(100.0, float(score)))
    except Exception:
        s = 0.0
    W, H = 260, 158
    cx, cy, r = 130, 132, 104
    sw = 20

    def _pt(frac):
        ang = math.radians(180 - frac * 180)
        return cx + r * math.cos(ang), cy - r * math.sin(ang)

    def _arc(f0, f1, col, width, op="1"):
        x0, y0 = _pt(f0); x1, y1 = _pt(f1)
        large = 1 if (f1 - f0) > 0.5 else 0
        return (f'<path d="M {x0:.2f} {y0:.2f} A {r} {r} 0 {large} 1 {x1:.2f} {y1:.2f}" '
                f'fill="none" stroke="{col}" stroke-width="{width}" stroke-linecap="round" '
                f'opacity="{op}"/>')

    p = [f'<svg viewBox="0 0 {W} {H}" width="100%" height="{H}" '
         f'preserveAspectRatio="xMidYMid meet" role="img" '
         f'style="max-width:260px;font-family:var(--font)" xmlns="http://www.w3.org/2000/svg">']
    # Zonas de fondo (tenues)
    zonas = [(0.00, 0.25, "#1f9d55"), (0.25, 0.50, "#f0a91b"),
             (0.50, 0.75, "#e0870b"), (0.75, 1.00, "#da291c")]
    for f0, f1, col in zonas:
        p.append(_arc(f0 + 0.004, f1 - 0.004, col, sw, "0.16"))
    # Arco de valor
    frac = s / 100.0
    if frac > 0.001:
        p.append(_arc(0.0, frac, color, sw, "1"))
    # Aguja / marcador
    nx, ny = _pt(frac)
    p.append(f'<circle cx="{nx:.2f}" cy="{ny:.2f}" r="8" fill="#fff" stroke="{color}" stroke-width="4"/>')
    # Valor central
    s_disp = f"{s:.1f}" if s != int(s) else str(int(s))
    p.append(f'<text x="{cx}" y="{cy-14}" text-anchor="middle" font-size="46" font-weight="800" '
             f'fill="{color}" font-variant-numeric="tabular-nums">{s_disp}</text>')
    p.append(f'<text x="{cx}" y="{cy+6}" text-anchor="middle" font-size="12" font-weight="700" '
             f'fill="{color}" letter-spacing="1.5">{_esc_min(nivel)}</text>')
    # Escala 0 / 100
    p.append(f'<text x="{cx-r}" y="{cy+16}" text-anchor="middle" font-size="10" fill="#9ca3af">0</text>')
    p.append(f'<text x="{cx+r}" y="{cy+16}" text-anchor="middle" font-size="10" fill="#9ca3af">100</text>')
    p.append('</svg>')
    return "".join(p)


def _svg_sparkline(values: list, color: str = "#da291c", w: int = 220, h: int = 46) -> str:
    """Mini-gráfico de línea (sparkline) autocontenido para la tendencia del score.
    Ignora valores None. Devuelve '' si hay menos de 2 puntos válidos."""
    pts = [(i, v) for i, v in enumerate(values) if v is not None]
    if len(pts) < 2:
        return ""
    ys = [float(v) for _, v in pts]
    vmin, vmax = min(ys), max(ys)
    rng = (vmax - vmin) or 1.0
    pad = 6
    n = len(values) - 1 or 1
    def _x(i): return pad + (w - 2 * pad) * (i / n)
    def _y(v): return pad + (h - 2 * pad) * (1 - (v - vmin) / rng)
    coords = [(_x(i), _y(v)) for i, v in pts]
    line = " ".join(f"{x:.1f},{y:.1f}" for x, y in coords)
    area = (f"M {coords[0][0]:.1f},{h-pad:.1f} "
            + " ".join(f"L {x:.1f},{y:.1f}" for x, y in coords)
            + f" L {coords[-1][0]:.1f},{h-pad:.1f} Z")
    lx, ly = coords[-1]
    return (f'<svg viewBox="0 0 {w} {h}" width="100%" height="{h}" preserveAspectRatio="none" '
            f'style="max-width:100%" xmlns="http://www.w3.org/2000/svg">'
            f'<path d="{area}" fill="{color}" opacity="0.08"/>'
            f'<polyline points="{line}" fill="none" stroke="{color}" stroke-width="2" '
            f'stroke-linejoin="round" stroke-linecap="round"/>'
            f'<circle cx="{lx:.1f}" cy="{ly:.1f}" r="3.2" fill="{color}"/></svg>')


# Word XML namespaces
_NS    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_X     = "{" + _NS + "}"
_R_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_R_X   = "{" + _R_NS + "}"
_HYPER = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

# Columnas esperadas por CSV — para validación.
# Un elemento puede ser una tupla de alternativas: basta con que exista UNA.
# (La exportación manual de Vision One fusiona entrada/objetivo en una sola
#  columna «Entry / target assets»; la extracción por API las mantiene separadas.)
_CSV_SCHEMA = {
    "cve-events":          ["Vulnerability ID","CVE impact score","Global exploit potential"],
    "cve-assets":          ["Device name","CVE event risk score","Total CVEs"],
    "threat-detections":   ["Risk event","Asset","Event risk level"],
    "anomaly-detections":  ["Risk event","Asset","Event risk level"],
    "security-conf":       ["Risk event","Asset","Event risk level"],
    "sys-conf":            ["Risk event","Asset","Event risk level"],
    "cloud-app":           ["Risk event","Asset","Event risk level","Detail info"],
    "account-compromise":  ["Risk event","Impact scope","Event risk level"],
    "predictive-analytics":["Risk event",
                            ("Entry assets","Entry / target assets","Detail info"),
                            ("Target assets","Entry / target assets","Detail info"),
                            "Attack path risk score","Detected"],
}

# ==============================================================================
# 3. LOGGING
# ==============================================================================
_log: logging.Logger = logging.getLogger("crem")

def _setup_log(mes_safe: str):
    ruta = CTX.dir_informe / f"log_{mes_safe}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(levelname)-7s  %(message)s",
        handlers=[logging.FileHandler(ruta, encoding="utf-8"),
                  logging.NullHandler()],
    )
    _log.info(f"Informe CREM v{VERSION} — período {mes_safe}")
    return ruta

# ==============================================================================
# 4. CONSOLE HELPERS
# ==============================================================================
def banner():
    console.print(Panel(
        f"[cyan]v{VERSION}  {AUTOR}[/]  ·  CVE diff · Caché · Gráficos · Modal · Export CSV · Log",
        title=f"[bold cyan]{PROYECTO}[/]", border_style="cyan", expand=False, padding=(0,3)))
    console.print()

def seccion(titulo, n, total):
    console.print()
    console.print(Rule(f"[bold blue]\\[{n}/{total}] {titulo}[/]", style="blue", align="left"))

def _plain(msg) -> str:
    """Quita el marcado Rich ([dim], [/], [bold]…) para escribirlo en el log.

    Sin esto el fichero log_*.txt se llenaba de etiquetas literales del tipo
    `[dim]…[/]`, ilegibles para quien revisa la ejecución.
    """
    s = str(msg)
    if "[" not in s:
        return s
    try:
        return RichText.from_markup(s).plain
    except Exception:
        return s

def ok(msg):   console.print(f"  [success]✓[/]  {msg}");   _log.info(_plain(msg))
def warn(msg): console.print(f"  [warning]⚠[/]  [warning]{msg}[/]"); _log.warning(_plain(msg))
def err(msg):  console.print(f"  [error]✗[/]  [error]{msg}[/]");     _log.error(_plain(msg))
def info(msg): console.print(f"  [info]→[/]  {msg}");      _log.info(_plain(msg))

# ── Degradaciones de la ejecución ─────────────────────────────────────────────
# El programa está lleno de `except Exception` que devuelven vacío y siguen. El
# informe sale igual, con menos datos, y nadie se entera hasta que un cliente
# pregunta. Todo fallo que recorte el contenido se anota aquí para: (1) un
# bloque-resumen al final de la ejecución y (2) un aviso visible en el HTML.
DEGRADACIONES: list[dict] = []

def degradado(ambito: str, detalle: str, impacto: str = "") -> None:
    """Registra que una parte del informe ha salido incompleta."""
    DEGRADACIONES.append({"ambito": _plain(ambito),
                          "detalle": _plain(detalle),
                          "impacto": _plain(impacto)})
    warn(f"{ambito}: {detalle}" + (f" [dim]→ {impacto}[/]" if impacto else ""))

def _resumen_degradaciones() -> None:
    """Bloque final con todo lo que salió incompleto. Se imprime siempre que
    haya algo: es la diferencia entre «terminó con ✓» y «terminó bien»."""
    if not DEGRADACIONES:
        return
    console.print()
    console.print(Panel(
        "\n".join(f"[warning]•[/] [bold]{d['ambito']}[/] — {d['detalle']}"
                  + (f"\n  [dim]{d['impacto']}[/]" if d["impacto"] else "")
                  for d in DEGRADACIONES),
        title=f"[warning] AVISOS DE ESTA EJECUCIÓN ({len(DEGRADACIONES)}) [/]",
        border_style="warning", padding=(1, 2)))
    _log.warning(f"La ejecución terminó con {len(DEGRADACIONES)} degradación(es):")
    for d in DEGRADACIONES:
        _log.warning(f"  · {d['ambito']} — {d['detalle']}"
                     + (f" ({d['impacto']})" if d["impacto"] else ""))

# ══════════════════════════════════════════════════════════════════════════════
# SISTEMA DE DISEÑO COMPARTIDO
# ══════════════════════════════════════════════════════════════════════════════
# Los dos informes repetían su propia copia de la paleta. Los mismos colores de
# marca y severidad declarados dos veces, libres de divergir en cuanto alguien
# tocara uno solo. Aquí se declaran UNA vez; cada informe añade después sus
# tokens propios (el técnico tiene navegación y modales; el ejecutivo, no).
CSS_RESET = "*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}"

CSS_TOKENS = """
  /* Marca */
  --red:#D52B1E;--red2:#b71c1c;--red-a:#fff0ee;
  /* Severidad — idéntica en ambos informes, por definición */
  --crit:#c62828;--crit-bg:#fdecea;--crit-border:#ef9a9a;
  --high:#e65100;--high-bg:#fff3e0;--high-border:#ffcc80;
  --med:#f57f17;--med-bg:#fffde7;--med-border:#ffe082;
  --low:#2e7d32;--low-bg:#f1f8e9;--low-border:#a5d6a7;
  /* Acento, fondo y grises comunes */
  --accent:#1e40af;--accent-bg:#eff6ff;
  --bg:#f3f4f8;--gray:#6b7280;--dgray:#9ca3af;
  /* Forma */
  --font:'Inter','Segoe UI',system-ui,-apple-system,sans-serif;
  --r:8px;--r2:12px;
  --sh:0 1px 3px rgba(0,0,0,.06),0 1px 2px rgba(0,0,0,.04);
  --sh2:0 4px 12px rgba(0,0,0,.08),0 2px 6px rgba(0,0,0,.05);
  --hdr-h:56px;"""

# CSS del aviso de degradaciones, común a los dos informes.
CSS_DEGRADADO = """
.degr{background:#fff8e1;border:1px solid #ffe082;border-left:4px solid #f57f17;
  border-radius:8px;padding:12px 16px;margin:0 0 18px}
.degr-t{font-size:12px;font-weight:700;color:#e65100;text-transform:uppercase;
  letter-spacing:.5px;margin-bottom:6px}
.degr ul{margin:0;padding-left:18px}
.degr li{font-size:12.5px;color:#5d4037;margin:3px 0}
.degr .degr-i{color:#8d6e63;font-style:italic}
"""

def _css_tecnico() -> str:
    """Hoja de estilos del informe técnico (navegación, tablas filtrables, modales)."""
    return f"""/* ── RESET & TOKENS ──────────────────────────────────────────────────────── */
{CSS_RESET}
:root{{{CSS_TOKENS}
  /* Propios del informe técnico (navegación, modales, código) */
  --dark:#0f1117;--dark2:#1a1d27;--dark3:#252836;
  --lgray:#f8f9fb;--mgray:#e5e7eb;--border:#e5e7eb;--panel:#ffffff;
  --new-bg:#f3e8ff;--new:#7c3aed;
  --mono:'Cascadia Code','Consolas',monospace;
  --sh3:0 20px 40px rgba(0,0,0,.12);
  --nav-w:228px;--tb-h:48px;
}}
body{{background:var(--bg);color:var(--dark);font-family:var(--font);font-size:13.5px;line-height:1.55;-webkit-font-smoothing:antialiased}}
a{{color:var(--accent);text-decoration:none}}a:hover{{text-decoration:underline}}
code{{font-family:var(--mono);background:var(--lgray);padding:1px 6px;border-radius:4px;font-size:11.5px;border:1px solid var(--border)}}

/* ── HEADER ──────────────────────────────────────────────────────────────── */
.hdr{{
  background:linear-gradient(180deg,#20232f 0%,#171922 100%);
  border-bottom:1px solid rgba(255,255,255,.07);
  position:sticky;top:0;z-index:200;
  box-shadow:0 1px 0 rgba(255,255,255,.04),0 6px 24px rgba(0,0,0,.28);
  height:var(--hdr-h);
}}
.hdr-in{{display:flex;align-items:center;gap:16px;padding:0 24px;height:100%}}
.nav-toggle-btn{{
  display:none;flex-shrink:0;
  background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.12);
  color:#fff;border-radius:8px;width:32px;height:32px;font-size:15px;
  cursor:pointer;align-items:center;justify-content:center;
}}
.nav-toggle-btn:hover{{background:rgba(255,255,255,.14)}}
.logo{{
  display:flex;align-items:center;gap:10px;
  text-decoration:none;flex-shrink:0;
}}
.logo-mark{{
  width:30px;height:30px;border-radius:8px;flex-shrink:0;
  background:linear-gradient(135deg,var(--red),var(--red2));
  display:flex;align-items:center;justify-content:center;color:#fff;
  box-shadow:0 2px 8px rgba(213,43,30,.4),inset 0 1px 0 rgba(255,255,255,.18);
}}
.logo-mark .ico{{width:17px;height:17px}}
.logo-word{{font-size:18.5px;font-weight:800;color:#fff;letter-spacing:-.4px;white-space:nowrap}}
.hdivider{{width:1px;height:22px;background:rgba(255,255,255,.14);flex-shrink:0}}
.hinfo{{display:flex;flex-direction:column;gap:1px;min-width:0}}
.htit{{color:#fff;font-size:13px;font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
.hsub{{color:rgba(255,255,255,.4);font-size:11px;white-space:nowrap}}
.kpis{{display:flex;gap:4px;margin-left:auto;flex-shrink:0}}
.kpi{{
  background:rgba(255,255,255,.05);
  border:1px solid rgba(255,255,255,.08);
  border-radius:var(--r);
  padding:6px 12px;text-align:center;min-width:62px;
  transition:background .15s;
}}
.kpi:hover{{background:rgba(255,255,255,.09)}}
.kn{{font-size:18px;font-weight:800;line-height:1;font-variant-numeric:tabular-nums;transition:color .3s}}
.kl{{font-size:9.5px;color:rgba(255,255,255,.35);text-transform:uppercase;letter-spacing:.5px;margin-top:2px}}
.k-tot .kn{{color:#93c5fd}}
.k-crit .kn{{color:#fca5a5}}
.k-med .kn{{color:#fde68a}}
.k-low .kn{{color:#86efac}}
.k-vis{{background:rgba(213,43,30,.12)!important;border-color:rgba(213,43,30,.25)!important}}
.k-vis .kn{{color:#fca5a5}}
.k-vis .kl{{color:rgba(252,165,165,.5)}}

/* ── RISK GAUGE (header) ────────────────────────────────────────────────── */
.rg-pill{{
  display:flex;align-items:center;gap:10px;
  background:rgba(255,255,255,.07);
  border:1px solid rgba(255,255,255,.12);
  border-radius:10px;padding:6px 14px;
  cursor:pointer;transition:background .15s;margin-left:auto;flex-shrink:0;
}}
.rg-pill:hover{{background:rgba(255,255,255,.12)}}
.rg-score{{font-size:28px;font-weight:900;line-height:1;font-variant-numeric:tabular-nums;min-width:44px;text-align:center}}
.rg-right{{display:flex;flex-direction:column;gap:1px}}
.rg-label{{font-size:9px;font-weight:700;text-transform:uppercase;letter-spacing:.6px;color:rgba(255,255,255,.4)}}
.rg-nivel{{font-size:11.5px;font-weight:700;line-height:1.1}}
.rg-trend{{font-size:10.5px;font-weight:600;line-height:1.1}}
.rg-better{{color:#86efac}}
.rg-worse{{color:#fca5a5}}
.rg-same{{color:rgba(255,255,255,.35)}}

/* ── TOP 3 INCIDENTS ────────────────────────────────────────────────────── */
.t3-label{{font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:.6px;color:var(--gray);margin-bottom:8px}}
.t3-row{{display:flex;gap:10px;margin-bottom:16px;flex-wrap:wrap}}
.t3-card{{
  flex:1;min-width:220px;max-width:380px;
  border:1px solid var(--border);border-radius:var(--r2);
  padding:12px 14px;background:var(--panel);
  border-left:4px solid var(--border);
  box-shadow:var(--sh);transition:box-shadow .15s;
}}
.t3-card:hover{{box-shadow:var(--sh2)}}
.t3-card.r-crit{{border-left-color:var(--crit);background:var(--crit-bg)}}
.t3-card.r-high{{border-left-color:var(--high);background:var(--high-bg)}}
.t3-card.r-med{{border-left-color:var(--med);background:var(--med-bg)}}
.t3-hdr{{display:flex;align-items:center;gap:6px;flex-wrap:wrap;margin-bottom:6px}}
.t3-ico{{font-size:16px;flex-shrink:0}}
.t3-cat{{font-size:11.5px;font-weight:700;color:var(--dark2)}}
.t3-date{{font-size:10.5px;color:var(--gray);margin-left:auto}}
.t3-wb{{font-size:11px}}
.t3-event{{font-size:12.5px;color:var(--dark);margin-bottom:4px;line-height:1.4;font-weight:500}}
.t3-asset{{font-size:11px;color:var(--gray)}}
.t3-id{{font-size:10px;color:var(--dgray)}}
.sec-empty-ok{{color:var(--low);font-size:13px;padding:14px 0;font-weight:600}}

/* ── CREM PANELS ────────────────────────────────────────────────────────── */
.crem-summary{{display:flex;gap:10px;flex-wrap:wrap;margin-bottom:12px}}
.cs-item{{
  display:flex;align-items:center;gap:6px;
  background:var(--panel);border:1px solid var(--border);
  border-radius:var(--r2);padding:7px 12px;font-size:12px;
}}
.cs-ico{{font-size:13px}}
.cs-lbl{{color:var(--dgray)}}
.cs-val{{font-weight:800;font-size:14px;margin-left:2px}}
.crem-row{{display:flex;gap:12px;flex-wrap:wrap;align-items:flex-start}}
.cp-panel{{
  flex:1;min-width:160px;
  background:var(--panel);border:1px solid var(--border);
  border-radius:var(--r2);overflow:hidden;box-shadow:var(--sh);
  transition:box-shadow .15s;
}}
.cp-panel:hover{{box-shadow:var(--sh2)}}
.cp-hdr{{
  display:flex;align-items:center;gap:7px;
  padding:10px 12px;
  background:var(--lgray);border-bottom:1px solid var(--border);
}}
.cp-ico{{font-size:14px;flex-shrink:0}}
.cp-title{{font-size:12px;font-weight:700;color:var(--dark2);flex:1}}
.cp-link{{font-size:11px;color:var(--accent);cursor:pointer;text-decoration:none;flex-shrink:0;padding:1px 6px;border-radius:4px;border:1px solid var(--accent);opacity:.7}}
.cp-link:hover{{opacity:1;background:var(--accent-bg)}}
.cp-body{{padding:10px 12px;display:flex;flex-direction:column;gap:8px;min-height:80px}}
.cp-item{{padding:7px 9px;border-radius:var(--r);background:var(--lgray);border:1px solid var(--border)}}
.cp-item:hover{{border-color:var(--mgray);background:#f0f1f4}}
.cp-name{{font-size:12px;font-weight:600;color:var(--dark2);margin-bottom:3px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
.cp-namelink{{color:var(--dark2);text-decoration:none;border-bottom:1px dashed transparent;transition:color .12s,border-color .12s}}
.cp-namelink:hover{{color:var(--accent);border-bottom-color:var(--accent)}}
.cp-meta{{display:flex;align-items:center;gap:5px;flex-wrap:wrap}}
.cp-tag{{font-size:10.5px;color:var(--gray);background:var(--panel);border:1px solid var(--border);border-radius:4px;padding:1px 5px}}
.cp-os{{font-size:10.5px;color:var(--gray);max-width:100px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}}
.cp-empty{{color:var(--dgray);font-size:11.5px;padding:10px 0;text-align:center}}

/* ── TOOLBAR ─────────────────────────────────────────────────────────────── */
.tb{{
  background:var(--panel);
  border-bottom:1px solid var(--border);
  padding:0 20px;
  display:flex;gap:6px;align-items:center;flex-wrap:wrap;
  min-height:var(--tb-h);
  box-shadow:var(--sh);
  position:sticky;top:var(--hdr-h);z-index:100;
}}
.tbg{{display:flex;gap:4px;align-items:center}}
.btn{{
  display:inline-flex;align-items:center;gap:5px;
  background:var(--lgray);color:var(--dark3);
  border:1px solid var(--border);border-radius:var(--r);
  padding:5px 12px;cursor:pointer;
  font-size:12px;font-family:var(--font);font-weight:600;
  transition:all .15s;white-space:nowrap;
}}
.btn:hover{{background:#fff;color:var(--red);border-color:var(--crit-border);box-shadow:0 1px 4px rgba(0,0,0,.06)}}
.btn .ico{{width:13px;height:13px}}
.tbsep{{width:1px;height:20px;background:var(--mgray);flex-shrink:0}}
.flbl{{font-size:11.5px;color:var(--gray);font-weight:600;white-space:nowrap}}
/* Severity chips */
.chip{{
  display:inline-flex;align-items:center;gap:5px;
  background:var(--lgray);border:1px solid var(--border);
  border-radius:20px;padding:4px 10px;
  font-size:11px;cursor:pointer;font-weight:500;
  transition:all .15s;white-space:nowrap;user-select:none;
}}
.chip:hover{{border-color:var(--red);color:var(--red);background:var(--red-a)}}
.chip.on{{background:var(--red);color:#fff;border-color:var(--red);box-shadow:0 0 0 2px rgba(213,43,30,.15)}}
.cdot{{width:5px;height:5px;border-radius:50%;background:currentColor;flex-shrink:0}}
/* Inv chips */
.inv-sep{{width:1px;height:20px;background:var(--mgray);flex-shrink:0}}
.inv-chip{{
  display:inline-flex;align-items:center;gap:5px;
  background:var(--lgray);border:1px solid var(--border);
  border-radius:20px;padding:4px 10px;
  font-size:11px;cursor:pointer;font-weight:500;
  transition:all .15s;white-space:nowrap;user-select:none;
}}
.inv-chip:hover{{border-color:#7f1d1d;color:#7f1d1d;background:#fff1f1}}
.inv-chip.on{{background:#7f1d1d;color:#fff;border-color:#7f1d1d;box-shadow:0 0 0 2px rgba(127,29,29,.15)}}
/* Filter info badge */
.flt-info{{
  display:none;align-items:center;gap:5px;
  background:rgba(213,43,30,.07);border:1px solid rgba(213,43,30,.2);
  border-radius:20px;padding:3px 10px;font-size:11px;color:var(--red);font-weight:600;
}}
.flt-info.on{{display:inline-flex}}
.flt-clear{{background:none;border:none;cursor:pointer;color:var(--red);font-size:12px;padding:0;margin-left:2px;opacity:.7}}
.flt-clear:hover{{opacity:1}}

/* ── SEARCH ──────────────────────────────────────────────────────────────── */
.sa{{margin-left:auto;display:flex;align-items:center;gap:6px}}
.sw{{position:relative}}
.sw input{{
  width:240px;padding:6px 12px 6px 34px;
  border:1px solid var(--border);border-radius:var(--r);
  font-size:13px;font-family:var(--font);outline:none;
  background:var(--lgray);
  transition:all .15s;
}}
.sw input:focus{{
  border-color:var(--accent);background:#fff;
  box-shadow:0 0 0 3px rgba(30,64,175,.1);
}}
.si{{position:absolute;left:10px;top:50%;transform:translateY(-50%);color:var(--dgray);font-size:13px;pointer-events:none}}
.sc{{cursor:pointer;color:var(--dgray);padding:3px 6px;border-radius:4px;border:none;background:none;font-size:13px}}
.sc:hover{{color:var(--red);background:var(--red-a)}}

/* ── SEARCH PANEL ────────────────────────────────────────────────────────── */
.srp{{
  display:none;
  position:sticky;top:calc(var(--hdr-h) + var(--tb-h));z-index:90;
  background:var(--panel);border-bottom:1px solid var(--border);
  box-shadow:var(--sh2);max-height:280px;overflow-y:auto;
}}
.srp-h{{
  padding:8px 20px;background:var(--lgray);
  border-bottom:1px solid var(--border);
  font-size:11.5px;font-weight:600;color:var(--gray);
  display:flex;justify-content:space-between;align-items:center;
  position:sticky;top:0;z-index:1;
}}
.sri{{
  padding:8px 20px;border-bottom:1px solid var(--border);
  cursor:pointer;transition:background .1s;
  display:flex;align-items:flex-start;gap:10px;
}}
.sri:hover{{background:var(--lgray)}}.sri:last-child{{border-bottom:none}}
.srlbl{{font-size:10px;background:var(--accent);color:#fff;border-radius:4px;padding:1px 7px;white-space:nowrap;margin-top:2px;flex-shrink:0}}
.srtxt{{font-size:12.5px;color:var(--dark);line-height:1.4}}
.srtxt mark{{background:#fef9c3;padding:0 2px;border-radius:2px;color:var(--dark)}}
.srnone{{padding:16px;text-align:center;color:var(--gray);font-size:13px}}

/* ── LAYOUT ──────────────────────────────────────────────────────────────── */
.lay{{
  display:flex;
  height:calc(100vh - var(--hdr-h) - var(--tb-h));
  overflow:hidden;
}}
.nav{{
  width:var(--nav-w);flex-shrink:0;
  background:var(--panel);border-right:1px solid var(--border);
  overflow-y:auto;overflow-x:hidden;
  padding:8px 0 16px;
  scrollbar-width:thin;scrollbar-color:var(--mgray) transparent;
}}
.nav::-webkit-scrollbar{{width:4px}}.nav::-webkit-scrollbar-thumb{{background:var(--mgray);border-radius:2px}}
.nav-backdrop{{display:none}}
.navt{{
  font-size:10px;font-weight:700;color:var(--dgray);
  text-transform:uppercase;letter-spacing:.8px;
  padding:10px 16px 4px;
}}
.ni{{
  display:flex;align-items:center;gap:8px;
  padding:7px 16px;cursor:pointer;
  font-size:12.5px;color:var(--gray);
  border-left:2px solid transparent;
  transition:all .12s;border-radius:0;
  position:relative;
}}
.ni:hover{{background:var(--lgray);color:var(--dark2)}}
.ni.on{{
  background:linear-gradient(90deg,rgba(213,43,30,.07),transparent);
  color:var(--red);border-left-color:var(--red);font-weight:600;
}}
.nico{{width:18px;flex-shrink:0}}
.ncnt{{
  margin-left:auto;font-size:11px;font-weight:600;
  background:var(--lgray);border:1px solid var(--border);
  border-radius:10px;padding:0px 7px;color:var(--gray);
  transition:all .15s;min-width:24px;text-align:center;
  font-variant-numeric:tabular-nums;
}}
.ni.on .ncnt{{background:var(--red);color:#fff;border-color:var(--red)}}
.ncnt.changed{{animation:countpop .3s ease}}
@keyframes countpop{{0%{{transform:scale(1)}}50%{{transform:scale(1.3)}}100%{{transform:scale(1)}}}}
.main{{flex:1;overflow-y:auto;padding:20px 24px;scrollbar-width:thin;scrollbar-color:var(--mgray) transparent}}
.main::-webkit-scrollbar{{width:6px}}.main::-webkit-scrollbar-thumb{{background:var(--mgray);border-radius:3px}}

/* ── SECTIONS ────────────────────────────────────────────────────────────── */
.sec{{
  background:var(--panel);border:1px solid var(--border);
  border-radius:var(--r2);margin-bottom:10px;
  box-shadow:var(--sh);overflow:hidden;
  transition:box-shadow .15s;
}}
.sec:focus-within{{box-shadow:var(--sh2)}}
.sec summary{{
  list-style:none;display:flex;align-items:center;
  justify-content:space-between;padding:13px 18px;
  cursor:pointer;user-select:none;transition:background .12s;
  gap:12px;
}}
.sec summary::-webkit-details-marker{{display:none}}
.sec summary:hover{{background:var(--lgray)}}
.sec[open]>summary{{border-bottom:1px solid var(--border);background:var(--lgray)}}
.sec[open] .sarr{{transform:rotate(90deg)}}
.sl{{display:flex;align-items:center;gap:10px;min-width:0}}
.sarr{{
  color:var(--dgray);font-size:9px;
  transition:transform .2s cubic-bezier(.4,0,.2,1);flex-shrink:0;
}}
.stit{{font-size:13.5px;font-weight:600;color:var(--dark2);white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
.scnt{{
  font-size:11.5px;font-weight:700;
  background:var(--lgray);color:var(--gray);
  border:1px solid var(--border);border-radius:20px;
  padding:2px 10px;flex-shrink:0;white-space:nowrap;
  font-variant-numeric:tabular-nums;transition:all .2s;
}}
.sec[open]>.summary>.scnt,
.sec[open] summary .scnt{{background:var(--red);color:#fff;border-color:var(--red)}}
.sbody{{padding:16px 18px;overflow-x:auto}}

/* ── CHARTS ──────────────────────────────────────────────────────────────── */
.charts-row{{display:flex;gap:14px;margin-bottom:16px;flex-wrap:wrap}}
.chart-box{{
  background:var(--lgray);border:1px solid var(--border);
  border-radius:var(--r);padding:14px;flex:0 0 auto;
}}
.chart-box-wide{{flex:1;min-width:280px}}
.chart-title{{font-size:11px;font-weight:700;color:var(--gray);margin-bottom:10px;text-transform:uppercase;letter-spacing:.5px}}
.chart-svg{{width:100%;overflow-x:auto}}
.chart-svg svg{{display:block;max-width:100%}}
.chart-legend{{display:flex;gap:16px;flex-wrap:wrap;justify-content:center;margin-top:10px;
  font-size:11px;color:var(--gray);font-weight:600}}
.chart-legend span{{display:inline-flex;align-items:center;gap:6px}}
.chart-legend i{{width:11px;height:11px;border-radius:3px;display:inline-block;flex:none}}
.donut-wrap{{display:flex;flex-direction:column;align-items:center;gap:4px}}
.donut-legend{{flex-direction:column;gap:8px;align-items:stretch;justify-content:flex-start;
  margin-top:12px;width:100%}}
.donut-legend span{{display:flex;align-items:center;gap:8px;font-size:12.5px;color:var(--dark3);font-weight:600}}
.donut-legend b{{margin-left:auto;color:var(--dark2);font-variant-numeric:tabular-nums;font-size:13px}}
.donut-legend small{{color:var(--gray);font-weight:600;font-variant-numeric:tabular-nums}}

/* ── COMPARATIVA DE CAMBIOS ──────────────────────────────────────────────── */
.chg-note{{display:flex;align-items:center;gap:8px;font-size:12px;color:var(--gray);
  background:var(--accent-bg);border:1px solid rgba(30,64,175,.13);border-radius:8px;
  padding:9px 12px;margin-bottom:12px;flex-wrap:wrap}}
.chg-note .ico{{color:var(--accent);flex:none}}
.chg-up-t{{color:var(--crit);font-weight:700}}
.chg-down-t{{color:var(--low);font-weight:700}}
.chg-tbl td{{vertical-align:middle}}
.chg-mod{{font-weight:600;color:var(--dark2)}}
.chg-num{{font-variant-numeric:tabular-nums;color:var(--gray);text-align:center}}
.chg-actn{{color:var(--dark2);font-weight:700}}
.chg-pill{{display:inline-flex;align-items:center;gap:4px;font-size:11.5px;font-weight:800;
  padding:3px 10px;border-radius:20px;white-space:nowrap;font-variant-numeric:tabular-nums}}
.chg-pill .ico{{width:13px;height:13px}}
.chg-up{{color:var(--crit);background:var(--crit-bg);border:1px solid var(--crit-border)}}
.chg-down{{color:var(--low);background:var(--low-bg);border:1px solid var(--low-border)}}
.chg-flat{{color:var(--gray);background:var(--lgray);border:1px solid var(--border)}}
.chg-na{{color:var(--dgray);background:var(--lgray);border:1px dashed var(--mgray);font-weight:600}}
.chg-crit{{font-variant-numeric:tabular-nums;color:var(--gray)}}
.chg-crit b{{color:var(--dark2)}}
.chg-cd{{font-size:10.5px;font-weight:800;padding:1px 6px;border-radius:10px;margin-left:4px}}
.chg-cd.chg-up{{color:var(--crit);background:var(--crit-bg)}}
.chg-cd.chg-down{{color:var(--low);background:var(--low-bg)}}
.chg-cd.chg-flat{{color:var(--gray);background:var(--lgray)}}
.chg-new{{color:var(--new);font-weight:700;text-align:center;font-variant-numeric:tabular-nums}}
.chg-res{{color:var(--low);font-weight:700;text-align:center;font-variant-numeric:tabular-nums}}

/* ── FRANJA RIESGO CREM (tendencia) ──────────────────────────────────────── */
.tr-wrap{{background:var(--lgray);border:1px solid var(--border);border-radius:var(--r);
  padding:14px 16px;margin-bottom:12px}}
.tr-title{{display:flex;align-items:center;gap:7px;font-size:12px;font-weight:700;color:var(--gray);
  text-transform:uppercase;letter-spacing:.5px;margin-bottom:12px}}
.tr-title .ico{{color:var(--red);width:15px;height:15px}}
.tr-strip{{display:flex;gap:8px;overflow-x:auto;padding-bottom:2px}}
.tr-cell{{flex:1;min-width:78px;background:var(--panel);border:1px solid var(--border);
  border-radius:10px;padding:10px 8px;text-align:center;position:relative}}
.tr-mes{{font-size:10.5px;color:var(--gray);font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;margin-bottom:6px}}
.tr-score{{font-size:22px;font-weight:800;line-height:1;font-variant-numeric:tabular-nums}}
.tr-none{{color:var(--dgray)}}
.sc-crit{{color:var(--crit)}}.sc-high{{color:var(--high)}}.sc-med{{color:#c79100}}.sc-low{{color:var(--low)}}
.tr-d{{display:inline-flex;align-items:center;gap:2px;margin-top:6px;font-size:11px;font-weight:800;
  padding:1px 7px;border-radius:10px;font-variant-numeric:tabular-nums}}
.tr-d .ico{{width:12px;height:12px}}
.tr-up{{color:var(--crit);background:var(--crit-bg)}}
.tr-down{{color:var(--low);background:var(--low-bg)}}
.tr-flat{{color:var(--gray);background:var(--lgray)}}

/* ── TABLES ──────────────────────────────────────────────────────────────── */
.tbl-hdr{{
  display:flex;align-items:center;justify-content:space-between;
  margin-bottom:8px;min-height:28px;gap:8px;flex-wrap:wrap;
}}
.tbl-info{{font-size:11.5px;color:var(--gray)}}
.tbl-actions{{display:flex;gap:6px;align-items:center;margin-left:auto}}
.tbl-filter{{
  padding:4px 10px 4px 28px;border:1px solid var(--border);border-radius:var(--r);
  font-size:12px;font-family:var(--font);outline:none;background:var(--lgray);
  transition:all .15s;width:180px;
}}
.tbl-filter:focus{{border-color:var(--accent);background:#fff;box-shadow:0 0 0 2px rgba(30,64,175,.1)}}
.tbl-filter-wrap{{position:relative;display:flex;align-items:center}}
.tbl-filter-ico{{position:absolute;left:8px;font-size:11px;color:var(--dgray);pointer-events:none}}
.btn-exp{{
  background:var(--lgray);color:var(--gray);
  border:1px solid var(--border);border-radius:var(--r);
  padding:4px 10px;cursor:pointer;font-size:11.5px;font-weight:500;
  display:inline-flex;align-items:center;gap:4px;transition:all .15s;
}}
.btn-exp:hover{{background:var(--dark);color:#fff;border-color:var(--dark)}}
.tbl-wrap{{
  overflow-x:auto;border:1px solid var(--border);
  border-radius:var(--r);
  scrollbar-width:thin;scrollbar-color:var(--mgray) transparent;
}}
/* Tablas virtuales: necesitan scroll VERTICAL propio, porque el cargador
   incremental de filas se dispara con el scroll de este contenedor. Sin
   max-height el div nunca desborda, el evento no salta nunca y la tabla se
   queda congelada en las primeras 200 filas. */
.tbl-wrap.vt{{max-height:70vh;overflow-y:auto}}
.tbl-wrap::-webkit-scrollbar{{height:5px;width:5px}}.tbl-wrap::-webkit-scrollbar-thumb{{background:var(--mgray);border-radius:3px}}
.dtbl{{border-collapse:collapse;width:100%;min-width:480px;font-size:12.5px}}
.dtbl thead{{position:sticky;top:0;z-index:2}}
.dtbl th{{
  background:var(--dark2);color:rgba(255,255,255,.8);
  padding:9px 12px;text-align:left;
  font-size:10.5px;font-weight:700;text-transform:uppercase;letter-spacing:.4px;
  white-space:nowrap;cursor:pointer;user-select:none;
  border-bottom:2px solid var(--red);
  transition:background .12s;
}}
.dtbl th:hover{{background:var(--dark3);color:#fff}}
.dtbl th:first-child{{border-radius:var(--r) 0 0 0}}
.dtbl th:last-child{{border-radius:0 var(--r) 0 0}}
.dtbl th.sa-::after{{content:" ↑";color:var(--red)}}
.dtbl th.sd-::after{{content:" ↓";color:var(--red)}}
.dtbl td{{
  padding:8px 12px;
  border-bottom:1px solid var(--border);
  vertical-align:top;color:var(--dark);
}}
.dtbl tbody tr:last-child td{{border-bottom:none}}
.dtbl tbody tr{{transition:background .1s}}
.dtbl tbody tr:hover td{{background:rgba(30,64,175,.03)!important}}
.dtbl tbody tr.hl td{{background:#fef9c3!important;transition:none}}
.r-crit td{{background:rgba(198,40,40,.05)!important}}
.r-high td{{background:rgba(230,81,0,.04)!important}}
.r-med td{{background:rgba(245,127,23,.03)!important}}
.r-new td{{background:rgba(124,58,237,.05)!important}}
tr.fo{{display:none}}
.ml-cell{{white-space:pre-wrap;min-width:160px;line-height:1.5}}
/* Empty state */
.tbl-empty{{text-align:center;padding:32px;color:var(--dgray);font-size:13px}}
.sec-empty-ok{{text-align:center;padding:22px 16px;color:var(--low,#16a34a);background:var(--low-bg,#f0fdf4);border:1px solid var(--low,#16a34a);border-radius:8px;font-size:13px;font-weight:600}}

/* ── BADGES & PILLS ──────────────────────────────────────────────────────── */
.badge{{
  display:inline-block;padding:2px 8px;border-radius:4px;
  font-size:10.5px;font-weight:700;text-transform:uppercase;
  letter-spacing:.3px;border:1px solid transparent;
}}
.score-pill{{
  display:inline-block;border-radius:5px;padding:2px 8px;
  font-size:11.5px;font-weight:800;color:#fff;
  min-width:34px;text-align:center;letter-spacing:.2px;
}}
.s-crit{{background:var(--crit)}}
.s-high{{background:var(--high)}}
.s-med{{background:var(--med);color:var(--dark)}}
.s-low{{background:var(--low)}}
.sev-dot{{display:inline-flex;align-items:center;line-height:0;margin-right:2px}}
.sev-dot .ico{{width:.72em;height:.72em}}
.cs-item .sev-dot .ico,.inv-chip .sev-dot .ico{{width:.85em;height:.85em}}
.prio-pill{{
  display:inline-flex;align-items:center;gap:4px;
  padding:3px 10px;border-radius:20px;
  font-size:11.5px;font-weight:700;white-space:nowrap;border:1px solid transparent;
}}
.p-crit{{background:var(--crit-bg);color:var(--crit);border-color:var(--crit-border)}}
.p-high{{background:var(--high-bg);color:var(--high);border-color:var(--high-border)}}
.p-med{{background:var(--med-bg);color:var(--med);border-color:var(--med-border)}}
.p-low{{background:var(--low-bg);color:var(--low);border-color:var(--low-border)}}
.aref{{
  display:inline-flex;align-items:center;gap:2px;
  font-size:10px;font-weight:700;letter-spacing:.2px;
  background:var(--accent-bg);color:var(--accent);
  border:1px solid rgba(30,64,175,.22);border-radius:5px;
  padding:1px 7px;margin:1px;cursor:pointer;text-decoration:none;
  transition:transform .12s,box-shadow .12s,background .12s,color .12s;
  vertical-align:middle;
}}
.aref::after{{content:"→";font-size:9px;opacity:.55;transition:transform .12s,opacity .12s}}
.aref:hover{{background:var(--accent);color:#fff;text-decoration:none;
  box-shadow:0 2px 7px rgba(30,64,175,.28);transform:translateY(-1px)}}
.aref:hover::after{{opacity:1;transform:translateX(2px)}}
.aref:active{{transform:translateY(0)}}
.aref-hit{{background:var(--accent)!important;color:#fff!important;
  box-shadow:0 0 0 3px rgba(30,64,175,.25)!important}}
.ext{{color:var(--accent);font-weight:600;font-size:12px}}
.ext:hover{{color:var(--red);text-decoration:none}}
.ico{{display:inline-block;vertical-align:-.14em;flex:none}}
/* Iconos de cabecera de sección: gris pizarra, rojo marca al abrir */
.sico{{display:inline-flex;align-items:center;justify-content:center;width:26px;height:26px;
  border-radius:7px;background:var(--lgray);border:1px solid var(--border);transition:all .15s}}
.sico .ico{{width:15px;height:15px;color:#64748b;transition:color .15s}}
.sec[open]>summary .sico{{background:var(--red-a);border-color:var(--crit-border)}}
.sec[open]>summary .sico .ico{{color:var(--red)}}
.sec summary:hover .sico .ico{{color:var(--red)}}
/* Iconos de navegación lateral */
.nico{{display:inline-flex;align-items:center;justify-content:center}}
.nico .ico{{width:15px;height:15px;color:#94a3b8;transition:color .15s}}
.ni:hover .nico .ico{{color:var(--accent)}}
.ni.on .nico .ico{{color:var(--red)}}
.cp-ico .ico{{width:15px;height:15px;color:var(--red)}}
.t3-ico .ico{{color:var(--red)}}
.section-icon .ico,.diff-new .ico,.diff-res .ico{{color:currentColor}}
.new-badge{{
  display:inline-block;background:var(--new);color:#fff;
  font-size:9px;font-weight:700;border-radius:3px;
  padding:1px 5px;margin-right:4px;
  text-transform:uppercase;vertical-align:middle;letter-spacing:.3px;
}}
.modal-btn{{
  background:none;border:none;cursor:pointer;
  font-size:14px;padding:3px 5px;color:var(--dgray);line-height:0;
  border-radius:5px;transition:all .12s;vertical-align:middle;
}}
.modal-btn:hover{{color:var(--accent);background:var(--accent-bg)}}
.cve-btn{{color:#1e3a5f!important}}
.cve-btn:hover{{color:var(--red)!important;background:var(--red-a)!important}}
.sol-btn{{
  display:inline-flex;align-items:center;justify-content:center;
  background:var(--accent-bg);color:var(--accent);border:1px solid rgba(30,64,175,.25);
  border-radius:4px;padding:3px 5px;margin-left:6px;
  cursor:pointer;transition:all .15s;vertical-align:middle;
  line-height:0;box-shadow:0 1px 2px rgba(0,0,0,0.05);
}}
.sol-btn:hover{{
  background:var(--accent);color:#fff;border-color:var(--accent);
  transform:scale(1.08);box-shadow:0 2px 4px rgba(30,64,175,0.2);
}}
.sol-btn .ico{{
  width:12px;height:12px;display:block;
}}
/* ── CVE MODAL ── */
.cve-meta{{display:flex;flex-wrap:wrap;gap:8px;padding:6px 0 10px}}
.cve-meta-pill{{font-size:11.5px;font-weight:600;color:var(--gray);background:var(--lgray);
  border:1px solid var(--border);border-radius:20px;padding:3px 11px;display:inline-flex;align-items:center;gap:5px}}
.cve-note{{font-size:12px;color:var(--gray);background:var(--accent-bg);border:1px solid rgba(30,64,175,.15);
  border-radius:7px;padding:8px 12px;margin-bottom:10px;display:flex;align-items:center;gap:7px}}
.cve-filter{{width:100%;padding:7px 12px;border:1px solid var(--border);border-radius:7px;
  font-family:var(--font);font-size:13px;outline:none;margin-bottom:10px;background:var(--lgray)}}
.cve-filter:focus{{border-color:var(--accent);background:#fff;box-shadow:0 0 0 2px rgba(30,64,175,.1)}}
.cve-empty{{padding:24px;text-align:center;color:var(--dark2);font-size:13.5px;line-height:1.7}}

/* ── INVENTORY BADGES ────────────────────────────────────────────────────── */
.inv-badge{{
  display:inline-flex;align-items:center;gap:3px;
  font-size:9.5px;font-weight:700;
  border-radius:4px;padding:1px 6px;
  white-space:nowrap;margin-left:4px;vertical-align:middle;
  text-transform:uppercase;letter-spacing:.3px;
  border:1px solid transparent;cursor:help;
  transition:opacity .15s;
}}
.inv-badge:hover{{opacity:.7}}
.inv-muy-crit{{background:#fee2e2;color:#991b1b;border-color:#fca5a5}}
.inv-crit-inv{{background:#fff0f0;color:#b91c1c;border-color:#fecaca}}
.inv-normal-inv{{background:#fff3e0;color:#e65100;border-color:#ffcc80}}
.inv-no-crit-inv{{background:var(--low-bg);color:var(--low);border-color:var(--low-border)}}
.muy-crit-row>td{{background:rgba(220,38,38,.04)!important}}
.crit-inv-row>td{{background:rgba(220,38,38,.02)!important}}
.normal-inv-row>td{{background:rgba(230,81,0,.02)!important}}
/* Sidebar inventory */
.nav-inv-hdr{{
  font-size:9.5px;font-weight:700;color:var(--dgray);
  text-transform:uppercase;letter-spacing:.8px;
  padding:12px 16px 4px;margin-top:4px;
  border-top:1px solid var(--border);
}}
.ni-inv{{
  display:flex;align-items:center;gap:7px;
  padding:4px 16px;font-size:11.5px;
  cursor:default;border-left:2px solid transparent;
  transition:background .1s;
}}
.ni-inv:hover{{background:var(--lgray)}}
.ni-inv-name{{font-weight:600;flex:0 0 auto;max-width:90px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}}
.ni-inv-desc{{color:var(--dgray);font-size:10px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;flex:1}}

/* ── BANNERS & NOTES ─────────────────────────────────────────────────────── */
.diff-banner{{
  display:flex;gap:14px;align-items:center;flex-wrap:wrap;
  background:linear-gradient(135deg,#fdf4ff,#f3e8ff);
  border:1px solid #e9d5ff;border-radius:var(--r);
  padding:9px 14px;margin-bottom:12px;font-size:12.5px;
}}
.diff-new{{color:var(--new);font-weight:600}}
.diff-res{{color:var(--low);font-weight:600}}
.diff-src{{color:var(--gray);font-size:11px}}
.note-info{{
  background:linear-gradient(135deg,#fdf4ff,#f3e8ff);
  border:1px solid #e9d5ff;border-radius:var(--r);
  padding:8px 13px;font-size:12px;margin-bottom:8px;
}}
.plan-leg{{
  display:flex;gap:10px;align-items:center;flex-wrap:wrap;
  padding:10px 14px;background:var(--lgray);border:1px solid var(--border);
  border-radius:var(--r);margin-bottom:12px;font-size:12px;color:var(--gray);
}}
.note{{
  background:#fefce8;border:1px solid #fef08a;
  border-radius:var(--r);padding:10px 14px;
  font-size:12.5px;color:#713f12;margin-bottom:8px;
}}

/* ── MODAL ───────────────────────────────────────────────────────────────── */
.modal-overlay{{display:none;position:fixed;inset:0;background:rgba(0,0,0,.45);z-index:500;align-items:center;justify-content:center;backdrop-filter:blur(2px)}}
.modal-overlay.on{{display:flex}}
.modal{{
  background:var(--panel);border-radius:var(--r2);
  box-shadow:var(--sh3);width:92%;max-width:700px;
  max-height:88vh;overflow:hidden;display:flex;flex-direction:column;
  animation:modalIn .2s cubic-bezier(.4,0,.2,1);
}}
@keyframes modalIn{{from{{opacity:0;transform:translateY(-12px) scale(.97)}}to{{opacity:1;transform:none}}}}
.modal-hdr{{
  display:flex;align-items:center;justify-content:space-between;
  padding:16px 20px;border-bottom:1px solid var(--border);
  background:var(--dark2);color:#fff;
}}
.modal-title{{font-size:15px;font-weight:700}}
.modal-close{{background:none;border:none;color:rgba(255,255,255,.6);font-size:20px;cursor:pointer;padding:2px 6px;border-radius:4px;transition:all .12s}}
.modal-close:hover{{color:#fff;background:rgba(255,255,255,.1)}}
.modal-body{{overflow-y:auto;padding:18px 20px;scrollbar-width:thin;scrollbar-color:var(--mgray) transparent}}
.modal-sec{{margin-bottom:16px}}
.modal-sec-title{{font-size:11px;font-weight:700;color:var(--dgray);text-transform:uppercase;letter-spacing:.6px;margin-bottom:10px;padding-bottom:6px;border-bottom:1px solid var(--border)}}
.modal-item{{display:flex;gap:10px;padding:8px 0;border-bottom:1px solid var(--border);font-size:12.5px}}
.modal-item:last-child{{border-bottom:none}}
.modal-mod{{background:var(--lgray);border:1px solid var(--border);border-radius:4px;padding:2px 8px;font-size:10.5px;font-weight:700;color:var(--gray);white-space:nowrap;flex-shrink:0;height:fit-content;margin-top:1px}}
.modal-prob{{flex:1;color:var(--dark)}}
.modal-accion{{color:var(--gray);font-size:11.5px;margin-top:3px;padding-left:8px;border-left:2px solid var(--accent)}}
.modal-xref{{display:flex;flex-wrap:wrap;gap:6px;padding:2px 0}}
.modal-xref .aref{{font-size:11px;padding:3px 10px}}
.aref-lg{{font-size:12px!important;padding:6px 12px!important;font-weight:700}}
.modal-prob a.ext,.modal-accion a.ext{{white-space:nowrap}}

/* ── ANIMATIONS ──────────────────────────────────────────────────────────── */
/* El blink se pinta con box-shadow inset (overlay) porque los fondos de fila
   usan background:!important y, en CSS, !important gana a las animaciones —
   un overlay de sombra se dibuja por encima del fondo y siempre es visible. */
@keyframes blinkrow{{
  0%,100%{{box-shadow:inset 0 0 0 9999px rgba(245,158,11,0)}}
  12%,48%,84%{{box-shadow:inset 0 0 0 9999px rgba(245,158,11,.52)}}
  30%,66%{{box-shadow:inset 0 0 0 9999px rgba(245,158,11,.16)}}
}}
.blink-hl td{{animation:blinkrow 2.6s ease-in-out both!important}}
.blink-hl{{outline:2px solid #f59e0b;outline-offset:-2px;border-radius:2px}}
.blink-hl td:first-child{{position:relative}}
.blink-hl td:first-child::before{{content:"";position:absolute;left:0;top:0;bottom:0;
  width:4px;background:#f59e0b;z-index:1}}
@keyframes fadeIn{{from{{opacity:0;transform:translateY(4px)}}to{{opacity:1;transform:none}}}}

/* ── SCROLL TOP ──────────────────────────────────────────────────────────── */
.stb{{
  position:fixed;bottom:24px;right:24px;
  background:var(--dark2);color:#fff;
  border:1px solid rgba(255,255,255,.1);
  border-radius:50%;width:40px;height:40px;font-size:16px;
  cursor:pointer;box-shadow:var(--sh2);
  display:none;align-items:center;justify-content:center;
  z-index:50;transition:all .15s;
}}
.stb:hover{{background:var(--red);transform:translateY(-2px);box-shadow:var(--sh3)}}
.stb.on{{display:flex}}

/* ── LOADING SCREEN ──────────────────────────────────────────────────────── */
#loading-screen{{
  position:fixed;inset:0;
  background:linear-gradient(135deg,#0f1117 0%,#1a1d27 100%);
  z-index:9999;display:flex;flex-direction:column;
  align-items:center;justify-content:center;gap:20px;
  transition:opacity .5s cubic-bezier(.4,0,.2,1);
}}
#loading-screen.hidden{{opacity:0;pointer-events:none}}
.ld-logo{{display:flex;align-items:center;gap:12px;font-size:32px;font-weight:900;color:#fff;letter-spacing:-1px}}
.ld-mark{{width:44px;height:44px;border-radius:12px;background:linear-gradient(135deg,var(--red),var(--red2));display:flex;align-items:center;justify-content:center;box-shadow:0 4px 16px rgba(213,43,30,.45),inset 0 1px 0 rgba(255,255,255,.2)}}
.ld-mark .ico{{width:24px;height:24px;color:#fff}}
.ld-word{{color:#fff}}
.ld-sub{{color:rgba(255,255,255,.3);font-size:12px;letter-spacing:.5px;margin-top:-10px}}
.ld-pct{{color:#fff;font-size:40px;font-weight:800;font-variant-numeric:tabular-nums;letter-spacing:-1px}}
.ld-bar-wrap{{width:300px;height:3px;background:rgba(255,255,255,.08);border-radius:2px;overflow:hidden}}
.ld-bar{{height:100%;width:0%;background:linear-gradient(90deg,var(--red),#ff6b6b);border-radius:2px;transition:width .25s ease}}
.ld-step{{color:rgba(255,255,255,.3);font-size:11px;min-height:16px;letter-spacing:.3px}}

/* ── RESPONSIVE ──────────────────────────────────────────────────────────── */
@media(max-width:1024px){{
  :root{{--nav-w:190px}}
}}
@media(max-width:900px){{
  :root{{--nav-w:0px}}
  .nav-toggle-btn{{display:flex}}
  .nav{{
    display:block;position:fixed;top:var(--hdr-h);left:0;bottom:0;
    width:260px;max-width:80vw;z-index:210;
    transform:translateX(-100%);transition:transform .22s ease;
    box-shadow:var(--sh3);
  }}
  .lay.nav-open .nav{{transform:translateX(0)}}
  .nav-backdrop{{
    display:none;position:fixed;top:var(--hdr-h);left:0;right:0;bottom:0;
    background:rgba(0,0,0,.45);z-index:205;
  }}
  .lay.nav-open .nav-backdrop{{display:block}}
  .kpis .kpi:nth-child(n+4){{display:none}}
  .hdr-in{{gap:10px;padding:0 12px}}
  .hsub{{display:none}}
}}
@media(max-width:640px){{
  .main{{padding:12px}}
  .sw input{{width:160px}}
  .kpis{{display:none}}
  .charts-row{{flex-direction:column}}
  .tb{{padding:0 12px;gap:4px}}
  .flbl{{display:none}}
  .inv-sep{{display:none}}
  .rg-score{{font-size:22px;min-width:34px}}
  .rg-pill{{padding:5px 10px;gap:7px}}
}}
@media(max-width:480px){{
  .hdr-in{{gap:6px}}
  .htit{{font-size:12px}}
  .hdivider{{display:none}}
  .crem-row{{flex-direction:column}}
  .cp-panel{{min-width:0}}
  .t3-row{{flex-direction:column}}
  .t3-card{{min-width:0;max-width:none}}
  .rg-label,.rg-trend{{display:none}}
  .sa{{margin-left:0;width:100%}}
  .sw input{{width:100%}}
}}
{CSS_DEGRADADO}
"""


def _css_ejecutivo() -> str:
    """Hoja de estilos del informe ejecutivo (vista clara, sin navegación)."""
    return f"""{CSS_RESET}
:root{{{CSS_TOKENS}
  /* Propios del informe ejecutivo (superficies y texto de la vista clara) */
  --red-s:#fff0ee;
  --s1:#ffffff;--s2:#f8f9fb;--s3:#f0f1f4;
  --bd:#e5e7eb;--bd2:#d8dce3;
  --t1:#1a1d27;--t2:#4b5563;--t3:#9ca3af;
  --green:#2e7d32;--green-s:#f1f8e9;
  --amber:#e65100;--amber-s:#fff3e0;
  --blue:#1e40af;
}}
body{{background:var(--bg);color:var(--t1);font-family:var(--font);font-size:13.5px;line-height:1.55;-webkit-font-smoothing:antialiased;padding:0 0 48px}}
a{{color:var(--accent);text-decoration:none}}a:hover{{text-decoration:underline}}

/* ── HEADER (estética informe técnico) ── */
.hdr{{background:linear-gradient(180deg,#20232f 0%,#171922 100%);
  border-bottom:1px solid rgba(255,255,255,.07);position:sticky;top:0;z-index:200;
  box-shadow:0 1px 0 rgba(255,255,255,.04),0 6px 24px rgba(0,0,0,.28);min-height:var(--hdr-h)}}
.hdr-in{{display:flex;align-items:center;gap:16px;padding:9px 24px;min-height:var(--hdr-h);flex-wrap:wrap}}
.logo{{display:flex;align-items:center;gap:10px;flex-shrink:0}}
.logo-mark{{width:30px;height:30px;border-radius:8px;flex-shrink:0;
  background:linear-gradient(135deg,var(--red),var(--red2));
  display:flex;align-items:center;justify-content:center;color:#fff;
  box-shadow:0 2px 8px rgba(213,43,30,.4),inset 0 1px 0 rgba(255,255,255,.18)}}
.logo-mark .ico{{width:17px;height:17px}}
.logo-word{{font-size:18.5px;font-weight:800;color:#fff;letter-spacing:-.4px;white-space:nowrap}}
.eje-badge{{display:inline-block;background:var(--red);color:#fff;font-size:9.5px;font-weight:700;
  padding:2px 8px;border-radius:10px;letter-spacing:.6px;text-transform:uppercase;margin-left:2px;vertical-align:middle}}
.hdivider{{width:1px;height:22px;background:rgba(255,255,255,.14);flex-shrink:0}}
.hinfo{{display:flex;flex-direction:column;gap:1px;min-width:0}}
.htit{{color:#fff;font-size:13px;font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}}
.hsub{{color:rgba(255,255,255,.4);font-size:11px;white-space:nowrap}}
/* Risk gauge pill (header) */
.rg-pill{{display:flex;align-items:center;gap:10px;background:rgba(255,255,255,.07);
  border:1px solid rgba(255,255,255,.12);border-radius:10px;padding:6px 14px;margin-left:auto;flex-shrink:0}}
.rg-score{{font-size:28px;font-weight:900;line-height:1;font-variant-numeric:tabular-nums;min-width:44px;text-align:center}}
.rg-right{{display:flex;flex-direction:column;gap:1px}}
.rg-label{{font-size:9px;font-weight:700;text-transform:uppercase;letter-spacing:.6px;color:rgba(255,255,255,.4)}}
.rg-nivel{{font-size:11.5px;font-weight:700;line-height:1.1}}
.rg-trend{{font-size:10.5px;font-weight:600;line-height:1.1}}
.rg-better{{color:#86efac}}.rg-worse{{color:#fca5a5}}.rg-same{{color:rgba(255,255,255,.35)}}
/* KPI tiles (header) */
.hdr .kpis{{display:flex;gap:4px;flex-shrink:0}}
.hdr .kpi{{background:rgba(255,255,255,.05);border:1px solid rgba(255,255,255,.08);
  border-radius:var(--r);padding:6px 12px;text-align:center;min-width:62px}}
.hdr .kn{{font-size:18px;font-weight:800;line-height:1;font-variant-numeric:tabular-nums}}
.hdr .kl{{font-size:9.5px;color:rgba(255,255,255,.35);text-transform:uppercase;letter-spacing:.5px;margin-top:2px}}
.hdr .k-tot .kn{{color:#93c5fd}}
.hdr .k-crit .kn{{color:#fca5a5}}
.hdr .k-med .kn{{color:#fde68a}}
.hdr .k-low .kn{{color:#86efac}}

/* ── CONTAINER ── */
.wrap{{max-width:1080px;margin:28px auto;padding:0 24px}}
.section-title{{font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:.8px;color:var(--t3);margin-bottom:14px;margin-top:28px;padding-bottom:6px;border-bottom:1px solid var(--bd);display:flex;align-items:center;gap:7px}}
.section-title .ico{{width:14px;height:14px;color:var(--red)}}
.ico{{display:inline-block;vertical-align:-.14em;flex:none}}

/* ── KPI GRID ── */
.kpi-grid{{display:grid;grid-template-columns:repeat(auto-fill,minmax(155px,1fr));gap:12px;margin-bottom:6px}}
.kpi{{background:var(--s1);border:1px solid var(--bd);border-radius:var(--r2);padding:16px 18px;box-shadow:var(--sh)}}
.kpi-val{{font-size:28px;font-weight:700;letter-spacing:-1px;color:var(--t1)}}
.kpi-val.red{{color:var(--red)}}
.kpi-val.green{{color:var(--green)}}
.kpi-val.amber{{color:var(--amber)}}
.kpi-label{{font-size:11.5px;color:var(--t3);margin-top:4px}}
.kpi-sub{{font-size:10.5px;color:var(--t3);margin-top:2px}}

/* ── CARDS ── */
.card{{background:var(--s1);border:1px solid var(--bd);border-radius:var(--r2);overflow:hidden;margin-bottom:16px;box-shadow:var(--sh)}}
.card-hdr{{padding:11px 18px;border-bottom:1px solid var(--bd);background:var(--s2);font-size:11.5px;font-weight:700;color:var(--t2);text-transform:uppercase;letter-spacing:.5px}}
.card-body{{padding:18px;overflow-x:auto}}

/* ── TABLE ── */
table{{width:100%;min-width:520px;border-collapse:collapse;font-size:13px}}
th{{font-size:10.5px;font-weight:600;text-transform:uppercase;letter-spacing:.5px;color:var(--t3);padding:8px 12px;border-bottom:2px solid var(--bd2);text-align:left;background:var(--s2)}}
td{{padding:9px 12px;border-bottom:1px solid var(--bd);color:var(--t1)}}
tr:last-child td{{border-bottom:none}}
tr:hover td{{background:var(--s3)}}
.num{{text-align:center;font-variant-numeric:tabular-nums}}
.bold{{font-weight:600}}
.empty-row{{text-align:center;color:var(--green);padding:18px;font-weight:500}}

/* ── RISK LEVELS ── */
.risk-crit{{color:var(--red);font-weight:700}}
.risk-med{{color:var(--amber)}}
.risk-ok{{color:var(--green)}}

/* ── TREND ── */
.tend-up{{color:var(--red);font-weight:600;font-size:12px}}
.tend-dn{{color:var(--green);font-weight:600;font-size:12px}}
.tend-eq{{color:var(--t3);font-size:12px}}

/* ── PRIO PILLS (estética técnico) ── */
.prio-pill{{display:inline-flex;align-items:center;gap:4px;padding:3px 10px;border-radius:20px;font-size:11.5px;font-weight:700;white-space:nowrap;border:1px solid transparent}}
.sev-dot{{display:inline-flex;align-items:center;line-height:0;margin-right:2px}}
.sev-dot .ico{{width:.72em;height:.72em}}
.p-crit{{background:var(--crit-bg);color:var(--crit);border-color:var(--crit-border)}}
.p-alto{{background:var(--high-bg);color:var(--high);border-color:var(--high-border)}}

.mod-name{{font-weight:600;color:var(--t1)}}
.mod-tag{{background:var(--s2);border:1px solid var(--bd);border-radius:5px;font-size:11px;padding:2px 7px;color:var(--gray)}}

/* ── CAMBIOS / REINCIDENTES ── */
.two-col{{display:grid;grid-template-columns:1fr 1fr;gap:14px}}
.chg-row{{display:grid;grid-template-columns:1fr auto auto;gap:12px;align-items:center;padding:7px 0;border-bottom:1px solid var(--bd);font-size:12.5px}}
.chg-row:last-child{{border-bottom:none}}
.chg-mod{{color:var(--t2);font-weight:600}}
.chg-vals{{color:var(--t3);font-variant-numeric:tabular-nums;font-size:12px;white-space:nowrap}}
.chg-vals b{{color:var(--t1)}}
.chg-up{{color:var(--red);font-weight:700;font-variant-numeric:tabular-nums;white-space:nowrap}}
.chg-dn{{color:var(--green);font-weight:700;font-variant-numeric:tabular-nums;white-space:nowrap}}
.chg-eq{{color:var(--t3);font-weight:600;white-space:nowrap}}
.reinc-item{{padding:6px 0;border-bottom:1px solid var(--bd);font-size:12.5px}}
.reinc-item:last-child{{border-bottom:none}}
.reinc-item.ok{{color:var(--green)}}

/* ── NOTAS ── */
.notes-box{{background:rgba(29,78,216,.04);border:1px solid rgba(29,78,216,.15);border-radius:var(--r);padding:14px 16px;font-size:13px;color:var(--t2);white-space:pre-wrap;line-height:1.7}}

/* ── FOOTER ── */
.footer{{margin-top:36px;padding:16px 24px;border-top:1px solid var(--bd);text-align:center;font-size:11px;color:var(--t3)}}
.footer a{{color:var(--t3)}}

/* ── CVE PILLS ── */
.cve-diff{{display:flex;gap:8px;flex-wrap:wrap;margin-top:10px}}
.cve-pill{{display:inline-flex;align-items:center;gap:5px;padding:4px 10px;border-radius:9px;font-size:12px;font-weight:600}}
.cve-new{{background:rgba(214,48,49,.1);color:var(--red);border:1px solid rgba(214,48,49,.2)}}
.cve-res{{background:rgba(21,128,61,.08);color:var(--green);border:1px solid rgba(21,128,61,.2)}}
.cve-per{{background:rgba(180,83,9,.07);color:var(--amber);border:1px solid rgba(180,83,9,.18)}}

/* ── CVEs PRIORITARIOS (enriquecidos) ── */
.cve-badge{{display:inline-block;font-size:9.5px;font-weight:700;padding:1px 6px;border-radius:8px;margin-left:4px;vertical-align:middle;letter-spacing:.3px;white-space:nowrap}}
.cve-badge.kev{{background:var(--crit);color:#fff}}
.cve-badge.epss-hi{{background:var(--high-bg);color:var(--high);border:1px solid var(--high-border)}}
.cve-badge.epss{{background:var(--s3);color:var(--t2);border:1px solid var(--bd2)}}
.cve-score{{display:inline-block;min-width:32px;text-align:center;font-weight:800;border-radius:5px;padding:1px 6px;background:var(--s3);color:var(--t2)}}
.cve-score.p-crit{{background:var(--crit-bg);color:var(--crit)}}
.cve-score.p-alto{{background:var(--high-bg);color:var(--high)}}
.cve-sol{{color:var(--t1);font-weight:500;min-width:220px}}
.cve-src{{font-size:11px;color:var(--t3);margin:-8px 0 20px 2px}}

/* ── CHART ── */
.chart-wrap{{position:relative;height:220px}}
.chart-svg{{width:100%;overflow-x:auto}}
.chart-svg svg{{display:block;max-width:100%}}
.chart-legend{{display:flex;gap:18px;flex-wrap:wrap;justify-content:center;margin-top:12px;
  font-size:11.5px;color:var(--t2);font-weight:600}}
.chart-legend span{{display:inline-flex;align-items:center;gap:6px}}
.chart-legend i{{width:11px;height:11px;border-radius:3px;display:inline-block}}

/* ── HERO / RIESGO CREM ── */
.hero{{display:grid;grid-template-columns:270px 1fr;gap:20px;align-items:stretch;
  background:var(--s1);border:1px solid var(--bd);border-radius:var(--r2);
  box-shadow:var(--sh2);padding:22px 26px;margin-bottom:6px}}
.hero-gauge{{display:flex;flex-direction:column;align-items:center;justify-content:center;
  border-right:1px solid var(--bd);padding-right:20px}}
.hero-gauge .gauge-cap{{font-size:10.5px;font-weight:700;text-transform:uppercase;letter-spacing:.8px;color:var(--t3);margin-bottom:2px}}
.hero-trend{{margin-top:6px;font-size:12px;font-weight:700}}
.rk-worse{{color:var(--red)}}.rk-better{{color:var(--green)}}.rk-same{{color:var(--t3)}}
.hero-body{{display:flex;flex-direction:column;justify-content:center;gap:12px}}
.hero-verdict{{font-size:14.5px;line-height:1.6;color:var(--t1)}}
.hero-verdict b{{font-weight:700;color:var(--t1)}}
.hero-chips{{display:flex;gap:8px;flex-wrap:wrap}}
.hero-chip{{display:inline-flex;align-items:center;gap:6px;background:var(--s2);
  border:1px solid var(--bd2);border-radius:20px;padding:5px 12px;font-size:12px;font-weight:600;color:var(--t2)}}
.hero-chip b{{color:var(--t1);font-variant-numeric:tabular-nums}}
.hero-chip.crit{{background:rgba(214,48,49,.07);border-color:rgba(214,48,49,.2);color:var(--red)}}
.hero-chip.ok{{background:var(--green-s);border-color:rgba(21,128,61,.2);color:var(--green)}}

/* ── TOP 3 INCIDENTES ── */
.t3-grid{{display:grid;grid-template-columns:repeat(auto-fit,minmax(230px,1fr));gap:12px}}
.t3-card{{background:var(--s1);border:1px solid var(--bd);border-left:4px solid var(--t3);
  border-radius:var(--r2);padding:13px 15px;box-shadow:var(--sh)}}
.t3-card.r-crit{{border-left-color:var(--red);background:rgba(214,48,49,.03)}}
.t3-card.r-high{{border-left-color:var(--amber);background:rgba(180,83,9,.03)}}
.t3-card.r-med{{border-left-color:var(--blue)}}
.t3-hdr{{display:flex;align-items:center;gap:8px;flex-wrap:wrap;margin-bottom:7px}}
.t3-cat{{font-size:11.5px;font-weight:700;color:var(--t2)}}
.t3-date{{font-size:10.5px;color:var(--t3);margin-left:auto}}
.t3-event{{font-size:13px;font-weight:500;color:var(--t1);line-height:1.45;margin-bottom:8px}}
.t3-foot{{display:flex;align-items:center;gap:8px;flex-wrap:wrap;font-size:11.5px}}
.t3-asset{{color:var(--t3)}}
.t3-wb{{margin-left:auto;font-weight:600}}
.t3-empty{{color:var(--green);font-weight:500;padding:14px 2px}}

/* ── SPARKLINE ── */
.spark-box{{display:flex;align-items:center;gap:14px;margin-top:2px}}
.spark-box .spark-svg{{flex:1;min-width:0}}
.spark-meta{{font-size:11.5px;color:var(--t3);white-space:nowrap}}
.spark-meta b{{font-size:18px;font-weight:800;color:var(--t1);display:block;font-variant-numeric:tabular-nums}}

@media(max-width:820px){{
  .hdr{{position:static}}
  .hdr-in{{gap:10px;padding:10px 14px}}
  .rg-pill{{margin-left:0}}
  .hdr .kpis{{width:100%;justify-content:space-between}}
}}
@media(max-width:700px){{.two-col{{grid-template-columns:1fr}}.kpi-grid{{grid-template-columns:1fr 1fr}}
  .hero{{grid-template-columns:1fr}}.hero-gauge{{border-right:none;border-bottom:1px solid var(--bd);padding-right:0;padding-bottom:14px}}}}
@media(max-width:480px){{
  .hinfo{{order:5}}
  .htit,.hsub{{white-space:normal}}
  .wrap{{padding:0 12px;margin:18px auto}}
  .card-body{{padding:12px}}
  .kpi-grid{{grid-template-columns:1fr 1fr;gap:8px}}
  .kpi{{padding:12px 14px}}
  .kpi-val{{font-size:22px}}
}}
@media print{{
  body{{background:#fff;padding:0;font-size:11px}}
  .hdr{{position:static;background:#171922!important;-webkit-print-color-adjust:exact!important;print-color-adjust:exact!important}}
  .card,.hero,.t3-card{{box-shadow:none;border:1px solid #ddd;break-inside:avoid}}
  .section-title{{break-after:avoid}}
  .hero{{break-inside:avoid}}
  .badge, .score-pill, .prio-pill, .cve-badge, .kpi-val, .kpi-label, .hero-g-lbl, .cve-score {{
    -webkit-print-color-adjust: exact !important;
    print-color-adjust: exact !important;
  }}
  a[href]::after{{content:""}}
}}
{CSS_DEGRADADO}
"""


def resumir_lineas(texto: str, maximo: int = 6) -> str:
    """
    Recorta una lista de viñetas dejando las `maximo` primeras.

    El informe ejecutivo llegó a publicar celdas de 45 líneas de «acción
    recomendada» para un solo equipo: en el PDF es un muro de texto que nadie
    lee. El detalle completo sigue estando en el informe técnico.
    """
    lineas = [l for l in str(texto).splitlines() if l.strip()]
    if len(lineas) <= maximo:
        return "\n".join(lineas)
    restantes = len(lineas) - maximo
    # «acción» pierde la tilde en plural: acciones, no «acciónes».
    palabra = "acciones" if restantes > 1 else "acción"
    return "\n".join(lineas[:maximo]) + \
           f"\n… y {restantes} {palabra} más (ver informe técnico)"


def _html_degradaciones() -> str:
    """
    Aviso dentro del propio informe cuando algo ha salido incompleto.

    Quien abre el HTML no ve la consola ni el log: sin esto, un informe al que
    le falta un módulo entero es indistinguible de uno correcto.
    """
    if not DEGRADACIONES:
        return ""
    items = "".join(
        f"<li><b>{_esc_min(d['ambito'])}</b>: {_esc_min(d['detalle'])}"
        + (f" <span class='degr-i'>— {_esc_min(d['impacto'])}</span>" if d["impacto"] else "")
        + "</li>"
        for d in DEGRADACIONES)
    n = len(DEGRADACIONES)
    return (f'<div class="degr"><div class="degr-t">Este informe se ha generado con '
            f'{n} aviso{"s" if n > 1 else ""}</div><ul>{items}</ul></div>')

def _progress():
    return Progress(
        SpinnerColumn(style="cyan"),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(bar_width=28, complete_style="cyan", finished_style="green"),
        TaskProgressColumn(), MofNCompleteColumn(),
        TimeElapsedColumn(), TimeRemainingColumn(),
        console=console, transient=False,
    )

# ==============================================================================
# 5. CSV / CACHE UTILS
# ==============================================================================
def normalizar_csvs(dir_csv: Path) -> dict:
    """
    Escanea dir_csv buscando archivos CSV con nombres en bruto de Vision One
    (ej: 'Account Compromise Indicators_20260723095557.csv') y los renombra
    automáticamente a su versión estandarizada ('account-compromise.csv').
    Devuelve un diccionario {nombre_estandar: ruta_final}.
    """
    if not isinstance(dir_csv, Path):
        dir_csv = Path(dir_csv)
    if not dir_csv.exists() or not dir_csv.is_dir():
        return {}

    standard_names = {
        "account-compromise.csv", "anomaly-detections.csv", "cloud-app.csv",
        "cve-events.csv", "cve-assets.csv", "predictive-analytics.csv",
        "security-conf.csv", "sys-conf.csv", "threat-detections.csv"
    }

    all_csvs = [f for f in dir_csv.glob("*.csv") if f.is_file()]
    if not all_csvs:
        return {}

    unmatched = []
    already_standard = set()
    for f in sorted(all_csvs, key=lambda x: x.stat().st_mtime, reverse=True):
        if f.name in standard_names:
            already_standard.add(f.name)
        else:
            unmatched.append(f)

    if not unmatched:
        return {f.name: f for f in dir_csv.glob("*.csv") if f.name in standard_names}

    renamed_map = {}

    def _detect_target(path: Path) -> Optional[str]:
        stem = path.stem.lower().replace("-", " ").replace("_", " ")
        
        header = ""
        try:
            with open(path, "r", encoding="utf-8-sig", errors="replace") as fh:
                header = fh.readline().lower()
        except Exception:
            pass

        if "account compromise" in stem:
            return "account-compromise.csv"
        if "anomaly" in stem:
            return "anomaly-detections.csv"
        if "cloud app" in stem:
            return "cloud-app.csv"
        if "predictive" in stem:
            return "predictive-analytics.csv"
        if "security configuration" in stem or "security conf" in stem:
            return "security-conf.csv"
        if "system configuration" in stem or "sys conf" in stem:
            return "sys-conf.csv"
        if "threat" in stem:
            return "threat-detections.csv"

        if any(k in stem for k in ["cve", "highly exploitable", "vulnerab"]):
            if "vulnerability id" in header or "cve impact score" in header:
                return "cve-events.csv"
            if "device name" in header or "cve event risk score" in header or "operating system" in header:
                return "cve-assets.csv"

        if "vulnerability id" in header:
            return "cve-events.csv"
        if "device name" in header and "total cves" in header:
            return "cve-assets.csv"
        if "attack path risk score" in header:
            return "predictive-analytics.csv"

        return None

    for f in unmatched:
        target_name = _detect_target(f)
        if target_name:
            target_path = dir_csv / target_name
            try:
                if target_name in already_standard:
                    f.unlink()
                else:
                    f.rename(target_path)
                    already_standard.add(target_name)
                    renamed_map[f.name] = target_name
                    if 'info' in globals():
                        info(f"Normalizado CSV en bruto: [bold]{f.name}[/] → [bold cyan]{target_name}[/]")
            except Exception as e:
                if 'warn' in globals():
                    warn(f"No se pudo renombrar {f.name} a {target_name}: {e}")

    return {f.name: f for f in dir_csv.glob("*.csv") if f.name in standard_names}

def _can_dec(raw: bytes, enc: str) -> bool:
    try: raw.decode(enc); return True
    except Exception: return False

def _leer_csv_raw(ruta: Path) -> pd.DataFrame:
    if not ruta.exists(): return pd.DataFrame()
    raw = ruta.read_bytes()[:512]
    enc = next((e for e in _ENCODINGS if _can_dec(raw, e)), "latin-1")
    try:
        df = pd.read_csv(ruta, encoding=enc, dtype=str, low_memory=False,
                         on_bad_lines="skip", skipinitialspace=True)
    except Exception as e:
        err(f"Error leyendo {ruta.name}: {e}"); return pd.DataFrame()
    if df is None or df.empty: return pd.DataFrame()
    df.columns = [str(c).strip().lstrip('\ufeff').strip() for c in df.columns]
    df.dropna(how="all", inplace=True)
    # Normalizamos TODAS las columnas (no solo las de dtype object): pandas >= 3
    # devuelve StringDtype al leer con dtype=str, así que el filtro por `object`
    # dejaba pasar los NaN y estos acababan impresos como el texto "nan" en el
    # informe (y rompían las máscaras booleanas del tipo `.str.len() > 0`).
    for c in df.columns:
        df[c] = df[c].fillna("").astype(str).str.strip()
    return df.reset_index(drop=True)

def _leer_csv(nombre: str, dir_csv: Path = None) -> pd.DataFrame:
    d = dir_csv or CTX.dir_csv
    normalizar_csvs(d)
    ruta = d / f"{nombre}.csv"
    if not ruta.exists():
        if d == CTX.dir_csv: warn(f"No encontrado: {ruta.name}")
        return pd.DataFrame()
    return _leer_csv_raw(ruta)

def _parse_detail(detail: str) -> dict:
    result = {}
    for part in str(detail).split("|"):
        if ": " in part:
            k, v = part.split(": ", 1)
            result[k.strip()] = v.strip()
    return result

def _nw(nivel: str) -> int:
    return NIVEL_PESO.get((nivel or "").strip(), 0)

def _serie(df: pd.DataFrame, col: str, default: str = "") -> pd.Series:
    """Devuelve df[col] como Series; si la columna no existe, una Series
    constante alineada al índice de df.

    Necesario porque `df.get(col, "")` NO devuelve una Series cuando falta la
    columna: devuelve el propio str "". Encadenar .map/.isin/.astype sobre eso
    revienta con AttributeError y tumba la generación entera del informe si un
    CSV del cliente viene sin alguna columna esperada.
    """
    if col in df.columns:
        return df[col]
    return pd.Series(default, index=df.index, dtype=object)

# ──────────────────────────────────────────────────────────────────────────────
# Mapa sección del informe → módulos Vision One que la alimentan.
# Una sección "aplica" si AL MENOS uno de sus módulos está provisionado en el
# tenant del cliente (según .api_meta.json). Las secciones derivadas/internas
# (None) siempre se muestran. Permite distinguir:
#   · módulo no contratado          → ocultar sección
#   · módulo activo sin incidencias → mostrar "Sin incidencias este mes"
SECCION_MODULOS = {
    "sec-cve":  ["asm_vuln"],
    "sec-ca":   ["asm_vuln"],
    "sec-sys":  ["asm_assessments", "oat", "endpoint_health",
                 "audit_logs", "suspicious_objects", "intel_reports"],
    "sec-sec":  ["asm_assessments", "cloud_email"],
    "sec-thr":  ["workbench"],
    "sec-ano":  ["workbench"],
    "sec-cld":  ["cloud_access"],
    "sec-acc":  ["workbench"],
    # derivadas/internas → siempre visibles
    "sec-res": None, "sec-pred": None, "sec-cambios": None,
    "sec-rein": None, "sec-tend": None, "sec-plan": None,
}

# _RESUMEN_SID (etiqueta del resumen → sid) se deriva del registro de módulos,
# más abajo, para no repetir la correspondencia en dos sitios.

def _seccion_aplica(modules: dict, sid: str) -> bool:
    """True si la sección corresponde a un módulo contratado (o si no se puede
    determinar). Sin .api_meta.json → no se oculta nada (compatibilidad)."""
    mods = SECCION_MODULOS.get(sid, None)
    if not mods:        # derivada/interna o sid desconocido
        return True
    if not modules:     # sin metadatos de módulos → comportamiento clásico
        return True
    return any(modules.get(m) for m in mods)

def _cache_path(nombre: str) -> Path:
    return CTX.dir_datos / f"{nombre}.pkl"

def _guardar_cache(nombre: str, df: pd.DataFrame):
    try:
        df.to_pickle(_cache_path(nombre))
    except Exception as e:
        warn(f"Cache write fail {nombre}: {e}")

def _cargar_cache(nombre: str) -> Optional[pd.DataFrame]:
    p = _cache_path(nombre)
    if not p.exists(): return None
    try:
        df = pd.read_pickle(p)
    except Exception:
        return None
    # Las cachés escritas por versiones anteriores pueden conservar NaN en
    # columnas de texto; sin sanear se imprimían como el literal "nan".
    try:
        for c in df.columns:
            if not pd.api.types.is_numeric_dtype(df[c]):
                df[c] = df[c].fillna("").astype(str)
    except Exception:
        pass
    return df

def validar_vigencia_csv(ruta: Path) -> dict:
    """Inspecciona las columnas de fecha de un CSV para determinar la vigencia de los datos."""
    if not ruta.exists(): return {"status": "missing", "max_date": None, "days_old": None}
    df = _leer_csv_raw(ruta)
    if df.empty: return {"status": "empty", "max_date": None, "days_old": None}
    
    date_cols = [c for c in ["Detected", "Last detected", "First seen time", "Publish date"] if c in df.columns]
    if not date_cols:
        return {"status": "ok", "max_date": None, "days_old": 0}
        
    dates = []
    for col in date_cols:
        parsed = pd.to_datetime(df[col], errors="coerce")
        parsed = parsed.dropna()
        if not parsed.empty:
            dates.append(parsed.max())
            
    if not dates:
        return {"status": "ok", "max_date": None, "days_old": 0}
        
    max_d = max(dates).to_pydatetime()
    now = datetime.now()
    days = (now - max_d).days
    
    status = "stale" if days > 35 else "ok"
    return {"status": status, "max_date": max_d.strftime("%Y-%m-%d"), "days_old": days}

# ==============================================================================
# 6. VALIDACIÓN DE CSVs
# ==============================================================================
def _validar_csvs() -> bool:
    normalizar_csvs(CTX.dir_csv)
    ok_count = 0; warn_count = 0; stale_count = 0
    for nombre, cols_req in _CSV_SCHEMA.items():
        ruta = CTX.dir_csv / f"{nombre}.csv"
        if not ruta.exists():
            warn(f"[Validación] Falta: {nombre}.csv")
            warn_count += 1; continue
        df = _leer_csv_raw(ruta)
        if df.empty:
            warn(f"[Validación] Vacío: {nombre}.csv")
            warn_count += 1; continue
        def _falta(spec) -> bool:
            alts = spec if isinstance(spec, (tuple, list)) else (spec,)
            return not any(a in df.columns for a in alts)
        missing = [(" o ".join(c) if isinstance(c, (tuple, list)) else c)
                   for c in cols_req if _falta(c)]
        if missing:
            warn(f"[Validación] {nombre}.csv — columnas no encontradas: {missing}")
            warn_count += 1
        else:
            ok_count += 1
            
        vig = validar_vigencia_csv(ruta)
        if vig["status"] == "stale":
            warn(f"[Vigencia] {nombre}.csv — datos desactualizados (última fecha: {vig['max_date']}, hace {vig['days_old']} días)")
            stale_count += 1
            
    info(f"Validación: {ok_count} OK, {warn_count} avisos ({stale_count} desactualizados)")
    _procedencia_csvs()
    return True  # no bloqueamos, solo avisamos


def _procedencia_csvs() -> list[dict]:
    """
    Deja en el log la ficha de cada CSV que entra en el informe: filas, tamaño,
    fecha del fichero y hash corto del contenido.

    Nace de un caso real: dos ejecuciones del mismo mes dieron Riesgo CREM 18 y
    33, y no había forma de saber si los CSV de entrada eran los mismos. Con
    esto, comparar dos ejecuciones es comparar dos bloques del log.
    """
    import hashlib
    filas = []
    for m in MODULOS:
        ruta = CTX.dir_csv / f"{m.csv}.csv"
        if not ruta.exists():
            continue
        try:
            h = hashlib.sha256()
            n_lineas = 0
            with open(ruta, "rb") as fh:
                for bloque in iter(lambda: fh.read(1 << 20), b""):
                    h.update(bloque)
                    n_lineas += bloque.count(b"\n")
            st = ruta.stat()
            filas.append({
                "csv":    m.csv,
                "filas":  max(0, n_lineas - 1),          # sin la cabecera
                "kb":     round(st.st_size / 1024),
                "fecha":  datetime.fromtimestamp(st.st_mtime).strftime("%Y-%m-%d %H:%M"),
                "sha256": h.hexdigest()[:12],
            })
        except Exception as ex:
            degradado("Procedencia", f"no se pudo inspeccionar {m.csv}.csv ({ex})",
                      "esa entrada quedará sin registrar en el log")
    if not filas:
        return filas
    _log.info("Procedencia de los CSV de entrada:")
    _log.info(f"  {'csv':<22}{'filas':>8}{'KB':>9}  {'modificado':<17}sha256")
    for f in filas:
        _log.info(f"  {f['csv']:<22}{f['filas']:>8,}{f['kb']:>9,}  {f['fecha']:<17}{f['sha256']}")
    return filas

# ==============================================================================
# 7. DETECCIÓN CVEs NUEVOS VS MES ANTERIOR
# ==============================================================================
def _buscar_dir_anterior(mes_es: str) -> Optional[Path]:
    """Busca carpeta [EMPRESA]/INFORMES/CSV/csv-MES-AÑO del mes anterior."""
    base = CTX.empresa_dir / "INFORMES" / "CSV"

    # Determinar mes anterior
    try:
        partes = mes_es.split()
        anio = int(partes[-1])
        mes_num = list(MESES_ES.values()).index(partes[0].capitalize()) + 1
        ultimo_anterior = date(anio, mes_num, 1) - timedelta(days=1)
        mes_ant_es = MESES_ES[ultimo_anterior.strftime("%B")].lower()
        anio_ant = ultimo_anterior.year
        candidato = base / f"csv-{mes_ant_es}-{anio_ant}"
        if candidato.exists() and any(candidato.glob("*.csv")):
            return candidato
    except Exception:
        pass

    # Fallback: el csv-*-AÑO más reciente dentro de [EMPRESA]/INFORMES/CSV/,
    # ordenado por fecha real (el orden alfabético colocaba "septiembre" por
    # delante de "octubre") y excluyendo la carpeta del mes que se genera:
    # si el informe se rehace tras archivar, se compararía consigo mismo.
    if base.exists():
        _dir_actual = f"csv-{mes_es.replace('/', '-').replace(' ', '-').lower()}"
        cands = [p for p in base.glob("csv-*-????")
                 if p.is_dir() and p.name.lower() != _dir_actual and any(p.glob("*.csv"))]
        if cands:
            return sorted(cands, key=_hist_sort_key)[-1]
    return None

def comparar_cves(datos_actual: dict, dir_anterior: Optional[Path]) -> dict:
    """
    Compara cve-events actual vs anterior.
    Devuelve:
        nuevos:     set de CVE IDs nuevos este mes
        resueltos:  set de CVE IDs que ya no aparecen
        persistentes: set de CVE IDs que siguen
    """
    # Solo el nombre de la carpeta (p. ej. "csv-julio-2026"): el informe se
    # entrega al cliente y no debe exponer rutas locales del equipo.
    resultado = {"nuevos": set(), "resueltos": set(), "persistentes": set(),
                 "dir_anterior": dir_anterior.name if dir_anterior else None}
    if not dir_anterior:
        return resultado

    df_actual = datos_actual.get("cve_events", pd.DataFrame())
    df_ant    = _leer_csv("cve-events", dir_anterior)

    if df_actual.empty or df_ant.empty:
        degradado("Comparación CVE", "faltan datos del mes actual o del anterior", "no habrá CVEs nuevos ni resueltos en el informe")
        return resultado

    ids_actual = set(df_actual.get("Vulnerability ID", pd.Series()).dropna().unique())
    ids_ant    = set(df_ant.get("Vulnerability ID", pd.Series()).dropna().unique())

    resultado["nuevos"]      = ids_actual - ids_ant
    resultado["resueltos"]   = ids_ant    - ids_actual
    resultado["persistentes"]= ids_actual & ids_ant

    ok(f"Diff CVE: [new]+{len(resultado['nuevos'])} nuevos[/]  "
       f"[success]-{len(resultado['resueltos'])} resueltos[/]  "
       f"{len(resultado['persistentes'])} persistentes")
    _log.info(f"CVE diff: +{len(resultado['nuevos'])} nuevos, -{len(resultado['resueltos'])} resueltos")
    return resultado

# ==============================================================================
# 8. DEDUPLICACIÓN sys-conf / sec-conf
# ==============================================================================
def _max_fecha(serie: pd.Series) -> str:
    """Devuelve el valor más reciente de una serie de fechas en texto,
    comparando por fecha real (no lexicográficamente)."""
    if serie.empty: return ""
    dt = pd.to_datetime(serie, errors="coerce", dayfirst=False)
    if dt.notna().any():
        return serie.loc[dt.idxmax()]
    return serie.max()

def _dedup_conf(df: pd.DataFrame) -> pd.DataFrame:
    """
    Agrupa por tipo de evento y cuenta activos únicos afectados.
    Devuelve una fila por tipo de evento con columna 'Activos afectados'.
    """
    if df.empty: return df
    col_evento = "Risk event"
    col_nivel  = "Event risk level"
    col_asset  = "Asset"
    if col_evento not in df.columns: return df

    grp = df.groupby(col_evento, sort=False).agg(
        Nivel      =(col_nivel,  lambda x: (lambda m: m.iloc[0] if len(m) else "")(x.mode())),
        n_activos  =(col_asset,  "nunique"),
        Activos    =(col_asset,  lambda x: ", ".join(sorted(x.dropna().unique()))),
        Detectado  =("Detected", _max_fecha),
        Detail     =("Detail info", "first") if "Detail info" in df.columns else (col_asset, "first"),
    ).reset_index()
    grp["_peso"] = grp["Nivel"].map(_nw)
    grp = grp.sort_values("_peso", ascending=False)
    return grp

def _dedup_cloud_app(df: pd.DataFrame) -> pd.DataFrame:
    """
    Una fila por (equipo, aplicación cloud).
    Función aparte para poder aplicar exactamente el mismo tratamiento al mes
    actual y al CSV del mes anterior en la comparativa mensual.
    """
    if df.empty: return df
    df = df.copy()
    df["_peso"]   = _serie(df, "Event risk level").map(_nw)
    _det          = _serie(df, "Detail info").map(_parse_detail)
    df["_app"]    = _det.map(lambda d: d.get("appName",""))
    # Muchas filas llegan sin «Asset»; el equipo real viene en Detail info.
    # Sin este respaldo la columna «Equipo» salía vacía y la deduplicación
    # colapsaba filas de equipos distintos bajo la misma clave ("", app).
    df["Asset"]   = [a if a else (d.get("endpointHostName","") or d.get("endpointIp",""))
                     for a, d in zip(_serie(df, "Asset").astype(str), _det)]
    return df.drop_duplicates(subset=["Asset","_app"]).sort_values("_peso", ascending=False)

# ==============================================================================
# 8b. REGISTRO DE MÓDULOS — única fuente de verdad
# ==============================================================================
# Antes cada módulo se identificaba por su etiqueta en castellano, repetida como
# literal en la carga, el resumen, la comparativa mensual y el mapa de secciones.
# El emparejamiento entre tablas se hacía por subcadena, así que en cuanto dos
# sitios llamaban distinto al mismo módulo («Detecciones Amenaza» vs «Amenazas»)
# la fila se quedaba sin datos en silencio.
#
# Ahora el módulo se referencia SIEMPRE por su `id` (estable, nunca visible) y
# los textos viven aquí. Añadir un módulo nuevo es añadir una entrada.
class Modulo(NamedTuple):
    id:        str                 # identificador estable; jamás se muestra
    clave:     str                 # clave dentro del dict `datos`
    csv:       str                 # nombre del CSV, sin extensión
    sid:       str                 # id de la sección del informe
    etiqueta:  str                 # nombre en el resumen y la tabla de módulos
    # Nombre en la comparativa mes a mes. Coincide con `etiqueta` salvo en
    # Amenazas, que el informe llama de dos formas distintas desde siempre.
    # Igualar ambos cambiaría texto visible para el cliente: decisión suya.
    etiqueta_cambios: str
    peso_col:  Optional[str] = None    # columna de severidad para ordenar
    dedup:     Optional[Callable[[pd.DataFrame], pd.DataFrame]] = None
    en_resumen: bool = False           # aparece en el resumen ejecutivo
    en_cambios: bool = False           # aparece en la comparativa mensual

    @property
    def cache_key(self) -> str:
        return self.csv.replace("-", "_")


# El orden de esta lista es el orden en el que salen los módulos en TODAS las
# vistas: carga, resumen, gráfico de barras y comparativa mensual.
MODULOS: tuple[Modulo, ...] = (
    Modulo("cve_events",   "cve_events",   "cve-events",           "sec-cve",
           "CVE Eventos",         "CVE Eventos",
           peso_col="CVE impact score",     en_resumen=True, en_cambios=True),
    Modulo("cve_assets",   "cve_assets",   "cve-assets",           "sec-ca",
           "CVE Activos",         "CVE Activos",
           peso_col="CVE event risk score"),
    Modulo("sys_conf",     "sys_conf",     "sys-conf",             "sec-sys",
           "Config. Sistema",     "Config. Sistema",
           peso_col="Event risk level", dedup=_dedup_conf,
           en_resumen=True, en_cambios=True),
    Modulo("sec_conf",     "sec_conf",     "security-conf",        "sec-sec",
           "Config. Seguridad",   "Config. Seguridad",
           peso_col="Event risk level", dedup=_dedup_conf,
           en_resumen=True, en_cambios=True),
    Modulo("threats",      "threats",      "threat-detections",    "sec-thr",
           "Detecciones Amenaza", "Amenazas",
           peso_col="Event risk level",     en_resumen=True, en_cambios=True),
    Modulo("anomaly",      "anomaly",      "anomaly-detections",   "sec-ano",
           "Anomalías",           "Anomalías",
           peso_col="Event risk level",     en_resumen=True, en_cambios=True),
    Modulo("cloud_app",    "cloud_app",    "cloud-app",            "sec-cld",
           "Cloud Apps Riesgo",   "Cloud Apps Riesgo",
           dedup=_dedup_cloud_app,          en_resumen=True, en_cambios=True),
    Modulo("accounts",     "accounts",     "account-compromise",   "sec-acc",
           "Compromiso Cuentas",  "Compromiso Cuentas",
           peso_col="Event risk level",     en_resumen=True, en_cambios=True),
    Modulo("attack_paths", "attack_paths", "predictive-analytics", "sec-pred",
           "Analítica Predictiva", "Analítica Predictiva",
           peso_col="Attack path risk score"),
)

MODULO_POR_ID  = {m.id: m for m in MODULOS}
# Etiqueta visible (de cualquiera de las dos vistas) → módulo. Sustituye al
# emparejamiento por subcadena entre el resumen y la comparativa mensual.
MODULO_POR_ETIQUETA = {}
for _m in MODULOS:
    MODULO_POR_ETIQUETA[_m.etiqueta.lower()] = _m
    MODULO_POR_ETIQUETA[_m.etiqueta_cambios.lower()] = _m
del _m

MODULOS_RESUMEN = tuple(m for m in MODULOS if m.en_resumen)
MODULOS_CAMBIOS = tuple(m for m in MODULOS if m.en_cambios)

# Etiqueta del resumen ejecutivo → sid (para coherencia con la visibilidad)
_RESUMEN_SID = {m.etiqueta: m.sid for m in MODULOS_RESUMEN}


def _procesar_modulo(m: Modulo, df: pd.DataFrame) -> pd.DataFrame:
    """
    Aplica a un DataFrame recién leído el tratamiento propio de su módulo:
    columna de peso, orden por severidad y deduplicación.

    Es el ÚNICO sitio donde se decide cómo queda un módulo, y por eso vale
    tanto para el mes que se genera como para los meses del histórico. Antes
    había dos caminos y la comparativa mensual acababa enfrentando filas
    agregadas contra filas crudas («Config. Seguridad 1923 → 17»).
    """
    if df is None or df.empty:
        return df if df is not None else pd.DataFrame()
    if m.peso_col:
        df = df.copy()
        df["_peso"] = (
            pd.to_numeric(_serie(df, m.peso_col, "0"), errors="coerce").fillna(0)
            if m.peso_col in ("CVE impact score", "CVE event risk score")
            else _serie(df, m.peso_col).map(_nw)
        )
        df = df.sort_values("_peso", ascending=False)
    if m.dedup:
        df = m.dedup(df)
    return df


def cargar_mes(dir_csv: Optional[Path] = None, ids: Optional[Iterable[str]] = None) -> dict:
    """
    Lee y procesa un mes completo desde `dir_csv` (por defecto, el mes que se
    está generando). Devuelve {id_modulo: DataFrame} ya procesado.

    Camino ÚNICO de carga: lo usan el mes actual, el mes anterior de la
    comparativa y los meses del histórico, de modo que las cifras de todos son
    directamente comparables.
    """
    mods = MODULOS if ids is None else [MODULO_POR_ID[i] for i in ids]
    return {m.id: _procesar_modulo(m, _leer_csv(m.csv, dir_csv)) for m in mods}


# ==============================================================================
# 9. CARGA Y PROCESAMIENTO
# ==============================================================================
def cargar_todos(usar_cache: bool = False) -> dict:
    normalizar_csvs(CTX.dir_csv)
    seccion("Cargando y procesando CSVs", 1, 4)
    datos = {}

    def _load(m: Modulo) -> pd.DataFrame:
        # cloud-app se cargaba aparte y guardaba caché sin leerla nunca, así que
        # con --solo-word el módulo Cloud Apps salía vacío. Ahora todos los
        # módulos pasan por aquí, sin excepciones.
        if usar_cache:
            cached = _cargar_cache(m.cache_key)
            if cached is not None:
                ok(f"{m.csv:<22} {len(cached):>6,} registros [dim](caché)[/]")
                return cached
        df = _procesar_modulo(m, _leer_csv(m.csv))
        if not df.empty:
            _guardar_cache(m.cache_key, df)
        return df

    for m in MODULOS:
        datos[m.id] = _load(m)

    # Módulos provisionados (de la extracción API). Si no existe, dict vacío →
    # el informe no oculta nada (flujo manual / CSV antiguos).
    modules = {}
    meta_p = CTX.dir_csv / ".api_meta.json"
    if meta_p.exists():
        try:
            modules = json.loads(meta_p.read_text(encoding="utf-8")).get("modules", {}) or {}
        except Exception as e:
            degradado("Módulos contratados", f"no se pudo leer .api_meta.json ({e})", "no se ocultará ninguna sección aunque el módulo no esté contratado")
    datos["_modules"] = modules
    if modules:
        n_act = sum(1 for v in modules.values() if v)
        ok(f"{'módulos API':<22} {n_act:>6} activos [dim](de .api_meta.json)[/]")

    for k, df in datos.items():
        if isinstance(df, pd.DataFrame):
            ok(f"{k:<22} {len(df):>6,} registros")

    return datos

# ==============================================================================
# 10. PLAN DE ACTUACIÓN
# ==============================================================================
ACCIONES_SYSCONF = {
    "Deprecated OS Version":               "Actualizar SO a versión soportada.",
    "MFA Disabled":                        "Activar MFA vía Azure AD / Entra ID.",
    "Inactive Service Account":            "Deshabilitar o revisar cuenta de servicio.",
    "Stale":                               "Revisar y deshabilitar cuenta obsoleta.",
    "Password Expiration Disabled":        "Configurar política de expiración de contraseña.",
    "Password Not Required":               "Forzar contraseña en la cuenta afectada.",
    "Encryption Disabled":                 "Habilitar BitLocker / cifrado de disco.",
    "Firewall on Windows Device Disabled": "Habilitar firewall de Windows.",
    "Antivirus":                           "Verificar y activar el antivirus.",
    "Non-Compliant":                       "Revisar configuración según política indicada.",
    "Excessive":                           "Revisar y reducir privilegios excesivos.",
    "Legacy Authentication":               "Bloquear protocolos de autenticación heredados.",
    "Machine Account Quota":               "Ajustar quota de cuentas de máquina en AD.",
    "Non-Domain Joined":                   "Unir el dispositivo al dominio corporativo.",
    "Service Principal Without Owner":     "Asignar propietario al Service Principal.",
    "Virtualization Based":                "Habilitar HVCI.",
    "Unexpected Service":                  "Investigar y deshabilitar servicio expuesto.",
}
ACCIONES_SECCONF = {
    "Phishing":          "Optimizar config. anti-phishing en Exchange Online.",
    "Malicious File":    "Habilitar análisis archivos maliciosos en Exchange/Defender.",
    "BEC":               "Configurar protección anti-BEC en Exchange Online.",
    "Malicious URL":     "Activar Safe Links para análisis de URLs en correo.",
    "Spam":              "Optimizar filtros anti-spam en Exchange Online.",
    "Ransomware":        "Habilitar reglas anti-ransomware en Exchange Online.",
    "Apex One Firewall": "Optimizar firewall en Apex One.",
    "Device Control":    "Configurar Device Control en EWS.",
    "Firewall Settings": "Habilitar y optimizar Firewall en EWS.",
    "Log Inspection":    "Habilitar Log Inspection en EWS.",
    "FIM":               "Activar File Integrity Monitoring en EWS.",
    "Self-Protection":   "Habilitar Agent Self-Protection en EWS.",
    "IPS Settings":      "Optimizar IPS en EWS.",
    "Anti-Malware":      "Revisar config. anti-malware en EWS.",
    "Web Reputation":    "Activar Web Reputation en EWS.",
    "Application Control": "Configurar Application Control en EWS.",
}
ACCIONES_THREATS = {
    "Invalid IP Datagram":    "Investigar tráfico anómalo. Revisar logs IDS/IPS.",
    "Security Risk Detection":"Analizar en Apex One/EWS. Aislar si necesario.",
    "Malware":                "Cuarentena. Escaneo completo. Cambiar credenciales.",
    "Ransomware":             "⚠️ Aislar equipo INMEDIATAMENTE. Contactar respuesta.",
}
ACCIONES_ANOMALY = {
    "Unusual Increase of Service Operation Failure Rate": "Revisar logs del servicio.",
    "Unusual Device Access Day": "Verificar acceso con propietario del activo.",
    "Unusual": "Verificar actividad con propietario.",
}
ACCIONES_ACCOUNTS = {
    "Impersonation": "Verificar identidad. Revisar logs. Resetear contraseña si sospechoso.",
    "default":       "Contactar propietario. Bloquear IPs sospechosas.",
}

def _accion(evento, tabla):
    for k, v in tabla.items():
        if k.lower() in evento.lower(): return v
    return None

def construir_plan(datos: dict, diff_cves: dict, enrich_map: dict = None) -> pd.DataFrame:
    activos: dict[str, dict] = {}

    def _up(equipo, modulo, problema, accion, peso, es_nuevo=False):
        eq = str(equipo).strip()
        if not eq or eq == "nan": return
        if eq not in activos:
            activos[eq] = {"_peso":0, "_problemas":[], "_modulos":set()}
        activos[eq]["_problemas"].append((modulo, problema, accion, peso, es_nuevo))
        activos[eq]["_modulos"].add(modulo)
        if peso > activos[eq]["_peso"]: activos[eq]["_peso"] = peso

    # CVE assets
    nuevos_cve = diff_cves.get("nuevos", set())
    df = datos.get("cve_assets", pd.DataFrame())
    if not df.empty:
        for _, r in df.iterrows():
            score = r.get("CVE event risk score","0"); total = r.get("Total CVEs","?")
            try: peso = int(float(score)) // 10
            except Exception: peso = 0
            
            dev_name = r.get("Device name","")
            os_name = r.get("Operating system","")
            
            accion = "Parchear CVEs pendientes. Priorizar score ≥80 y explotabilidad activa."
            if enrich_map:
                sig = _os_sig(os_name)
                device_sols = []
                cve_df = datos.get("cve_events", pd.DataFrame())
                if not cve_df.empty and "OS/Application" in cve_df.columns:
                    sig_lower = sig.lower() if sig else ""
                    for _, rc in cve_df.iterrows():
                        cve_os = str(rc.get("OS/Application", "")).lower()
                        cve_id = str(rc.get("Vulnerability ID", ""))
                        if sig_lower and sig_lower in cve_os:
                            rec = enrich_map.get(cve_id, {})
                            if rec and rec.get("found"):
                                import cve_enrich
                                sol = cve_enrich.solucion_para_producto(rec, os_name) or rec.get("solucion", "")
                                if sol:
                                    device_sols.append(sol)
                if device_sols:
                    grouped_sols = agrupar_soluciones(device_sols)
                    accion = "\n".join(f"- {s}" for s in grouped_sols)
                    
            _up(dev_name, "CVE",
                f"Score {score}, {total} CVEs sin parchear",
                accion, peso)

    # Sys-conf / sec-conf (ya deduplicados)
    for key, mod, tabla, fb in [
        ("sys_conf", "Sys-Conf", ACCIONES_SYSCONF, "Revisar configuración del sistema."),
        ("sec_conf", "Sec-Conf", ACCIONES_SECCONF, "Revisar módulo de seguridad indicado."),
    ]:
        df = datos.get(key, pd.DataFrame())
        if df.empty: continue
        # columna Activos viene del dedup
        col_asset = "Activos" if "Activos" in df.columns else "Asset"
        col_evento = "Risk event"
        col_nivel  = "Nivel" if "Nivel" in df.columns else "Event risk level"
        for _, r in df.iterrows():
            evento = str(r.get(col_evento,""))
            nivel  = str(r.get(col_nivel,""))
            accion = _accion(evento, tabla) or fb
            # El dedup tiene una fila por tipo de evento con multiples activos
            activos_str = str(r.get(col_asset,""))
            for eq in [a.strip() for a in activos_str.split(",") if a.strip()]:
                _up(eq, mod, f"[{nivel}] {evento[:60]}", accion, _nw(nivel))

    # Threats / anomaly
    for key, mod, tabla, fb in [
        ("threats",  "Amenaza", ACCIONES_THREATS,  "Analizar en consola TrendAI."),
        ("anomaly",  "Anomalía",ACCIONES_ANOMALY,  "Verificar con propietario."),
    ]:
        df = datos.get(key, pd.DataFrame())
        if df.empty: continue
        for _, r in df.iterrows():
            evento = str(r.get("Risk event","")); nivel = str(r.get("Event risk level",""))
            _up(r.get("Asset",""), mod, f"[{nivel}] {evento[:60]}",
                _accion(evento, tabla) or fb, _nw(nivel))

    # Cloud
    df = datos.get("cloud_app", pd.DataFrame())
    if not df.empty:
        for _, r in df.iterrows():
            app = str(r.get("_app","")) or "App desconocida"
            nivel = str(r.get("Event risk level",""))
            _up(r.get("Asset",""), "Cloud", f"[{nivel}] Acceso a {app}",
                f"Verificar si '{app}' está sancionada. Si no, bloquear vía Zero Trust.", _nw(nivel))

    # Accounts
    df = datos.get("accounts", pd.DataFrame())
    if not df.empty:
        for _, r in df.iterrows():
            scope = str(r.get("Impact scope",""))
            evento = str(r.get("Risk event",""))
            nivel  = str(r.get("Event risk level",""))
            cuenta = scope.split(":")[-1].strip() if ":" in scope else scope
            _up(cuenta, "Cuenta", f"[{nivel}] {evento[:50]}",
                _accion(evento, ACCIONES_ACCOUNTS) or ACCIONES_ACCOUNTS["default"], _nw(nivel))

    rows = []
    for eq, data in sorted(activos.items(), key=lambda x: x[1]["_peso"], reverse=True):
        vistos, prob_unicos = set(), []
        for m, p, a, pe, es_n in sorted(data["_problemas"], key=lambda x: x[3], reverse=True):
            key = (m, p[:40])
            if key not in vistos:
                vistos.add(key); prob_unicos.append((m, p, a, pe, es_n))
        probs_items = []
        for m, p, _, _, es_n in prob_unicos:
            prefix = "[NUEVO] " if es_n else ""
            probs_items.append(f"• {prefix}{m}: {p.strip()}")
        probs_str = "\n".join(probs_items)

        acc_vistos = set()
        acc_items = []
        for m, p, a, pe, es_n in prob_unicos:
            if not a: continue
            for line in str(a).splitlines():
                line_clean = line.strip().lstrip("-").lstrip("•").strip()
                if line_clean and line_clean not in acc_vistos:
                    acc_vistos.add(line_clean)
                    acc_items.append(f"• {line_clean}")
        acc_str = "\n".join(acc_items) if acc_items else "• Seguir recomendaciones generales de seguridad."

        rows.append({
            "Activo / Equipo":      eq,
            "Prioridad":            _prioridad_label(data["_peso"]),
            "Problemas detectados": probs_str,
            "Acciones a realizar":  acc_str,
            "Módulos afectados":    ", ".join(sorted(data["_modulos"])),
            "_peso":                data["_peso"],
            "_detalles_json":       json.dumps([
                {"m":m,"p":p,"a":a,"pe":pe,"nuevo":es_n}
                for m,p,a,pe,es_n in prob_unicos], ensure_ascii=False),
        })
    return pd.DataFrame(rows) if rows else pd.DataFrame()

def _prioridad_label(p: int) -> str:
    if p >= 8: return "🔴 CRÍTICO"
    if p >= 4: return "🟠 ALTO"
    if p >= 3: return "🟡 MEDIO"
    return "🟢 BAJO"

def construir_resumen(datos: dict) -> list[dict]:
    filas = []
    _mods = datos.get("_modules", {}) or {}

    for m in MODULOS_RESUMEN:
        if not _seccion_aplica(_mods, m.sid):
            continue  # módulo no contratado → no aparece en resumen ni gráficos
        label = m.etiqueta
        df = datos.get(m.id, pd.DataFrame())

        # CVE Eventos se clasifica por «CVE impact score» (0-100), no por el
        # nivel textual que traen los demás módulos.
        if m.id == "cve_events":
            if not df.empty and "CVE impact score" in df.columns:
                scores = pd.to_numeric(df["CVE impact score"], errors="coerce").fillna(0)
                filas.append({"Módulo de Seguridad":label,"Total":len(df),
                              "Alto / Crítico":int((scores>=80).sum()),
                              "Medio":int(((scores>=60)&(scores<80)).sum()),
                              "Bajo":int((scores<60).sum())})
            else:
                filas.append({"Módulo de Seguridad":label,"Total":0,"Alto / Crítico":0,"Medio":0,"Bajo":0})
            continue

        if df.empty:
            filas.append({"Módulo de Seguridad":label,"Total":0,"Alto / Crítico":0,"Medio":0,"Bajo":0}); continue
        col_nivel = "Nivel" if "Nivel" in df.columns else "Event risk level"
        lv = df.get(col_nivel, pd.Series())
        vc = lv.value_counts() if len(lv) > 0 else pd.Series()
        filas.append({"Módulo de Seguridad":label,"Total":len(df),
                      "Alto / Crítico":int(vc.get("Critical",0))+int(vc.get("High",0)),
                      "Medio":int(vc.get("Medium",0)),"Bajo":int(vc.get("Low",0))})

    return filas

# ==============================================================================
# 11b. GENERADOR DE PLANTILLA WORD (Python puro, sin Node.js)
# ==============================================================================
def _crear_plantilla_python() -> bool:
    """
    Genera Revisión_CREM_MES_AÑO.docx con python-docx.
    No requiere Node.js ni archivos externos.
    Reproduce el branding corporativo de EMPRESA.
    """
    try:
        from docx import Document as _Doc
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        err("python-docx no disponible para generar plantilla.")
        return False

    # ── Colores EMPRESA ──────────────────────────────────────────────────────────
    R_RED  = RGBColor(0xDA, 0x29, 0x1C)
    R_GRAY = RGBColor(0x88, 0x8B, 0x8D)
    R_DARK = RGBColor(0x1A, 0x1A, 0x1A)
    R_WHT  = RGBColor(0xFF, 0xFF, 0xFF)

    PAGE_W_CM = 21.0; PAGE_H_CM = 29.7
    M_CM = 2.0

    doc = _Doc()

    # ── Página ──────────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width   = Cm(PAGE_W_CM)
    sec.page_height  = Cm(PAGE_H_CM)
    sec.top_margin   = Cm(M_CM)
    sec.bottom_margin= Cm(M_CM)
    sec.left_margin  = Cm(M_CM)
    sec.right_margin = Cm(M_CM)

    def _run(para, text, bold=False, size=11, color=None, italic=False):
        run = para.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color: run.font.color.rgb = color
        run.font.name = "Calibri"
        return run

    def _heading(text, level=1, color=None):
        p = doc.add_heading("", level=level)
        _run(p, text, bold=True,
             size={1:18,2:14,3:12}.get(level,12),
             color=color or R_DARK)
        # Red bottom border for H1
        if level == 1:
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), 'DA291C')
            pBdr.append(bottom)
            pPr.append(pBdr)
        return p

    def _body(text, color=None):
        p = doc.add_paragraph()
        _run(p, text, size=10, color=color or R_GRAY)
        return p

    def _space(): doc.add_paragraph()

    def _pb(): doc.add_page_break()

    def _make_table(headers, col_widths_cm):
        """Crea tabla con cabecera roja EMPRESA y fila de datos vacía."""
        n = len(headers)
        t = doc.add_table(rows=2, cols=n)
        t.style = 'Table Grid'
        # Header row
        hdr_row = t.rows[0]
        for ci, (hdr, w) in enumerate(zip(headers, col_widths_cm)):
            cell = hdr_row.cells[ci]
            # Background red
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'DA291C')
            tcPr.append(shd)
            # Width
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(int(w * 567)))  # cm to twips (approx)
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            # Text
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(hdr)
            run.bold = True
            run.font.color.rgb = R_WHT
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        # Data row — empty
        data_row = t.rows[1]
        for ci in range(n):
            cell = data_row.cells[ci]
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run("")
            run.font.size = Pt(9)
            run.font.name = "Calibri"
        return t

    def _infobox(icon, text):
        """Crea una fila de info con icono a la izquierda."""
        t = doc.add_table(rows=1, cols=2)
        t.style = 'Table Grid'
        # Icon cell
        c0 = t.rows[0].cells[0]
        tc0 = c0._tc
        tcPr0 = tc0.get_or_add_tcPr()
        shd0 = OxmlElement('w:shd')
        shd0.set(qn('w:val'),'clear'); shd0.set(qn('w:color'),'auto')
        shd0.set(qn('w:fill'),'DA291C'); tcPr0.append(shd0)
        p0 = c0.paragraphs[0]; p0.clear()
        r0 = p0.add_run(icon); r0.bold=True; r0.font.color.rgb=R_WHT
        r0.font.size=Pt(10); r0.font.name="Calibri"
        # Text cell
        c1 = t.rows[0].cells[1]
        tc1 = c1._tc
        tcPr1 = tc1.get_or_add_tcPr()
        shd1 = OxmlElement('w:shd')
        shd1.set(qn('w:val'),'clear'); shd1.set(qn('w:color'),'auto')
        shd1.set(qn('w:fill'),'F4F4F4'); tcPr1.append(shd1)
        p1 = c1.paragraphs[0]; p1.clear()
        r1 = p1.add_run(text); r1.font.size=Pt(9)
        r1.font.color.rgb=R_GRAY; r1.italic=True; r1.font.name="Calibri"
        return t

    # ── PORTADA ─────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(p, "Informe de Seguridad", bold=True, size=40, color=R_DARK)
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    _run(p2, "TrendAI CREM", bold=False, size=28, color=R_RED)
    p3 = doc.add_paragraph()
    _run(p3, "EMPRESA_NOMBRE", size=20, color=R_GRAY)
    p4 = doc.add_paragraph()
    _run(p4, "Mes 2026", size=16, color=R_GRAY)
    _pb()

    # ── ÍNDICE ──────────────────────────────────────────────────────────────────
    _heading("Índice", 1)
    for item in ["1. Resumen Ejecutivo","2. Cambios vs Mes Anterior",
                 "3. Activos Reincidentes","4. Vulnerabilidades CVE",
                 "5. Configuración del Sistema","6. Configuración de Seguridad",
                 "7. Detecciones de Amenazas","8. Anomalías",
                 "9. Cloud Apps de Riesgo","10. Compromiso de Cuentas",
                 "11. Plan de Actuación por Equipo"]:
        _body(item)
    _pb()

    # ── 1. RESUMEN EJECUTIVO ────────────────────────────────────────────────────
    _heading("1. Resumen Ejecutivo", 1)
    _body("Visión global del período. Los 5 problemas más críticos y acciones inmediatas.")
    _space()
    # T0: Resumen (6 cols con Tendencia)
    _make_table(
        ["Módulo de Seguridad","Total","Alto / Crítico","Medio","Bajo","Tendencia vs mes anterior"],
        [5.0, 1.5, 2.5, 1.8, 1.8, 4.4]
    )
    _space()
    _heading("1.1 Top 5 Acciones Inmediatas", 2)
    # T1-T5: infoboxes acciones ejecutivas
    for i in range(5):
        _infobox("🔴" if i < 2 else "🟠" if i < 4 else "🟡", f"[ACCIÓN {i+1} — rellena automáticamente]")
        _space()
    _pb()

    # ── 2. CAMBIOS VS MES ANTERIOR ──────────────────────────────────────────────
    _heading("2. Cambios vs Mes Anterior", 1)
    _body("Comparativa de eventos respecto al mes anterior.")
    _space()
    # T6
    _make_table(
        ["Módulo","Mes anterior","Mes actual","Variación","Activos nuevos","Activos resueltos"],
        [3.0, 2.5, 2.5, 2.0, 2.8, 4.2]
    )
    _space()
    _pb()

    # ── 3. ACTIVOS REINCIDENTES ─────────────────────────────────────────────────
    _heading("3. Activos Reincidentes — Sin Parchear N Meses", 1)
    _body("CVEs con 2 o más meses consecutivos sin resolver.")
    _space()
    # T7: infobox advertencia
    _infobox("⚠️", "Un activo reincidente 3+ meses sin resolver indica riesgo crítico acumulado.")
    _space()
    # T8
    _make_table(
        ["Activo / Equipo","CVE ID","Score","Meses sin resolver","Acción urgente"],
        [3.5, 2.8, 1.5, 3.0, 6.2]
    )
    _space()
    _pb()

    # ── 4. CVE ──────────────────────────────────────────────────────────────────
    _heading("4. Vulnerabilidades CVE", 1)
    _heading("4.1 Eventos CVE — Ordenados por Score", 2)
    _body("Cada CVE enlaza a su ficha oficial en cve.org. CVEs nuevos marcados en lila.")
    _space()
    # T9
    _make_table(
        ["CVE ID","Score","Explotabilidad","SO / Aplicación","Equipos Afect.","1ª Detección"],
        [2.8, 1.5, 3.5, 3.8, 2.6, 2.8]
    )
    _space()
    _heading("4.2 Activos con Mayor Exposición", 2)
    _body("AUT = Average Unpatched Time. Días medios sin parchear.")
    _space()
    # T10
    _make_table(
        ["Dispositivo","Sistema Operativo","IP","Score","Total CVEs","AUT (días)"],
        [3.0, 4.0, 2.5, 1.5, 2.2, 3.8]
    )
    _space()
    _pb()

    # ── 5. SYS-CONF ─────────────────────────────────────────────────────────────
    _heading("5. Configuración del Sistema", 1)
    _body("Problemas de configuración: SO obsoletos, MFA desactivado, cuentas inactivas…")
    _space()
    # T11
    _make_table(
        ["Evento de Riesgo","Activos afectados","Nivel","Detectado","Detalle"],
        [5.0, 3.8, 1.5, 2.8, 4.9]
    )
    _space()
    _pb()

    # ── 6. SEC-CONF ─────────────────────────────────────────────────────────────
    _heading("6. Configuración de Seguridad", 1)
    _body("Módulos de seguridad con configuraciones no optimizadas.")
    _space()
    # T12
    _make_table(
        ["Evento de Riesgo","Activos afectados","Nivel","Detectado","Módulo"],
        [5.0, 3.8, 1.5, 2.8, 4.9]
    )
    _space()
    _pb()

    # ── 7. THREATS ──────────────────────────────────────────────────────────────
    _heading("7. Detecciones de Amenazas", 1)
    _body("Detecciones activas: malware, IPS, reglas de Endpoint & Workload Security.")
    _space()
    # T13
    _make_table(
        ["Evento","Activo","Nivel","Detectado","Regla / Detalle"],
        [4.3, 3.5, 1.5, 2.8, 5.9]
    )
    _space()
    _pb()

    # ── 8. ANOMALY ──────────────────────────────────────────────────────────────
    _heading("8. Anomalías Detectadas", 1)
    _body("Comportamientos anómalos en dispositivos e identidades.")
    _space()
    # T14
    _make_table(
        ["Evento","Activo / Identidad","Nivel","Detectado","Detalle"],
        [4.3, 3.5, 1.5, 2.8, 5.9]
    )
    _space()
    _pb()

    # ── 9. CLOUD ────────────────────────────────────────────────────────────────
    _heading("9. Aplicaciones Cloud de Riesgo", 1)
    _body("Accesos a aplicaciones cloud no sancionadas.")
    _space()
    # T15
    _make_table(
        ["Aplicación","Equipo","Categoría","Nivel","Detectado"],
        [3.2, 3.2, 3.8, 1.5, 6.3]
    )
    _space()
    _pb()

    # ── 10. ACCOUNTS ────────────────────────────────────────────────────────────
    _heading("10. Compromiso de Cuentas", 1)
    _body("Intentos de suplantación, accesos inusuales, autenticación heredada.")
    _space()
    # T16
    _make_table(
        ["Evento","Cuenta / Activo","Nivel","Detectado","IPs Origen"],
        [4.0, 3.8, 1.5, 2.8, 5.9]
    )
    _space()
    _pb()

    # ── 11. PLAN DE ACTUACIÓN ────────────────────────────────────────────────────
    _heading("11. Plan de Actuación por Equipo", 1)
    _body("Tabla consolidada para el técnico en campo. Resolver primero los 🔴 CRÍTICO.")
    _space()
    # T17-T20: infoboxes prioridad
    _infobox("🔴", "CRÍTICO — Intervención inmediata. Score CVE ≥ 80 o nivel Critical/High.")
    _space()
    _infobox("🟠", "ALTO — Resolver en 72h.")
    _space()
    _infobox("🟡", "MEDIO — Resolver en 7 días.")
    _space()
    _infobox("🟢", "BAJO — Ciclo de mantenimiento habitual.")
    _space()
    # T21
    _make_table(
        ["Activo / Equipo","Prioridad","Problemas detectados","Acciones a realizar","Módulos afectados"],
        [3.2, 1.8, 4.5, 4.8, 2.7]
    )
    _space()
    _pb()

    # ── RECOMENDACIONES ──────────────────────────────────────────────────────────
    _heading("Recomendaciones Generales", 1)
    for rec in [
        "Parchear CVEs con score ≥ 80 y explotabilidad activa de forma inmediata.",
        "Activar MFA en todas las cuentas con política débil identificadas.",
        "Actualizar urgentemente sistemas operativos con fecha EOL superada.",
        "Revisar y sancionar aplicaciones cloud si su uso es legítimo.",
        "Investigar compromisos de cuentas y resetear credenciales afectadas.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        _run(p, rec, size=10, color=R_GRAY)

    # ── Guardar ──────────────────────────────────────────────────────────────────
    PLANTILLA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(PLANTILLA))
    ok(f"Plantilla generada: [dim]{PLANTILLA}[/] ({len(doc.tables)} tablas)")
    return True


# ==============================================================================
# 11. WORD — XML UTILS
# ==============================================================================
def _get_tcPr(tc):
    tcPr = tc.find(f"{_X}tcPr")
    if tcPr is None:
        tcPr = etree.SubElement(tc, f"{_X}tcPr"); tc.insert(0, tcPr)
    return tcPr

def _shading(tc, color):
    tcPr = _get_tcPr(tc)
    for s in tcPr.findall(f"{_X}shd"): tcPr.remove(s)
    shd = etree.SubElement(tcPr, f"{_X}shd")
    shd.set(f"{_X}val","clear"); shd.set(f"{_X}color","auto"); shd.set(f"{_X}fill", color)

def _bordes(tc):
    tcPr = _get_tcPr(tc)
    for b in tcPr.findall(f"{_X}tcBorders"): tcPr.remove(b)
    borders = etree.SubElement(tcPr, f"{_X}tcBorders")
    for lado in ("top","left","bottom","right"):
        el = etree.SubElement(borders, f"{_X}{lado}")
        el.set(f"{_X}val","single"); el.set(f"{_X}sz","4")
        el.set(f"{_X}space","0"); el.set(f"{_X}color", C_BORDE)

def _nivel_color_hex(nivel: str) -> Optional[str]:
    n = (nivel or "").lower().strip()
    return {"critical":C_CRIT,"high":C_HIGH,"medium":C_MED}.get(n)

def _estilizar_tabla(tabla, color_rows: Optional[dict] = None):
    color_rows = color_rows or {}
    for i, fila in enumerate(tabla.rows):
        color = C_RED if i == 0 else color_rows.get(i, "FFFFFF" if i%2==0 else "F8F8F8")
        for celda in fila.cells:
            tc = celda._tc
            _shading(tc, color); _bordes(tc)
            if i == 0:
                for run in tc.iter(f"{_X}r"):
                    rPr = run.find(f"{_X}rPr")
                    if rPr is None: rPr = etree.Element(f"{_X}rPr")
                    run.insert(0, rPr)
                    for c in rPr.findall(f"{_X}color"): rPr.remove(c)
                    etree.SubElement(rPr, f"{_X}color").set(f"{_X}val", C_WHITE)
                    if rPr.find(f"{_X}b") is None: etree.SubElement(rPr, f"{_X}b")

def _escribir_celda(fila_xml, idx: int, texto: str):
    celdas = fila_xml.findall(f"{_X}tc")
    if idx >= len(celdas): return
    tc = celdas[idx]
    for extra in tc.findall(f"{_X}p")[1:]: tc.remove(extra)
    p = tc.find(f"{_X}p")
    if p is None: return
    rpr_xml = None
    for r in p.findall(f"{_X}r"):
        rpr = r.find(f"{_X}rPr")
        if rpr is not None: rpr_xml = etree.tostring(rpr).decode(); break
    for child in list(p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r","proofErr","hyperlink"): p.remove(child)
    for li, line in enumerate(str(texto).split("\n")):
        if li > 0:
            etree.SubElement(etree.SubElement(p, f"{_X}r"), f"{_X}br")
        run = etree.SubElement(p, f"{_X}r")
        if rpr_xml:
            try: run.append(etree.fromstring(rpr_xml))
            except Exception: pass
        t = etree.SubElement(run, f"{_X}t")
        t.text = line
        t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")

def _escribir_celda_hyperlink(doc, fila_xml, idx: int, texto: str, url: str):
    celdas = fila_xml.findall(f"{_X}tc")
    if idx >= len(celdas): return
    tc = celdas[idx]
    for extra in tc.findall(f"{_X}p")[1:]: tc.remove(extra)
    p = tc.find(f"{_X}p")
    if p is None: return
    for child in list(p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r","proofErr","hyperlink"): p.remove(child)
    rid = doc.part.relate_to(url, _HYPER, is_external=True)
    hl = etree.SubElement(p, f"{_X}hyperlink")
    hl.set(f"{_R_X}id", rid); hl.set(f"{_X}history","1")
    run = etree.SubElement(hl, f"{_X}r")
    rPr = etree.SubElement(run, f"{_X}rPr")
    etree.SubElement(rPr, f"{_X}rStyle").set(f"{_X}val","Hyperlink")
    etree.SubElement(rPr, f"{_X}color").set(f"{_X}val","2C5F9E")
    etree.SubElement(rPr, f"{_X}u").set(f"{_X}val","single")
    t = etree.SubElement(run, f"{_X}t")
    t.text = str(texto)
    t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")

def _rellenar_tabla(tabla, filas_data: list, col_nivel: Optional[int] = None,
                    col_hyperlink: Optional[int] = None,
                    hyperlink_fn=None, doc=None, cve_nuevos: Optional[set] = None):
    tbl = tabla._tbl
    if not filas_data:
        _estilizar_tabla(tabla); return
    tmpl = copy.deepcopy(tabla.rows[1]._tr)
    for fila in tabla.rows[1:]: tbl.remove(fila._tr)
    color_rows = {}
    for row_idx, cols in enumerate(filas_data, start=1):
        nueva = copy.deepcopy(tmpl)
        for ci, val in enumerate(cols):
            if col_hyperlink is not None and ci == col_hyperlink and hyperlink_fn and doc:
                url = hyperlink_fn(cols)
                if url:
                    _escribir_celda_hyperlink(doc, nueva, ci, val, url)
                else:
                    _escribir_celda(nueva, ci, val)
            else:
                _escribir_celda(nueva, ci, val)
        # Color por nivel
        if col_nivel is not None and col_nivel < len(cols):
            c = _nivel_color_hex(cols[col_nivel])
            if c:
                color_rows[row_idx] = c
                for tc in nueva.findall(f"{_X}tc"): _shading(tc, c)
        # CVE nuevo → fondo lila claro
        if cve_nuevos and cols[0] in cve_nuevos:
            color_rows[row_idx] = "EDE7F6"
            for tc in nueva.findall(f"{_X}tc"): _shading(tc, "EDE7F6")
        tbl.append(nueva)
    _estilizar_tabla(tabla, color_rows)

def _rellenar_seccion(tabla, filas_data: list, sid: str, modules: dict, **kw):
    """Como _rellenar_tabla pero, si no hay filas, escribe una fila-nota que
    distingue 'módulo no contratado' de 'sin incidencias este mes'."""
    if not filas_data:
        try: ncols = len(tabla.columns)
        except Exception: ncols = 1
        msg = ("Sin incidencias detectadas este mes" if _seccion_aplica(modules, sid)
               else "Módulo no contratado — sin licencia en el tenant")
        fila = [msg] + ["—"] * max(0, ncols - 1)
        _rellenar_tabla(tabla, [fila])   # sin col_nivel: no colorear la nota
        return
    _rellenar_tabla(tabla, filas_data, **kw)

def _reemplazar_doc(doc, buscar: str, por: str):
    def _en(elem):
        for p in elem.iter(f"{_X}p"):
            runs = p.findall(f"{_X}r")
            texto = "".join((r.find(f"{_X}t").text or "") for r in runs if r.find(f"{_X}t") is not None)
            if buscar not in texto: continue
            rpr_xml = None
            for r in runs:
                rpr = r.find(f"{_X}rPr")
                if rpr is not None: rpr_xml = etree.tostring(rpr).decode(); break
            for r in runs: p.remove(r)
            run = etree.SubElement(p, f"{_X}r")
            if rpr_xml:
                try: run.append(etree.fromstring(rpr_xml))
                except Exception: pass
            t = etree.SubElement(run, f"{_X}t")
            t.text = texto.replace(buscar, por)
            t.set("{http://www.w3.org/XML/1998/namespace}space","preserve")
    _en(doc.element.body)
    for sec in doc.sections:
        for elem in (sec.header, sec.first_page_header, sec.even_page_header,
                     sec.footer, sec.first_page_footer, sec.even_page_footer):
            if elem: _en(elem._element)

# ==============================================================================
# 12. GENERACIÓN WORD
# ==============================================================================

# ==============================================================================
# CONFIG.JSON POR EMPRESA
# ==============================================================================
DEFAULT_CONFIG = {
    "empresa":        "",
    "sla_critico_dias": 1,
    "sla_alto_dias":    3,
    "sla_medio_dias":   7,
    "meses_reincidente": 2,
    "modulos_ignorar": [],
    "notas_adicionales": "",
    "contacto_tecnico": "",
    "abrir_html_al_terminar": False,
    "inventario_activos": {
        "_comentario": "Criticidad: MUY CRITICO | CRITICO | NORMAL | NO CRITICO | (vacío=sin catalogar). config.json se guarda en [EMPRESA]/config.json",
        "_ejemplo": {"Servidor01": {"descripcion": "BBDD Producción", "criticidad": "MUY CRITICO"}}
    }
}

def _get_inventario(cfg: dict) -> dict:
    """Devuelve {nombre: {descripcion, criticidad}} filtrando claves _comentario/_ejemplo."""
    return {k: v for k, v in cfg.get("inventario_activos", {}).items()
            if not k.startswith("_") and isinstance(v, dict)}

# Mapa estilo por criticidad — usado tanto en Python (badges) como en JS
# El primer elemento es un punto-icono SVG (antes emojis 💀🔴🟢⬜) que se propaga
# a badges de inventario, chips de criticidad y resumen CREM.
_CRIT_META = {
    "MUY CRITICO": (_dot("#8B0000"),     "#8B0000",      "#fff0f0", "muy-crit"),
    "CRITICO":     (_dot("#c62828"),     "#8B0000",      "#fff5f5", "crit-inv"),
    "NORMAL":      (_dot("#f57f17"),     "#e65100",      "#fff3e0", "normal-inv"),
    "NO CRITICO":  (_dot("#2e7d32"),     "var(--low)",   "var(--low-bg)", "no-crit-inv"),
    "":            (_dot("#9ca3af"),     "var(--gray)",  "var(--lgray)",  ""),
}

# Ejemplo de formato de inventario_activos — cada empresa define el suyo propio
# en su config.json (ver sección 7 del README). No se aplica ningún inventario
# por defecto: si config.json no trae inventario_activos, se generan activos
# "sin catalogar" (⬜) hasta que el usuario lo rellene desde el dashboard.
_INVENTARIO_EJEMPLO = {
    "_comentario": "Criticidad: MUY CRITICO | CRITICO | NO CRITICO | (vacío=sin catalogar)",
    "servidor-erp":      {"descripcion": "ERP Principal",                 "criticidad": "MUY CRITICO"},
    "dc-01":             {"descripcion": "Controlador de Dominio",        "criticidad": "CRITICO"},
    "pc-recepcion":      {"descripcion": "PC Recepción",                  "criticidad": "NO CRITICO"},
}

def _leer_config(empresa_dir: Path) -> dict:
    """
    Lee config.json de la empresa.
    📁 RUTA: [EMPRESA]/config.json
         ej: MIEMPRESA/config.json  (misma carpeta raíz de la empresa, junto a CSV/ e INFORMES/)
    """
    ruta = empresa_dir / "config.json"
    cfg = dict(DEFAULT_CONFIG)
    if ruta.exists():
        try:
            data = json.loads(ruta.read_text(encoding="utf-8"))
            cfg.update(data)
            ok(f"config.json leído: [dim]{ruta}[/]")
        except Exception as e:
            err(f"config.json inválido en {ruta} ({e}) — usando valores por defecto "
                f"(SLA e inventario de criticidad NO se están aplicando este mes, revisa el archivo)")
    else:
        # Crear config.json por defecto
        cfg["empresa"] = str(empresa_dir.name)
        ruta.write_text(json.dumps(cfg, indent=2, ensure_ascii=False), encoding="utf-8")
        ok(f"config.json creado: [dim]{ruta}[/]  ← guárdalo aquí para esta empresa")
    return cfg

# ==============================================================================
# MODO ACTUALIZACIÓN CSV
# ==============================================================================
def _actualizar_csv_historico(empresa_dir: Path, mes_es: str, risk_info: Optional[dict] = None,
                              mover: bool = True):
    """Archiva los CSV del mes en [EMPRESA]/INFORMES/CSV/csv-mes-año/ y persiste el
    Riesgo CREM final en risk_score.json.

    mover=True  → MUEVE los CSV (la carpeta CSV/ queda vacía para la próxima extracción).
    mover=False → los copia (se conservan en CSV/).
    """
    import shutil
    # Normalizar: "Mayo 2026" → "mayo-2026"
    mes_safe = mes_es.replace("/", "-").replace(" ", "-").lower()

    dir_csv_actual = empresa_dir / "CSV"
    dir_destino    = empresa_dir / "INFORMES" / "CSV" / f"csv-{mes_safe}"

    csvs = sorted(dir_csv_actual.glob("*.csv")) if dir_csv_actual.exists() else []
    if not csvs:
        info(f"No hay CSVs en [dim]{dir_csv_actual}[/] para archivar (¿ya movidos a histórico?).")
    else:
        dir_destino.mkdir(parents=True, exist_ok=True)
        n = 0
        # CSVs + metadatos de extracción (.api_meta.json) van juntos al histórico
        extra = [dir_csv_actual / ".api_meta.json"]
        for f in csvs + [e for e in extra if e.exists()]:
            dest = dir_destino / f.name
            try:
                if mover:
                    if dest.exists():
                        dest.unlink()
                    shutil.move(str(f), str(dest))
                else:
                    shutil.copy2(f, dest)
                n += 1
            except Exception as ex:
                degradado("Archivado de CSV", f"no se pudo mover {f.name} ({ex})", "ese CSV no estará en el histórico del mes que viene")
        verbo = "Movidos" if mover else "Copiados"
        ok(f"{verbo} {n} archivos → [dim]{dir_destino}[/]")
        if mover:
            info("Carpeta CSV/ vaciada — lista para la próxima extracción de la API.")

    if risk_info is not None:
        try:
            dir_destino.mkdir(parents=True, exist_ok=True)
            (dir_destino / "risk_score.json").write_text(
                json.dumps(risk_info, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception as ex:
            degradado("Riesgo CREM", f"no se pudo persistir risk_score.json ({ex})", "el mes que viene no podrá comparar contra este score")

# ==============================================================================
# ANÁLISIS DE CAMBIOS MES A MES
# ==============================================================================
def analizar_cambios(datos_actual: dict, dir_anterior: Optional[Path], diff_cves: dict) -> list[dict]:
    """
    Genera tabla comparativa módulo por módulo vs mes anterior.
    Retorna lista de dicts con: modulo, anterior, actual, variacion, nuevos, resueltos
    """
    filas = []

    def _count_nivel(df, nivel_col="Event risk level"):
        if df is None or df.empty: return 0
        col = nivel_col if nivel_col in df.columns else ("Nivel" if "Nivel" in df.columns else None)
        if col is None: return len(df)
        return int((df[col].isin(["Critical","High"])).sum())

    # El mes anterior se carga por el MISMO camino que el actual (`cargar_mes`),
    # así que llega con el mismo tratamiento: sys/sec-conf agrupados por tipo de
    # evento y cloud-app deduplicado por equipo+app. Cuando eran dos caminos
    # distintos se comparaban filas crudas contra filas agregadas y salían
    # caídas ficticias del tipo "Config. Seguridad 1923 → 17 (-99%)".
    datos_ant = (cargar_mes(dir_anterior, ids=[m.id for m in MODULOS_CAMBIOS])
                 if dir_anterior else {})

    def _cargar_anterior(mid: str) -> pd.DataFrame:
        return datos_ant.get(mid, pd.DataFrame())

    def _pct(act, ant):
        """Variación porcentual act vs ant como texto (o '' si no aplica)."""
        if not ant:
            return ""
        return f"{(act-ant)/ant*100:+.0f}%"

    # CVE
    act_cve = len(datos_actual.get("cve_events", pd.DataFrame()))
    ant_cve_df = _cargar_anterior("cve_events")
    ant_cve = len(ant_cve_df)
    nue_cve = len(diff_cves.get("nuevos", set()))
    res_cve = len(diff_cves.get("resueltos", set()))
    var_cve = act_cve - ant_cve
    _sc_act = pd.to_numeric(datos_actual.get("cve_events", pd.DataFrame()).get("CVE impact score", pd.Series()), errors="coerce").fillna(0)
    _sc_ant = pd.to_numeric(ant_cve_df.get("CVE impact score", pd.Series()), errors="coerce").fillna(0) if not ant_cve_df.empty else pd.Series()
    crit_act = int((_sc_act >= 80).sum())
    crit_ant = int((_sc_ant >= 80).sum()) if len(_sc_ant) else 0
    filas.append({
        "Módulo": MODULO_POR_ID["cve_events"].etiqueta_cambios,
        "Mes anterior": str(ant_cve) if ant_cve else "—",
        "Mes actual": str(act_cve),
        "Variación": f"{'↑' if var_cve>0 else '↓' if var_cve<0 else '='} {abs(var_cve):+d}" if ant_cve else "—",
        "Activos nuevos": str(nue_cve) if nue_cve else "—",
        "Activos resueltos": str(res_cve) if res_cve else "—",
        "_var_num": var_cve if ant_cve else 0,
        "_pct": _pct(act_cve, ant_cve),
        "_crit_ant": crit_ant, "_crit_act": crit_act,
        "_ant_n": ant_cve if dir_anterior else None, "_act_n": act_cve,
    })

    for m in MODULOS_CAMBIOS:
        if m.id == "cve_events":
            continue  # ya añadido arriba, con su diff de vulnerabilidades
        label  = m.etiqueta_cambios
        df_act = datos_actual.get(m.id, pd.DataFrame())
        df_ant = _cargar_anterior(m.id)

        act_n = len(df_act)
        ant_n = len(df_ant)
        var   = act_n - ant_n

        # Activos nuevos/resueltos (por columna Asset o Device name)
        col = "Device name" if m.id == "cve_assets" else "Asset"
        act_assets = set(df_act[col].dropna().unique()) if not df_act.empty and col in df_act.columns else set()
        ant_assets = set(df_ant[col].dropna().unique()) if not df_ant.empty and col in df_ant.columns else set()
        nuevos    = len(act_assets - ant_assets) if ant_assets else 0
        resueltos = len(ant_assets - act_assets) if ant_assets else 0

        # Con histórico, un módulo que pasa de 0 a N filas SÍ tiene variación
        # (+N). Antes se anulaba por el filtro `ant_n > 0` y se publicaba
        # "Compromiso Cuentas 0 → 4 · = 0".
        _hay_ant = bool(dir_anterior)
        filas.append({
            "Módulo": label,
            "Mes anterior": str(ant_n) if _hay_ant else "—",
            "Mes actual": str(act_n),
            "Variación": f"{'↑' if var>0 else '↓' if var<0 else '='} {abs(var):+d}" if _hay_ant else "—",
            "Activos nuevos": str(nuevos) if nuevos and _hay_ant else "—",
            "Activos resueltos": str(resueltos) if resueltos and _hay_ant else "—",
            "_var_num": var if _hay_ant else 0,
            "_pct": _pct(act_n, ant_n) if _hay_ant else "",   # _pct() ya devuelve "" si ant_n == 0
            "_crit_ant": _count_nivel(df_ant), "_crit_act": _count_nivel(df_act),
            "_ant_n": ant_n if _hay_ant else None, "_act_n": act_n,
        })

    return filas

# ==============================================================================
# ACTIVOS REINCIDENTES (N meses sin parchear)
# ==============================================================================
def _hist_sort_key(p: Path) -> tuple:
    """Clave de orden cronológico real para carpetas 'csv-<mes_es>-<año>'
    (el orden alfabético de meses en español no coincide con el cronológico)."""
    m = re.match(r"csv-([a-záéíóúñ]+)-(\d{4})$", p.name, re.IGNORECASE)
    if not m:
        return (0, 0)
    mes_en = MESES_ES_INV.get(m.group(1).lower(), "")
    orden_meses = list(MESES_ES.keys())  # January..December en orden
    mes_num = orden_meses.index(mes_en) + 1 if mes_en in orden_meses else 0
    return (int(m.group(2)), mes_num)

def detectar_reincidentes(empresa_dir: Path, datos_actual: dict, mes_es: str, meses_min: int = 2) -> list[dict]:
    """
    Detecta CVEs que llevan meses_min o más meses consecutivos sin resolverse.
    Lee los históricos de INFORMES/CSV/csv-*/cve-events.csv para comparar.
    """
    base_hist = empresa_dir / "INFORMES" / "CSV"
    if not base_hist.exists():
        return []

    # Leer todos los históricos ordenados cronológicamente (no alfabéticamente).
    # Se excluye la carpeta del mes que se está generando: si el informe se
    # rehace después de archivar los CSV, el mes actual aparecería también como
    # histórico y se contaría dos veces (más abajo ya se suma con `meses + 1`),
    # inflando los reincidentes y, con ellos, el Riesgo CREM.
    _dir_actual = f"csv-{mes_es.replace('/', '-').replace(' ', '-').lower()}"
    historicos = [p for p in sorted(base_hist.glob("csv-*-????"), key=_hist_sort_key)
                  if p.name.lower() != _dir_actual]
    if not historicos:
        return []

    # CVEs actuales
    df_act = datos_actual.get("cve_events", pd.DataFrame())
    if df_act.empty or "Vulnerability ID" not in df_act.columns:
        return []
    cves_actuales = set(df_act["Vulnerability ID"].dropna().unique())

    # Para cada CVE actual, contar meses CONSECUTIVOS hacia atrás en los que también aparece
    # (un hueco en un mes intermedio corta la racha, en vez de sumar apariciones sueltas)
    cve_sets_por_mes = []
    for hist_dir in reversed(historicos[-6:]):  # más reciente primero, máximo 6 meses atrás
        # `_leer_csv` (y no `_leer_csv_raw`) para que un histórico con los CSV
        # aún en bruto de Vision One también se lea. Con la lectura cruda, un
        # mes ilegible cortaba TODAS las rachas y los reincidentes caían a 0 en
        # silencio, restando 15 puntos al Riesgo CREM sin un solo aviso.
        df_h = _leer_csv("cve-events", hist_dir)
        if df_h.empty or "Vulnerability ID" not in df_h.columns:
            cve_sets_por_mes.append(None)  # sin datos ese mes: corta la racha (no se puede confirmar)
            degradado("Reincidentes",
                      f"no se pudo leer cve-events de {hist_dir.name}",
                      "la racha de meses se corta ahí; puede haber menos reincidentes de los reales")
            continue
        cve_sets_por_mes.append(set(df_h["Vulnerability ID"].dropna().unique()))

    cve_apariciones: dict[str, int] = {}
    for cve in cves_actuales:
        racha = 0
        for mes_set in cve_sets_por_mes:
            if mes_set is not None and cve in mes_set:
                racha += 1
            else:
                break
        if racha > 0:
            cve_apariciones[cve] = racha

    # Filtrar los que superan meses_min
    reincidentes_ids = {cve for cve, n in cve_apariciones.items() if n >= meses_min}
    if not reincidentes_ids:
        return []

    # Construir filas: cruzar con activos
    df_assets = datos_actual.get("cve_assets", pd.DataFrame())
    resultados = []
    for _, row in df_act.iterrows():
        cve_id = str(row.get("Vulnerability ID",""))
        if cve_id not in reincidentes_ids:
            continue
        meses = cve_apariciones[cve_id]
        score = str(row.get("CVE impact score",""))
        scope = str(row.get("Impact scope",""))
        # Acción según score
        try: s = int(float(score))
        except Exception: s = 0
        if s >= 80: accion = "⚠️ URGENTE — parchear inmediatamente"
        elif s >= 60: accion = "Parchear en próximo mantenimiento"
        else: accion = "Planificar parcheado"
        resultados.append({
            "activo":  scope,
            "cve_id":  cve_id,
            "score":   score,
            "meses":   str(meses + 1),  # +1 porque el actual también cuenta
            "accion":  accion,
        })

    # Ordenar por score descendente
    resultados.sort(key=lambda x: float(x["score"]) if x["score"].replace(".","").isdigit() else 0, reverse=True)
    return resultados

# ==============================================================================
# CVEs PRIORITARIOS ENRIQUECIDOS (NVD + KEV + EPSS)
# ==============================================================================
def _leer_env(path: Path) -> dict:
    """Parser .env mínimo (KEY=VALUE), misma convención que trendai_api."""
    out = {}
    try:
        for line in path.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, _, v = line.partition("=")
            out[k.strip()] = v.strip().strip('"').strip("'")
    except Exception:
        pass
    return out


def _resolver_nvd_key(cfg: dict, empresa_dir: Optional[Path] = None) -> Optional[str]:
    """
    Resuelve la API key de NVD por orden de prioridad:
      1) variable de entorno       NVD_API_KEY
      2) .env global del proyecto  BASE_DIR/.env        → NVD_API_KEY   (vale para todos)
      3) .env de la empresa        [EMPRESA]/.env       → NVD_API_KEY
      4) config.json de la empresa → "nvd_api_key"      (compatibilidad)
    Devuelve None si no hay ninguna.
    """
    env = os.environ.get("NVD_API_KEY", "").strip()
    if env:
        return env
    g = _leer_env(BASE_DIR / ".env").get("NVD_API_KEY", "").strip()
    if g:
        return g
    if empresa_dir:
        e = _leer_env(empresa_dir / ".env").get("NVD_API_KEY", "").strip()
        if e:
            return e
    val = str(cfg.get("nvd_api_key", "") or "").strip()
    return val or None


def enriquecer_cves(datos: dict, cache_dir, api_key=None, max_nvd=None) -> dict:
    """
    Enriquece TODOS los CVE de cve_events con NVD + KEV + EPSS (vía cve_enrich).
    Devuelve {cve_id: registro} o {} si el módulo no está disponible / sin CVEs.
    """
    if cve_enrich is None:
        degradado("Enriquecimiento CVE", "el módulo cve_enrich no está disponible",
                  "sin CVSS, sin CISA KEV y sin EPSS: la tabla de CVEs prioritarios saldrá vacía")
        return {}
    df = datos.get("cve_events", pd.DataFrame())
    if df.empty or "Vulnerability ID" not in df.columns:
        return {}
    ids = sorted({str(x) for x in df["Vulnerability ID"].dropna().unique() if str(x).startswith("CVE-")})
    if not ids:
        return {}
    info(f"Enriqueciendo [cyan]{len(ids)}[/] CVEs (NVD + CISA KEV + EPSS)…")
    try:
        return cve_enrich.enrich(ids, cache_dir=cache_dir, nvd_api_key=api_key,
                                 max_nvd=max_nvd, log=lambda m: info(f"[dim]{m}[/]"))
    except Exception as ex:
        degradado("Enriquecimiento CVE", f"la consulta falló ({ex})",
                  "sin CVSS, sin CISA KEV y sin EPSS: la tabla de CVEs prioritarios saldrá vacía")
        return {}


def construir_cves_prioritarios(datos: dict, enrich_map: dict, top_n: int = 8) -> list[dict]:
    """
    Construye la lista de CVEs prioritarios con su solución detallada.
    Orden: explotación activa (KEV) → EPSS ↓ → score TrendAI ↓.
    Cada item: id, score, activos, kev, epss, cvss, severidad, descripcion, solucion.
    """
    df = datos.get("cve_events", pd.DataFrame())
    if df.empty or not enrich_map or "Vulnerability ID" not in df.columns:
        return []

    # Contexto (producto + nº de activos) por CVE desde el CSV de TrendAI
    ctx: dict = {}
    for _, r in df.iterrows():
        cid = str(r.get("Vulnerability ID", ""))
        if not cid.startswith("CVE-"):
            continue
        if cid not in ctx:
            try:    scope = int(float(r.get("Impact scope", 0) or 0))
            except Exception: scope = 0
            try:    score = int(float(r.get("CVE impact score", 0) or 0))
            except Exception: score = 0
            ctx[cid] = {"os_app": str(r.get("OS/Application", "")),
                        "activos": scope, "score": score}

    items = []
    for cid, c in ctx.items():
        rec = enrich_map.get(cid, {})
        if not rec.get("found"):
            continue
        kev  = "kev" in rec
        epss = float(rec.get("epss", 0) or 0)
        sol  = cve_enrich.solucion_para_producto(rec, c["os_app"]) if cve_enrich else rec.get("solucion", "")
        items.append({
            "id": cid, "score": c["score"], "activos": c["activos"],
            "kev": kev, "epss": epss, "cvss": rec.get("cvss"),
            "severidad": rec.get("severidad", ""),
            "descripcion": rec.get("descripcion_es", ""),
            "solucion": sol or rec.get("solucion", ""),
            "ref": (rec.get("refs") or [{}])[0].get("url", ""),
        })
    items.sort(key=lambda x: (x["kev"], x["epss"], x["score"]), reverse=True)
    return items[:top_n]


# ==============================================================================
# TOP 5 ACCIONES EJECUTIVAS
# ==============================================================================
def construir_top5(datos: dict, diff_cves: dict, reincidentes: list[dict]) -> list[tuple[str,str,str]]:
    """
    Genera las 5 acciones más urgentes. Devuelve list de (emoji, nivel, texto).
    """
    acciones = []  # (peso, emoji, texto)

    # Reincidentes (máxima prioridad)
    for r in reincidentes[:2]:
        acciones.append((
            100 + float(r.get("score","0") or 0),
            "🔴",
            f"Activo reincidente {r['meses']} meses: {r['activo'][:40]} — {r['cve_id']} (score {r['score']}) — {r['accion']}"
        ))

    # CVEs nuevos críticos
    df_cve = datos.get("cve_events", pd.DataFrame())
    nuevos = diff_cves.get("nuevos", set())
    if not df_cve.empty and nuevos:
        df_new = df_cve[df_cve.get("Vulnerability ID", pd.Series()).isin(nuevos)]
        df_new_crit = df_new[pd.to_numeric(_serie(df_new, "CVE impact score", "0"), errors="coerce").fillna(0) >= 80]
        if not df_new_crit.empty:
            n = len(df_new_crit)
            top = df_new_crit.iloc[0]
            acciones.append((90,
                "🔴",
                f"{n} CVE{'s' if n>1 else ''} nuevos con score ≥80 este mes — Parchear urgente. Top: {top.get('Vulnerability ID','')} (score {top.get('CVE impact score','')})"
            ))

    # Cuentas comprometidas
    df_acc = datos.get("accounts", pd.DataFrame())
    if not df_acc.empty:
        hi = df_acc[_serie(df_acc, "Event risk level").isin(["High","Critical"])]
        if not hi.empty:
            acciones.append((85,
                "🔴" if len(hi) > 3 else "🟠",
                f"{len(hi)} compromiso{'s' if len(hi)>1 else ''} de cuenta{'s' if len(hi)>1 else ''} de nivel Alto/Crítico — Verificar identidades y resetear credenciales"
            ))

    # Sys-conf High
    df_sys = datos.get("sys_conf", pd.DataFrame())
    if not df_sys.empty:
        col_n = "Nivel" if "Nivel" in df_sys.columns else "Event risk level"
        hi = df_sys[_serie(df_sys, col_n).isin(["High","Critical"])]
        if not hi.empty:
            top_ev = str(hi.iloc[0].get("Risk event",""))[:50]
            acciones.append((75,
                "🟠",
                f"{len(hi)} configuración{'es' if len(hi)>1 else ''} de sistema de nivel Alto — Prioridad: {top_ev}"
            ))

    # Amenazas
    df_thr = datos.get("threats", pd.DataFrame())
    if not df_thr.empty:
        hi = df_thr[_serie(df_thr, "Event risk level").isin(["High","Critical","Medium"])]
        if not hi.empty:
            top_ev = str(hi.iloc[0].get("Risk event",""))[:50]
            acciones.append((70,
                "🟠",
                f"{len(hi)} detección{'es' if len(hi)>1 else ''} de amenaza activa — Revisar en consola TrendAI. Top: {top_ev}"
            ))

    # CVEs totales sin parchear críticos
    if not df_cve.empty:
        scores = pd.to_numeric(_serie(df_cve, "CVE impact score", "0"), errors="coerce").fillna(0)
        n_crit = int((scores >= 80).sum())
        if n_crit > 0:
            acciones.append((65,
                "🟡",
                f"{n_crit} CVE{'s' if n_crit>1 else ''} con score ≥80 pendientes de parchear — Revisar plan de actuación por equipo"
            ))

    # Ordenar y devolver top 5
    acciones.sort(key=lambda x: x[0], reverse=True)
    return [(emoji, txt) for _, emoji, txt in acciones[:5]]

# ==============================================================================
# RISK SCORE — KPI ejecutivo (0–100) para CIO/CISO
# ==============================================================================
def calcular_risk_score(datos: dict, diff_cves: dict, reincidentes: list,
                        tendencia_hist: list = None) -> dict:
    """
    Calcula el índice de riesgo consolidado (0-100) del tenant.
    Devuelve: {score, nivel, color, delta, trend, prev_score}
    """
    score = 0

    # CVEs críticos / altos (hasta 35 pts)
    df_cve = datos.get("cve_events", pd.DataFrame())
    if not df_cve.empty and "CVE impact score" in df_cve.columns:
        nums = pd.to_numeric(df_cve["CVE impact score"], errors="coerce").fillna(0)
        n_crit = int((nums >= 90).sum())
        n_high = int(((nums >= 70) & (nums < 90)).sum())
        n_med  = int(((nums >= 50) & (nums < 70)).sum())
        score += min(20, n_crit * 5) + min(10, n_high * 2) + min(5, n_med)

    # Amenazas activas High/Critical (hasta 25 pts)
    df_thr = datos.get("threats", pd.DataFrame())
    if not df_thr.empty:
        n_hi = int(df_thr.get("Event risk level", pd.Series()).isin(["Critical","High"]).sum())
        score += min(25, n_hi * 5)

    # Cuentas comprometidas (hasta 20 pts)
    df_acc = datos.get("accounts", pd.DataFrame())
    if not df_acc.empty:
        n_hi = int(df_acc.get("Event risk level", pd.Series()).isin(["Critical","High"]).sum())
        score += min(20, n_hi * 5)

    # Activos reincidentes sin parchear (hasta 15 pts)
    score += min(15, len(reincidentes) * 3)

    # Config crítica (hasta 5 pts)
    for key in ["sys_conf","sec_conf"]:
        df = datos.get(key, pd.DataFrame())
        if not df.empty:
            col_n = "Nivel" if "Nivel" in df.columns else "Event risk level"
            n_hi = int(df.get(col_n, pd.Series()).isin(["Critical","High"]).sum())
            score += min(2, n_hi)

    score = min(100, score)

    if   score >= 75: nivel, color = "CRÍTICO", "#c62828"
    elif score >= 50: nivel, color = "ALTO",    "#e65100"
    elif score >= 25: nivel, color = "MEDIO",   "#f57f17"
    else:             nivel, color = "BAJO",    "#2e7d32"

    # Tendencia vs mes anterior (usa el score real persistido si existe, si no aproxima).
    # `tendencia_hist` ya excluye el mes que se está generando, así que su último
    # elemento ES el mes anterior (antes se cogía [-2], un mes de más hacia atrás).
    delta = 0; trend = "—"; prev_score = None
    if tendencia_hist:
        prev = tendencia_hist[-1]
        if prev.get("risk_score") is not None:
            prev_score = float(prev["risk_score"])
        else:
            prev_crit = int(prev.get("cve_crit", 0))
            prev_thr  = int(prev.get("threats", 0))
            prev_score = min(100, prev_crit * 5 + min(25, prev_thr * 3) + 5)
        delta, trend = _delta_riesgo(score, prev_score)

    return {"score": score, "nivel": nivel, "color": color,
            "delta": delta, "trend": trend, "prev_score": prev_score}


def _delta_riesgo(score: float, prev_score: Optional[float]) -> tuple:
    """Variación del Riesgo CREM vs mes anterior → (delta, texto)."""
    if prev_score is None:
        return 0, "—"
    d = float(score) - float(prev_score)
    if   d >  5: return d, f"▲ +{d:.0f}"
    elif d < -5: return d, f"▼ {abs(d):.0f}"
    return d, "● Estable"


def _aplicar_riesgo_externo(risk_data: dict, score: float) -> dict:
    """
    Sustituye el score heurístico por el valor manual / de la API de Vision One.
    Recalcula también nivel, color y variación: antes solo se pisaba el número y
    el informe mostraba el score real junto a la tendencia del score automático,
    que podían contradecirse.
    """
    if   score >= 75: nivel, color = "CRÍTICO", "#c62828"
    elif score >= 50: nivel, color = "ALTO",    "#e65100"
    elif score >= 25: nivel, color = "MEDIO",   "#f57f17"
    else:             nivel, color = "BAJO",    "#2e7d32"
    delta, trend = _delta_riesgo(score, risk_data.get("prev_score"))
    risk_data.update({"score": score, "nivel": nivel, "color": color,
                      "delta": delta, "trend": trend, "manual": True})
    return risk_data


# ==============================================================================
# TOP 3 INCIDENTES CRÍTICOS (Workbench)
# ==============================================================================
def construir_top3_incidentes(datos: dict) -> list[dict]:
    """
    Devuelve los 3 incidentes más críticos/altos del Workbench.
    Sirve como tabla de ACCIONES PRIORITARIAS.
    """
    incidentes = []
    cats = [
        ("threats",  "🚨", "Amenaza"),
        ("anomaly",  "📡", "Anomalía"),
        ("accounts", "👤", "Cuenta"),
    ]
    for key, ico, cat_label in cats:
        df = datos.get(key, pd.DataFrame())
        if df.empty: continue
        col_a = "Impact scope" if key == "accounts" else "Asset"
        for _, r in df.iterrows():
            nivel = str(r.get("Event risk level",""))
            if nivel.lower() not in ("critical", "high"): continue
            det = _parse_detail(str(r.get("Detail info","")))
            wb_id   = det.get("workbenchId","")
            wb_link = det.get("link","")
            incidentes.append({
                "ico": ico, "cat": cat_label,
                "event":   str(r.get("Risk event",""))[:90],
                "asset":   str(r.get(col_a,""))[:50],
                "nivel":   nivel,
                "date":    str(r.get("Detected",""))[:16],
                "wb_id":   wb_id,
                "wb_link": wb_link,
                "_peso":   _nw(nivel),
            })

    incidentes.sort(key=lambda x: x["_peso"], reverse=True)
    return incidentes[:3]


# ==============================================================================
# CREM 5 PANELS — TOP 3 por dimensión (Devices/Internet/Accounts/Apps/Cloud)
# ==============================================================================
def construir_crem_panels(datos: dict, inv: dict = None) -> dict:
    """
    Construye los 5 paneles CREM con TOP 3 activos de riesgo por dimensión.
    """
    _inv = inv or {}

    def _crit_of(name_s: str) -> str:
        n = name_s.lower()
        for k, v in _inv.items():
            kl = k.lower()
            if kl and (kl in n or n in kl):
                return v.get("criticidad","").upper()
        return ""

    def _crit_bonus(crit: str) -> float:
        if crit == "MUY CRITICO": return 15.0
        if crit == "CRITICO":     return 8.0
        return 0.0

    panels: dict = {}

    # 1. Devices — TOP 3 por CVE risk score + bonus de criticidad de inventario
    df_ca = datos.get("cve_assets", pd.DataFrame())
    top_dev = []
    if not df_ca.empty:
        df_s = df_ca.copy()
        df_s["_num"]  = pd.to_numeric(_serie(df_s, "CVE event risk score", "0"), errors="coerce").fillna(0)
        df_s["_crit"] = _serie(df_s, "Device name").astype(str).map(_crit_of)
        df_s["_rank"] = df_s["_num"] + df_s["_crit"].map(_crit_bonus)
        try:
            for _, r in df_s.nlargest(3, "_rank").iterrows():
                top_dev.append({
                    "name":  str(r.get("Device name",""))[:40],
                    "score": str(r.get("CVE event risk score","")),
                    "cves":  str(r.get("Total CVEs","")),
                    "os":    str(r.get("Operating system",""))[:30],
                    "crit":  r["_crit"],
                })
        except Exception as ex:
            degradado("Panel CREM", f"la dimensión Equipos falló al construirse ({ex})", "el panel saldrá vacío")
    panels["devices"] = top_dev

    # 2. Internet — TOP 3 activos con IP (expuestos a internet, heurística) + bonus de criticidad
    top_inet = []
    if not df_ca.empty and "IP address" in df_ca.columns:
        df_ip = df_ca[df_ca["IP address"].str.strip().str.len() > 0].copy()
        df_ip["_num"]  = pd.to_numeric(_serie(df_ip, "CVE event risk score", "0"), errors="coerce").fillna(0)
        df_ip["_crit"] = _serie(df_ip, "Device name").astype(str).map(_crit_of)
        df_ip["_rank"] = df_ip["_num"] + df_ip["_crit"].map(_crit_bonus)
        try:
            for _, r in df_ip.nlargest(3,"_rank").iterrows():
                top_inet.append({
                    "name":  str(r.get("Device name",""))[:40],
                    "ip":    str(r.get("IP address","")),
                    "score": str(r.get("CVE event risk score","")),
                    "cves":  str(r.get("Total CVEs","")),
                    "crit":  r["_crit"],
                })
        except Exception as ex:
            degradado("Panel CREM", f"la dimensión Internet falló al construirse ({ex})", "el panel saldrá vacío")
    panels["internet"] = top_inet

    # 3. Accounts — TOP 3 cuentas comprometidas/alto riesgo
    df_acc = datos.get("accounts", pd.DataFrame())
    top_acc = []
    if not df_acc.empty:
        df_a = df_acc.copy()
        df_a["_peso"] = _serie(df_a, "Event risk level").map(_nw)
        try:
            for _, r in df_a.nlargest(3,"_peso").iterrows():
                top_acc.append({
                    "account": str(r.get("Impact scope","") or r.get("Asset",""))[:50],
                    "event":   str(r.get("Risk event",""))[:60],
                    "nivel":   str(r.get("Event risk level","")),
                    "date":    str(r.get("Detected",""))[:16],
                })
        except Exception as ex:
            degradado("Panel CREM", f"la dimensión Cuentas falló al construirse ({ex})", "el panel saldrá vacío")
    panels["accounts"] = top_acc

    # 4. Applications — TOP 3 aplicaciones con CVEs más críticos
    df_cve = datos.get("cve_events", pd.DataFrame())
    top_apps = []
    if not df_cve.empty and "OS/Application" in df_cve.columns:
        df_c2 = df_cve.copy()
        df_c2["_score"] = pd.to_numeric(_serie(df_c2, "CVE impact score", "0"), errors="coerce").fillna(0)
        try:
            grp = df_c2.groupby("OS/Application").agg(
                count=("Vulnerability ID","count"),
                max_score=("_score","max"),
            ).reset_index()
            grp = grp[grp["OS/Application"].str.strip().str.len() > 0]
            for _, r in grp.nlargest(3,"max_score").iterrows():
                top_apps.append({
                    "app":       str(r["OS/Application"])[:50],
                    "cves":      str(int(r.get("count",0))),
                    "max_score": str(int(r.get("max_score",0))),
                })
        except Exception as ex:
            degradado("Panel CREM", f"la dimensión Aplicaciones falló al construirse ({ex})", "el panel saldrá vacío")
    panels["applications"] = top_apps

    # 5. Cloud Assets — TOP 3 apps cloud de riesgo
    df_cld = datos.get("cloud_app", pd.DataFrame())
    top_cld = []
    if not df_cld.empty:
        df_c3 = df_cld.copy()
        df_c3["_peso"] = _serie(df_c3, "Event risk level").map(_nw)
        try:
            for _, r in df_c3.nlargest(3,"_peso").iterrows():
                top_cld.append({
                    "app":   str(r.get("_app","") or r.get("Risk event",""))[:40],
                    "asset": str(r.get("Asset",""))[:40],
                    "nivel": str(r.get("Event risk level","")),
                    "date":  str(r.get("Detected",""))[:16],
                })
        except Exception as ex:
            degradado("Panel CREM", f"la dimensión Cloud falló al construirse ({ex})", "el panel saldrá vacío")
    panels["cloud"] = top_cld

    return panels


# ==============================================================================
# TENDENCIA MENSUAL (para HTML)
# ==============================================================================
def leer_tendencia_historica(empresa_dir: Path, mes_es: str, n_meses: int = 6) -> list[dict]:
    """
    Lee los informes anteriores buscando datos de tendencia.
    Devuelve lista [{mes, cve_total, alto_total, ...}] ordenada cronológicamente.
    """
    base_hist = empresa_dir / "INFORMES" / "CSV"
    if not base_hist.exists():
        return []

    # Orden cronológico real, no alfabético: `sorted()` a secas colocaba
    # "Agosto" y "Julio" delante de "Junio"/"Mayo" y el gráfico de tendencia
    # salía desordenado (Abril → Agosto → Julio → Junio → Mayo).
    # Se excluye además la carpeta del mes que se está generando: si el informe
    # se rehace tras archivar los CSV, el mes actual aparecía como "histórico".
    _dir_actual = f"csv-{mes_es.replace('/', '-').replace(' ', '-').lower()}"
    historicos = [p for p in sorted(base_hist.glob("csv-*-????"), key=_hist_sort_key)
                  if p.name.lower() != _dir_actual][-n_meses:]
    tendencia = []

    for hist_dir in historicos:
        # Nombre del directorio: csv-mes-año → extraer mes y año
        parts = hist_dir.name.replace("csv-","").split("-")
        if len(parts) < 2: continue
        mes_label = parts[0].capitalize() + " " + parts[-1]

        # Mismo camino de carga que el mes actual: si se leyera en crudo, las
        # series de sys-conf compararían filas agregadas contra filas sueltas y
        # el gráfico caería en vertical sin que nada hubiera mejorado.
        hist = cargar_mes(hist_dir, ids=("cve_events", "threats", "sys_conf"))
        df_cve = hist["cve_events"]
        df_thr = hist["threats"]
        df_sys = hist["sys_conf"]

        if df_cve.empty:
            degradado("Tendencia", f"sin datos de CVE en {hist_dir.name}",
                      "ese mes sale a cero en el gráfico de evolución")

        scores = pd.to_numeric(df_cve.get("CVE impact score", pd.Series()), errors="coerce").fillna(0) if not df_cve.empty else pd.Series()

        risk_score_real = None
        risk_json = hist_dir / "risk_score.json"
        if risk_json.exists():
            try:
                risk_score_real = json.loads(risk_json.read_text(encoding="utf-8")).get("score")
            except Exception as ex:
                degradado("Tendencia", f"risk_score.json ilegible en {hist_dir.name} ({ex})",
                          "ese mes queda sin punto en la evolución del Riesgo CREM")

        tendencia.append({
            "mes":       mes_label,
            "cve":       len(df_cve),
            "cve_crit":  int((scores >= 80).sum()),
            "threats":   len(df_thr),
            "sys_issues":len(df_sys),
            "risk_score":risk_score_real,
        })

    return tendencia

def paso_generar_word(mes_es: str, datos: dict, diff_cves: dict, empresa: str = "", cfg: dict = None, cambios: list = None, reincidentes: list = None, enrich_map: dict = None) -> Optional[Path]:
    seccion("Generando informe Word", 2, 4)

    # Auto-generar plantilla si no existe o está desactualizada
    if not PLANTILLA.exists():
        info("Plantilla no encontrada — generando automáticamente…")
        if not _crear_plantilla_python():
            err("No se pudo generar la plantilla."); return None
    else:
        # Verificar que tiene EXACTAMENTE el número de tablas esperado. Las tablas
        # se referencian por índice fijo, así que un recuento distinto (más o menos)
        # obliga a regenerar la plantilla para mantener la coherencia de índices.
        _tmp = Document(str(PLANTILLA))
        if len(_tmp.tables) != N_TABLAS:
            info(f"Plantilla desactualizada ({len(_tmp.tables)} tablas, se esperan {N_TABLAS}) — regenerando…")
            if not _crear_plantilla_python():
                err("No se pudo regenerar la plantilla."); return None

    doc = Document(str(PLANTILLA))
    tablas = doc.tables
    if len(tablas) != N_TABLAS:
        err(f"Plantilla tiene {len(tablas)} tablas; se necesitan {N_TABLAS}."); return None

    _reemplazar_doc(doc, "Mes 2026", mes_es)
    _reemplazar_doc(doc, "EMPRESA_NOMBRE", empresa if empresa else "")
    if empresa:
        _reemplazar_doc(doc, "CREM", empresa)

    E = pd.DataFrame()
    nuevos_cve = diff_cves.get("nuevos", set())

    def _cve_url(row): return f"{CVE_BASE_URL}{row[0]}" if str(row[0]).startswith("CVE-") else None

    # Resumen
    # Resumen ejecutivo (ahora con columna Tendencia)
    res_filas = construir_resumen(datos)
    tendencia_hist = leer_tendencia_historica(CTX.empresa_dir, mes_es)
    def _tendencia_str(modulo):
        if len(tendencia_hist) < 2: return "—"
        # Match simple por nombre
        mod_l = modulo.lower()
        for key in ["cve","threats","sys","anomaly","cloud","accounts"]:
            if key in mod_l or any(k in mod_l for k in ["config","amenaza","anomal","cloud","cuentas"]):
                prev = tendencia_hist[-2]
                curr_f = [f for f in res_filas if modulo == f.get("Módulo de Seguridad","")]
                if curr_f and key == "cve":
                    c_curr = curr_f[0].get("Total",0); c_prev = prev.get("cve",0)
                    if c_prev == 0: return "—"
                    d = c_curr - c_prev
                    return f"{'↑' if d>0 else '↓' if d<0 else '='} {abs(d):+d}"
        return "—"
    rows_res = [[f["Módulo de Seguridad"],str(f["Total"]),str(f["Alto / Crítico"]),
                 str(f["Medio"]),str(f["Bajo"]), _tendencia_str(f["Módulo de Seguridad"])]
                for f in res_filas]
    _rellenar_tabla(tablas[TBL_RESUMEN], rows_res)

    # Top 5 acciones ejecutivas (T1-T5)
    if reincidentes is None: reincidentes = []
    top5 = construir_top5(datos, diff_cves, reincidentes)
    for i, (emoji, texto) in enumerate(top5):
        if i < len(TBL_ACCION):
            # Escribir texto en celda 1 (col 1) de la infobox
            tbl_inf = tablas[TBL_ACCION[i]]
            if tbl_inf.rows:
                _escribir_celda(tbl_inf.rows[0]._tr, 1, texto)

    # Cambios mes a mes
    if cambios:
        rows_cambios = [[c["Módulo"],c["Mes anterior"],c["Mes actual"],c["Variación"],c["Activos nuevos"],c["Activos resueltos"]] for c in cambios]
        _rellenar_tabla(tablas[TBL_CAMBIOS], rows_cambios)
        ok(f"Cambios mes a mes    {len(rows_cambios):>4,} módulos")

    # Activos reincidentes
    if reincidentes:
        rows_rein = [[r["activo"][:40],r["cve_id"],r["score"],r["meses"]+" mes(es)",r["accion"]] for r in reincidentes]
        _rellenar_tabla(tablas[TBL_REINCID], rows_rein)
        ok(f"Reincidentes         {len(rows_rein):>4,} CVEs")
    else:
        _rellenar_tabla(tablas[TBL_REINCID], [["Sin reincidentes detectados","—","—","—","—"]])

    # CVE diff note en resumen si hay nuevos
    if nuevos_cve:
        ok(f"CVE nuevos este mes: [new]{len(nuevos_cve)}[/] (marcados en lila en la tabla)")

    _wm = datos.get("_modules", {}) or {}   # módulos provisionados (visibilidad)

    # CVE events
    df = datos.get("cve_events", E)
    rows = [[str(r.get("Vulnerability ID","")),str(r.get("CVE impact score","")),
             str(r.get("Global exploit potential","")),str(r.get("OS/Application",""))[:40],
             str(r.get("Impact scope","")),str(r.get("First seen time",""))[:10]]
            for _, r in df.iterrows()] if not df.empty else []
    _rellenar_seccion(tablas[TBL_CVE], rows, "sec-cve", _wm, col_hyperlink=0, hyperlink_fn=_cve_url,
                    doc=doc, cve_nuevos=nuevos_cve)
    ok(f"CVE eventos   {len(rows):>5,} filas  ({len(nuevos_cve)} nuevos marcados)")

    # CVE assets
    df = datos.get("cve_assets", E)
    rows = [[str(r.get("Device name","")),str(r.get("Operating system",""))[:35],
             str(r.get("IP address","")),str(r.get("CVE event risk score","")),
             str(r.get("Total CVEs","")),str(r.get("Average Unpatched Time (AUT)",""))[:6]]
            for _, r in df.iterrows()] if not df.empty else []
    _rellenar_seccion(tablas[TBL_CVE_ASSET], rows, "sec-ca", _wm)
    ok(f"CVE activos   {len(rows):>5,} filas")

    # Sys-conf (deduplicado)
    df = datos.get("sys_conf", E)
    rows = []
    col_e = "Risk event"; col_n = "Nivel" if "Nivel" in df.columns else "Event risk level"
    col_a = "Activos" if "Activos" in df.columns else "Asset"
    for _, r in df.iterrows():
        det = _parse_detail(str(r.get("Detail info","") if "Detail info" in r.index else ""))
        extra = det.get("osName", det.get("modulesNotOptimized",""))[:40]
        rows.append([str(r.get(col_e,""))[:55], str(r.get(col_a,""))[:40],
                     str(r.get(col_n,"")), str(r.get("Detectado","") or r.get("Detected",""))[:16],
                     extra or str(r.get("n_activos",""))])
    _rellenar_seccion(tablas[TBL_SYSCONF], rows, "sec-sys", _wm, col_nivel=2)
    ok(f"Sys-conf      {len(rows):>5,} tipos únicos")

    # Sec-conf (deduplicado)
    df = datos.get("sec_conf", E)
    rows = []
    col_n2 = "Nivel" if "Nivel" in df.columns else "Event risk level"
    col_a2 = "Activos" if "Activos" in df.columns else "Asset"
    for _, r in df.iterrows():
        extra = _parse_detail(str(r.get("Detail info","") if "Detail info" in r.index else "")).get("modulesNotOptimized","")[:40]
        rows.append([str(r.get(col_e,""))[:55], str(r.get(col_a2,""))[:40],
                     str(r.get(col_n2,"")), str(r.get("Detectado","") or r.get("Detected",""))[:16],
                     extra or str(r.get("n_activos",""))])
    _rellenar_seccion(tablas[TBL_SECCONF], rows, "sec-sec", _wm, col_nivel=2)
    ok(f"Sec-conf      {len(rows):>5,} tipos únicos")

    # Threats
    df = datos.get("threats", E)
    rows = []
    for _, r in df.iterrows():
        det = _parse_detail(str(r.get("Detail info","")))
        rows.append([str(r.get("Risk event",""))[:55], str(r.get("Asset","")),
                     str(r.get("Event risk level","")), str(r.get("Detected",""))[:16],
                     det.get("ruleName", det.get("uuid",""))[:40]])
    _rellenar_seccion(tablas[TBL_THREATS], rows, "sec-thr", _wm, col_nivel=2)
    ok(f"Amenazas      {len(rows):>5,} filas")

    # Anomaly
    df = datos.get("anomaly", E)
    rows = []
    for _, r in df.iterrows():
        asset = str(r.get("Asset","")); ident = str(r.get("Identity type",""))
        rows.append([str(r.get("Risk event",""))[:55],
                     f"{asset} ({ident})" if ident else asset,
                     str(r.get("Event risk level","")), str(r.get("Detected",""))[:16],
                     str(r.get("Detail info",""))[:40]])
    _rellenar_seccion(tablas[TBL_ANOMALY], rows, "sec-ano", _wm, col_nivel=2)
    ok(f"Anomalías     {len(rows):>5,} filas")

    # Cloud
    df = datos.get("cloud_app", E)
    rows = []
    for _, r in df.iterrows():
        cat = _parse_detail(str(r.get("Detail info",""))).get("appCategory","")[:35]
        rows.append([str(r.get("_app",""))[:30], str(r.get("Asset","")), cat,
                     str(r.get("Event risk level","")), str(r.get("Detected",""))[:16]])
    _rellenar_seccion(tablas[TBL_CLOUD], rows, "sec-cld", _wm, col_nivel=3)
    ok(f"Cloud apps    {len(rows):>5,} filas")

    # Accounts
    df = datos.get("accounts", E)
    rows = []
    for _, r in df.iterrows():
        ips = _parse_detail(str(r.get("Detail info",""))).get("ips","")[:40]
        rows.append([str(r.get("Risk event",""))[:45], str(r.get("Impact scope",""))[:35],
                     str(r.get("Event risk level","")), str(r.get("Detected",""))[:16], ips])
    _rellenar_seccion(tablas[TBL_ACCOUNTS], rows, "sec-acc", _wm, col_nivel=2)
    ok(f"Cuentas       {len(rows):>5,} filas")

    # Plan
    df_plan = construir_plan(datos, diff_cves, enrich_map=enrich_map)
    rows_plan = []
    if not df_plan.empty:
        for _, r in df_plan.iterrows():
            rows_plan.append([str(r.get("Activo / Equipo","")),str(r.get("Prioridad","")),
                              str(r.get("Problemas detectados","")),str(r.get("Acciones a realizar","")),
                              str(r.get("Módulos afectados",""))])
    _rellenar_tabla(tablas[TBL_PLAN], rows_plan)
    ok(f"Plan actuación {len(rows_plan):>4,} activos")

    mes_safe = mes_es.replace("/","-").replace(" ","_")
    ruta = CTX.dir_informe / f"Revisión_CREM_{mes_safe}.docx"
    doc.save(str(ruta))
    ok(f"Word guardado → [bold]{ruta}[/]")
    return ruta

# ==============================================================================
# 13. GENERACIÓN HTML
# ==============================================================================

def _build_tendencia_html(tendencia: list[dict], mes_actual: str) -> str:
    """Genera HTML con la franja de evolución del Riesgo CREM + gráfico canvas de tendencia."""
    meses_json = json.dumps([t["mes"] for t in tendencia], ensure_ascii=False)
    cve_json   = json.dumps([t["cve"] for t in tendencia])
    crit_json  = json.dumps([t["cve_crit"] for t in tendencia])
    thr_json   = json.dumps([t["threats"] for t in tendencia])
    sys_json   = json.dumps([t["sys_issues"] for t in tendencia])

    # ── Franja de evolución del Riesgo CREM (score histórico persistido) ──────
    risk_strip = ""
    if any(t.get("risk_score") is not None for t in tendencia):
        cells, prev = [], None
        for t in tendencia:
            rs = t.get("risk_score")
            mes_lbl = _esc_min(str(t.get("mes", "")))
            if rs is None:
                cells.append(f'<div class="tr-cell"><div class="tr-mes">{mes_lbl}</div>'
                             f'<div class="tr-score tr-none">—</div></div>')
                continue
            rs = float(rs)
            lvl = "crit" if rs >= 75 else "high" if rs >= 50 else "med" if rs >= 25 else "low"
            delta = ""
            if prev is not None:
                d = rs - prev
                dcls = "up" if d > 0 else "down" if d < 0 else "flat"
                di = _ico("trend-up") if d > 0 else (_ico("trend-down") if d < 0 else "")
                delta = f'<span class="tr-d tr-{dcls}">{di}{d:+.0f}</span>'
            cells.append(f'<div class="tr-cell"><div class="tr-mes">{mes_lbl}</div>'
                         f'<div class="tr-score sc-{lvl}">{rs:.0f}</div>{delta}</div>')
            prev = rs
        risk_strip = (
            f'<div class="tr-wrap">'
            f'<div class="tr-title">{_ico("target")} Evolución del Riesgo CREM</div>'
            f'<div class="tr-strip">{"".join(cells)}</div></div>')

    return f"""
{risk_strip}
<div style="background:var(--lgray);border-radius:var(--r);padding:16px;margin-bottom:12px">
  <div style="font-size:12px;font-weight:700;color:var(--gray);text-transform:uppercase;letter-spacing:.5px;margin-bottom:12px">
    Evolución de incidencias · últimos {len(tendencia)} meses</div>
  <canvas id="trendChart" width="700" height="180" style="max-width:100%"></canvas>
</div>
<script>
(function(){{
  var meses={meses_json};
  var cve={cve_json};
  var crit={crit_json};
  var thr={thr_json};
  var sysI={sys_json};
  function drawTrend(){{
    var c=document.getElementById('trendChart');
    if(!c)return;
    var W=c.parentElement.offsetWidth-32||700; if(W<300)W=300;
    c.width=W; c.height=180;
    var ctx=c.getContext('2d');
    ctx.clearRect(0,0,W,180);
    var pad={{l:40,r:16,t:16,b:40}};
    var pw=W-pad.l-pad.r, ph=180-pad.t-pad.b;
    var n=meses.length; if(n<2)return;
    var datasets=[
      {{data:cve,   color:'#1565c0',label:'CVE Total'}},
      {{data:crit,  color:'#c62828',label:'CVE Crítico'}},
      {{data:thr,   color:'#e65100',label:'Amenazas'}},
      {{data:sysI,  color:'#888B8D',label:'Config. Sistema'}},
    ];
    var allVals=datasets.flatMap(function(d){{return d.data;}});
    var maxV=Math.max.apply(null,allVals)||1;
    // Grid
    ctx.strokeStyle='#e8e8e8'; ctx.lineWidth=1;
    for(var i=0;i<=4;i++){{
      var gy=pad.t+ph*(1-i/4);
      ctx.beginPath();ctx.moveTo(pad.l,gy);ctx.lineTo(pad.l+pw,gy);ctx.stroke();
      ctx.fillStyle='#888'; ctx.font='10px Segoe UI,sans-serif';ctx.textAlign='right';
      ctx.fillText(Math.round(maxV*i/4),pad.l-4,gy+3);
    }}
    // X labels
    meses.forEach(function(m,i){{
      var x=pad.l+i*(pw/(n-1));
      ctx.fillStyle='#888';ctx.font='10px Segoe UI,sans-serif';ctx.textAlign='center';
      var lbl=m.length>8?m.slice(0,3)+' '+m.split(' ').pop():m;
      ctx.fillText(lbl,x,180-8);
    }});
    // Lines
    datasets.forEach(function(ds){{
      ctx.strokeStyle=ds.color; ctx.lineWidth=2; ctx.lineJoin='round';
      ctx.beginPath();
      ds.data.forEach(function(v,i){{
        var x=pad.l+i*(pw/(n-1));
        var y=pad.t+ph*(1-v/maxV);
        if(i===0)ctx.moveTo(x,y);else ctx.lineTo(x,y);
      }});
      ctx.stroke();
      // Dots
      ds.data.forEach(function(v,i){{
        var x=pad.l+i*(pw/(n-1));
        var y=pad.t+ph*(1-v/maxV);
        ctx.beginPath();ctx.arc(x,y,3,0,Math.PI*2);
        ctx.fillStyle=ds.color;ctx.fill();
      }});
    }});
    // Legend
    var lx=pad.l;
    datasets.forEach(function(ds){{
      ctx.fillStyle=ds.color; ctx.fillRect(lx,4,10,10);
      ctx.fillStyle='#595959';ctx.font='10px Segoe UI,sans-serif';ctx.textAlign='left';
      ctx.fillText(ds.label,lx+13,13);
      lx+=ctx.measureText(ds.label).width+30;
    }});
  }}
  if(document.readyState==='loading')
    document.addEventListener('DOMContentLoaded',function(){{setTimeout(drawTrend,100);}});
  else setTimeout(drawTrend,100);
  window.addEventListener('resize',drawTrend);
  var secT=document.getElementById('sec-tend');
  if(secT)secT.addEventListener('toggle',function(){{setTimeout(drawTrend,50);}});
}})();
</script>
"""

def _js_tecnico(activos_json: str,
                 axref_json: str,
                 cve_por_activo_json: str,
                 inventario_js: str,
                 sidx_json: str,
                 slabels_json: str,
                 chart_bar_labels: str,
                 chart_bar_alto: str,
                 chart_bar_medio: str,
                 chart_bar_bajo: str,
                 chart_donut_data: str) -> str:
    """
    JavaScript del informe técnico: filtros, búsqueda, gráficos y modales.

    Recibe los datos ya serializados a JSON. Tenerlos en la firma y no
    sueltos como variables locales deja escrito qué necesita esta capa.
    """
    return f"""const IDX={sidx_json};
window._ldStep&&window._ldStep(7);
const SL={slabels_json};
const ICO={json.dumps(_ICONS_JS, ensure_ascii=False)};
const ACTIVOS={activos_json};
const AXREF={axref_json};
const CVE_ACTIVO={cve_por_activo_json};
const INV={inventario_js};
const INV_KEYS=Object.keys(INV).map(k=>{{
  const low=k.toLowerCase();
  return{{orig:k,low:low,short:low.split('.')[0]}};
}});
// Busca un nombre en el inventario con match parcial (FQDN vs nombre corto)
function invMatch(name){{
  if(!name)return null;
  const nl=name.toLowerCase();
  const ns=nl.split('.')[0];
  for(const k of INV_KEYS){{
    if(nl===k.low||ns===k.short||nl.includes(k.short)||k.short.includes(ns))
      return INV[k.orig];
  }}
  return null;
}}
const CBL={chart_bar_labels};
const CBAL={chart_bar_alto};
const CBME={chart_bar_medio};
const CBBA={chart_bar_bajo};
const CDD={chart_donut_data};
let SR=[],AF=new Set();

// Los gráficos del resumen (donut de severidad y barras por módulo) se renderizan
// como SVG en el servidor: no requieren JavaScript ni canvas.

// ── NAV ───────────────────────────────────────────────────────────────────────
window._ldStep&&window._ldStep(8);
function navTo(id,el){{
  document.querySelectorAll('.ni').forEach(n=>n.classList.remove('on'));
  if(el)el.classList.add('on');
  const t=document.getElementById(id);
  closeNav();
  if(!t)return;t.open=true;
  setTimeout(()=>t.scrollIntoView({{behavior:'smooth',block:'start'}}),50);
}}
function toggleNav(){{const l=document.getElementById('lay');if(l)l.classList.toggle('nav-open');}}
function closeNav(){{const l=document.getElementById('lay');if(l)l.classList.remove('nav-open');}}
function openSec(s){{const e=document.getElementById(s);if(e){{e.open=true;e.scrollIntoView({{behavior:'smooth',block:'start'}});}}}}
function _blinkRows(rows){{
  rows.forEach(r=>{{
    // Asegurar visibilidad aunque haya un filtro activo que la ocultara
    r.classList.remove('fo');
    if(r.style.display==='none')r.style.display='';
    r.classList.remove('blink-hl');
    void r.offsetWidth;              // reinicia la animación si ya estaba activa
    r.classList.add('blink-hl');
    setTimeout(()=>r.classList.remove('blink-hl'),2600);
  }});
}}
// Navega a la sección referenciada y resalta (blink) las filas que citan `term`.
// Soporta tablas normales y virtuales (se filtran para forzar el render de la fila).
function openSecHL(s,term,el){{
  const e=document.getElementById(s);
  if(!e)return false;
  // Abrir la sección y cualquier <details> ancestro por si está anidada/colapsada
  let p=e;
  while(p){{if(p.tagName==='DETAILS')p.open=true;p=p.parentElement;}}
  // Marcar la sección como activa en la navegación lateral
  document.querySelectorAll('.ni').forEach(n=>n.classList.remove('on'));
  const nav=document.querySelector('[data-t="'+s+'"]');
  if(nav)nav.classList.add('on');
  if(el){{el.classList.add('aref-hit');setTimeout(()=>el.classList.remove('aref-hit'),700);}}
  const tl=String(term).toLowerCase().trim();
  // Tabla virtual: rellenar su buscador para que renderice la fila objetivo
  const vf=e.querySelector('input[id$="-filter"]');
  if(vf){{vf.value=term;vf.dispatchEvent(new Event('input',{{bubbles:true}}));}}
  e.scrollIntoView({{behavior:'smooth',block:'start'}});
  let tries=0;
  (function find(){{
    const rows=Array.from(e.querySelectorAll('.dtbl tbody tr'));
    const matched=rows.filter(r=>{{
      const da=(r.getAttribute('data-assets')||'').toLowerCase();
      return da.includes(tl)||r.textContent.toLowerCase().includes(tl);
    }});
    if(matched.length){{
      _blinkRows(matched.slice(0,60));
      setTimeout(()=>matched[0].scrollIntoView({{behavior:'smooth',block:'center'}}),60);
    }} else if(tries++<12){{
      setTimeout(find,110);          // esperar al render (virtual / sección recién abierta)
    }}
  }})();
  return false;
}}
function secClick(s){{
  const n=document.querySelector('[data-t="'+s+'"]');
  document.querySelectorAll('.ni').forEach(x=>x.classList.remove('on'));
  if(n)n.classList.add('on');
}}
function expAll(){{document.querySelectorAll('details.sec').forEach(d=>d.open=true);}}
function colAll(){{document.querySelectorAll('details.sec').forEach(d=>d.open=false);}}

// ── FILTERS ───────────────────────────────────────────────────────────────────
const AF_INV=new Set();
function tglF(chip,lv){{chip.classList.toggle('on');AF.has(lv)?AF.delete(lv):AF.add(lv);applyF();}}
function tglInv(chip,crit){{chip.classList.toggle('on');AF_INV.has(crit)?AF_INV.delete(crit):AF_INV.add(crit);applyF();}}
const PRIO_LV={{'🔴 CRÍTICO':'critical','🟠 ALTO':'high','🟡 MEDIO':'medium','🟢 BAJO':'low'}};
function getRowCrit(tr){{
  // 1. Badge con data-crit (fiable, generado por Python con match parcial)
  const b=tr.querySelector('.inv-badge[data-crit]');
  if(b){{const c=(b.dataset.crit||'').trim();return c||'SIN_INV';}}
  // 2. Badge legacy sin data-crit
  const b2=tr.querySelector('.inv-badge');
  if(b2){{
    const t=b2.textContent.trim().toUpperCase();
    if(t.includes('MUY'))return 'MUY CRITICO';
    if(t.includes('NO'))return 'NO CRITICO';
    if(t.includes('CRIT')||t.includes('🔴'))return 'CRITICO';
  }}
  // 3. Tablas virtuales (sin badges): buscar en INV por texto de celda
  if(INV_KEYS&&INV_KEYS.length){{
    const cells=Array.from(tr.cells).slice(0,3).map(c=>c.textContent.toLowerCase()).join(' ');
    for(const k of INV_KEYS){{
      if(k.short&&cells.includes(k.short))return (INV[k.orig]||{{}}).crit||'SIN_INV';
      if(k.low&&cells.includes(k.low))return (INV[k.orig]||{{}}).crit||'SIN_INV';
    }}
  }}
  return 'SIN_INV';
}}
var _filterActive=false;
function applyF(){{
  const active=AF.size>0||AF_INV.size>0;
  _filterActive=active;
  // Show/hide filter info badge
  const fi=document.getElementById('flt-info');
  if(fi)fi.classList.toggle('on',active);
  // Filter rows
  let totVis=0,critVis=0,medVis=0,lowVis=0;
  document.querySelectorAll('.dtbl tbody tr').forEach(tr=>{{
    var lvOk=true,invOk=true;
    if(AF.size){{
      const b=tr.querySelector('.badge');
      if(b){{lvOk=AF.has(b.textContent.trim().toLowerCase());}}
      else{{const pp=tr.querySelector('.prio-pill');if(pp){{const cm={{'p-crit':'critical','p-high':'high','p-med':'medium','p-low':'low'}};let lv='';pp.classList.forEach(c=>{{if(cm[c])lv=cm[c];}});if(!lv)lv=PRIO_LV[pp.textContent.trim()]||'';lvOk=AF.has(lv);}}else lvOk=true;}}
    }}
    if(AF_INV.size){{invOk=AF_INV.has(getRowCrit(tr));}}
    const show=lvOk&&invOk;
    tr.classList.toggle('fo',!show);
    if(show){{
      totVis++;
      const b2=tr.querySelector('.badge');
      if(b2){{
        const lv=b2.textContent.trim().toLowerCase();
        if(lv==='critical'||lv==='high')critVis++;
        else if(lv==='medium')medVis++;
        else lowVis++;
      }}
    }}
  }});
  updateCounters(totVis,critVis,medVis,lowVis,active);
}}
function updateCounters(tot,crit,med,low,filtered){{
  // Update KPIs
  const kv=document.getElementById('kpi-vis');
  const kt=document.getElementById('kpi-tot');
  const kc=document.getElementById('kpi-crit');
  const km=document.getElementById('kpi-med');
  const kl=document.getElementById('kpi-low');
  if(filtered){{
    if(kv){{kv.textContent=tot;animateNum(kv);}}
    if(kc){{kc.textContent=crit;animateNum(kc);}}
    if(km){{km.textContent=med;animateNum(km);}}
    if(kl){{kl.textContent=low;animateNum(kl);}}
  }} else {{
    // Restore originals from data attributes
    restoreKpis();
  }}
  // Update sidebar section counters
  document.querySelectorAll('details.sec').forEach(sec=>{{
    const sid=sec.id;
    const cnt=document.getElementById('cnt-'+sid);
    if(!cnt)return;
    const total=parseInt(sec.dataset.total)||0;
    if(!filtered){{cnt.textContent=total;animateNum(cnt);return;}}
    const visible=sec.querySelectorAll('.dtbl tbody tr:not(.fo)').length;
    const prev=parseInt(cnt.dataset.vis??total);
    if(visible!==prev){{
      cnt.dataset.vis=visible;
      cnt.textContent=visible<total?visible+' / '+total:total;
      animateNum(cnt);
    }}
    // Update sidebar nav counter too
    const ni=document.querySelector('.ni[data-t="'+sid+'"] .ncnt');
    if(ni){{ni.textContent=filtered?(visible<total?visible+' / '+total:total):total;}}
  }});
}}
function animateNum(el){{
  el.classList.remove('changed');
  void el.offsetWidth;
  el.classList.add('changed');
}}
function _storeOrigKpis(){{
  ['kpi-tot','kpi-crit','kpi-med','kpi-low'].forEach(id=>{{
    const el=document.getElementById(id);
    if(el&&!el.dataset.orig)el.dataset.orig=el.textContent;
  }});
}}
function restoreKpis(){{
  ['kpi-tot','kpi-crit','kpi-med','kpi-low'].forEach(id=>{{
    const el=document.getElementById(id);
    if(el&&el.dataset.orig){{el.textContent=el.dataset.orig;animateNum(el);}}
  }});
  // Restore nav counters
  document.querySelectorAll('details.sec').forEach(sec=>{{
    const cnt=document.getElementById('cnt-'+sec.id);
    const total=parseInt(sec.dataset.total)||0;
    if(cnt){{cnt.textContent=total;cnt.dataset.vis=total;}}
    const ni=document.querySelector('.ni[data-t="'+sec.id+'"] .ncnt');
    if(ni)ni.textContent=total;
  }});
}}
function clearAllFilters(){{
  AF.clear();AF_INV.clear();
  document.querySelectorAll('.chip.on,.inv-chip.on').forEach(c=>c.classList.remove('on'));
  applyF();
}}

// ── INVENTORY SIDEBAR + ROW COLORS ───────────────────────────────────────────
function buildInvSidebar(){{
  const nav=document.querySelector('.nav');
  if(!nav||!INV||!Object.keys(INV).length)return;
  const order={{'MUY CRITICO':0,'CRITICO':1,'NORMAL':2,'NO CRITICO':3,'':4}};
  const sorted=Object.entries(INV).sort((a,b)=>(order[a[1].crit]??4)-(order[b[1].crit]??4));
  let html='<div class="nav-inv-hdr">Inventario activos</div>';
  sorted.forEach(([name,d])=>{{
    const lbl=name.length>16?name.slice(0,15)+'…':name;
    html+=`<div class="ni-inv" style="border-left-color:${{d.color}}" title="${{name}} — ${{d.desc}}">
      <span style="font-size:12px;flex-shrink:0">${{d.emoji}}</span>
      <span class="ni-inv-name" style="color:${{d.color}}">${{lbl}}</span>
      <span class="ni-inv-desc">${{d.desc}}</span>
    </div>`;
  }});
  nav.insertAdjacentHTML('beforeend',html);
}}
function applyInvRowColors(){{
  document.querySelectorAll('.dtbl tbody tr').forEach(tr=>{{
    tr.classList.remove('muy-crit-row','crit-inv-row','normal-inv-row');
    const crit=getRowCrit(tr);
    if(crit==='MUY CRITICO')tr.classList.add('muy-crit-row');
    else if(crit==='CRITICO')tr.classList.add('crit-inv-row');
    else if(crit==='NORMAL')tr.classList.add('normal-inv-row');
  }});
}}

// ── SEARCH ────────────────────────────────────────────────────────────────────
function doSearch(q){{
  const panel=document.getElementById('SRP'),list=document.getElementById('SRL'),cnt=document.getElementById('SRC');
  q=q.trim().toLowerCase();
  if(q.length<2){{panel.style.display='none';clrHL();return;}}
  const terms=q.split(/[ \t]+/);

  // Resultados del índice pre-construido
  SR=IDX.filter(x=>terms.every(t=>x.t.includes(t))).slice(0,40);

  // Búsqueda en tablas visibles (para secciones con más filas que el límite del índice)
  const liveResults=[];
  document.querySelectorAll('details.sec[open] .dtbl tbody tr').forEach(tr=>{{
    const txt=tr.textContent.toLowerCase();
    if(!terms.every(t=>txt.includes(t)))return;
    const sid=tr.closest('details.sec').id;
    // Evitar duplicar lo que ya está en IDX
    if(SR.some(x=>x.s===sid))return;
    const cells=Array.from(tr.cells).slice(0,3).map(c=>c.textContent.trim().slice(0,60));
    liveResults.push({{s:sid,row:tr,d:cells,live:true}});
    if(liveResults.length>=20)return;
  }});

  const totalRes=SR.length+liveResults.length;
  cnt.textContent=totalRes+' resultado'+(totalRes!==1?'s':'')+
    (liveResults.length?' ('+liveResults.length+' en tabla abierta)':'');

  const idxHtml=SR.map((x,i)=>{{
    const sl=SL[x.s]||x.s;
    const prev=x.d.join(' — ');
    const mk=prev.replace(new RegExp('('+terms.map(re).join('|')+')','gi'),'<mark>$1</mark>');
    return'<div class="sri" onclick="goR('+i+')"><span class="srlbl">'+sl+'</span><div class="srtxt">'+mk+'</div></div>';
  }}).join('');

  const liveHtml=liveResults.map((x,i)=>{{
    const sl=SL[x.s]||x.s;
    const prev=x.d.join(' — ');
    const mk=prev.replace(new RegExp('('+terms.map(re).join('|')+')','gi'),'<mark>$1</mark>');
    const li=SR.length+i;
    return'<div class="sri" onclick="goLive('+li+')"><span class="srlbl" style="background:#1565c0">'+sl+' ↓</span><div class="srtxt">'+mk+'</div></div>';
  }}).join('');

  list.innerHTML=(totalRes>0)?(idxHtml+liveHtml):'<div class="srnone">Sin resultados para "'+q+'"</div>';
  panel.style.display='block';

  // Guardar live results para goLive
  window._liveResults=liveResults;
}}
function re(s){{var sp=['\\\\','.','+','*','?','^','$','(',')','{','}','|','[',']'];for(var i=0;i<sp.length;i++)s=s.split(sp[i]).join('\\\\'+sp[i]);return s;}}
function goR(i){{
  const x=SR[i];if(!x)return;
  openSec(x.s);clrHL();
  const tb=document.querySelector('#'+x.s+' .dtbl tbody');
  if(tb){{const rows=tb.querySelectorAll('tr');if(rows[x.r]){{rows[x.r].classList.add('hl');setTimeout(()=>rows[x.r].scrollIntoView({{block:'center',behavior:'smooth'}}),200);}}}}
  document.getElementById('SRP').style.display='none';
}}
function goLive(i){{
  const lr=window._liveResults;if(!lr)return;
  const x=lr[i-SR.length];if(!x)return;
  clrHL();
  x.row.classList.add('hl');
  setTimeout(()=>x.row.scrollIntoView({{block:'center',behavior:'smooth'}}),100);
  document.getElementById('SRP').style.display='none';
}}
function clrHL(){{document.querySelectorAll('tr.hl').forEach(r=>r.classList.remove('hl'));}}
function skd(e){{if(e.key==='Escape')clrSearch();if(e.key==='Enter'&&SR.length)goR(0);}}
function clrSearch(){{document.getElementById('SI').value='';document.getElementById('SRP').style.display='none';clrHL();}}

// ── TABLE SORT ────────────────────────────────────────────────────────────────
document.querySelectorAll('.dtbl th').forEach((th,ci)=>{{
  let asc=true;
  th.addEventListener('click',()=>{{
    const tb=th.closest('table').querySelector('tbody');
    const rows=Array.from(tb.rows);
    rows.sort((a,b)=>{{
      const av=a.cells[ci]?.textContent.trim()||'';const bv=b.cells[ci]?.textContent.trim()||'';
      const an=parseFloat(av),bn=parseFloat(bv);
      if(!isNaN(an)&&!isNaN(bn))return asc?an-bn:bn-an;
      return asc?av.localeCompare(bv):bv.localeCompare(av);
    }});
    rows.forEach(r=>tb.appendChild(r));
    th.closest('tr').querySelectorAll('th').forEach(h=>h.classList.remove('sa-','sd-'));
    th.classList.add(asc?'sa-':'sd-');asc=!asc;
  }});
}});

// ── EXPORT CSV ────────────────────────────────────────────────────────────────
function exportCSV(tid, b64){{
  const bytes=atob(b64);
  const arr=new Uint8Array(bytes.length);
  for(let i=0;i<bytes.length;i++)arr[i]=bytes.charCodeAt(i);
  const blob=new Blob([arr],{{type:'text/csv;charset=utf-8;'}});
  const a=document.createElement('a');
  a.href=URL.createObjectURL(blob);a.download=tid+'.csv';a.click();
}}

// ── MODAL DETALLE ACTIVO ──────────────────────────────────────────────────────
function linkifyCVE(t){{
  return esc(t).replace(/CVE-\\d{{4}}-\\d{{4,7}}/g,
    m=>'<a href="https://www.cve.org/CVERecord?id='+m+'" target="_blank" class="ext">'+m+' '+ICO.ext+'</a>');
}}
function showModal(eq){{
  const data=ACTIVOS[eq];if(!data)return;
  document.getElementById('MT').innerHTML=ICO.search+' '+esc(eq);
  // Hipervínculos: dónde aparece este activo (salta + blink a cada sección)
  const refs=AXREF[eq]||[];
  const xref=refs.length?
    `<div class="modal-sec"><div class="modal-sec-title">Ver este activo en</div>`+
    `<div class="modal-xref">`+refs.map(x=>
      `<a class="aref" href="#${{x[0]}}" data-sid="${{x[0]}}">${{esc(x[1])}}</a>`).join(' ')+
    `</div></div>`:'';
  const cveBtn=CVE_ACTIVO[eq]?
    `<div class="modal-sec"><a class="aref aref-lg" href="#sec-ca" data-cve="1">`+
    ICO.cve+` Ver todos los CVEs de este activo</a></div>`:'';
  const prio=`<div class="modal-sec"><div class="modal-sec-title">Prioridad</div>`+
    `<div style="padding:4px 0">${{prio2pill(data.prioridad)}}</div>`+
    `<div style="font-size:12px;color:var(--gray);margin-top:4px">Módulos: ${{data.modulos}}</div></div>`;
  const items=data.detalles.map(d=>
    `<div class="modal-item">`+
    `<span class="modal-mod">${{esc(d.m)}}</span>`+
    `<div><div class="modal-prob">${{linkifyCVE(d.p)}}${{d.nuevo?' <span class="new-badge">NUEVO</span>':''}}</div>`+
    `<div class="modal-accion">→ ${{linkifyCVE(d.a)}}</div></div></div>`
  ).join('');
  const mb=document.getElementById('MB');
  mb.innerHTML=xref+cveBtn+prio+
    `<div class="modal-sec"><div class="modal-sec-title">Problemas y acciones (${{data.detalles.length}})</div>${{items}}</div>`;
  // Enlazar los chips por JS (evita problemas de escape en atributos)
  mb.querySelectorAll('.modal-xref .aref').forEach(a=>{{
    a.addEventListener('click',ev=>{{ev.preventDefault();closeModal();openSecHL(a.dataset.sid,eq,a);}});
  }});
  const cb=mb.querySelector('[data-cve]');
  if(cb)cb.addEventListener('click',ev=>{{ev.preventDefault();closeModal();showCVEModal(eq);}});
  document.getElementById('MO').classList.add('on');
}}
function closeModal(){{document.getElementById('MO').classList.remove('on');}}
function showCVEModal(eq){{
  const data=CVE_ACTIVO[eq];
  if(!data){{
    alert('Sin datos CVE para '+eq);return;
  }}
  document.getElementById('MCOT').innerHTML=ICO.cve+' CVEs que afectan a '+esc(eq);
  const os=data.os?esc(data.os):'—';
  const cves=data.cves||[];
  if(cves.length===0){{
    document.getElementById('MCOB').innerHTML=
      '<div class="cve-empty">'+ICO.info+
      ' No se identificaron CVEs por sistema operativo para este activo.<br>'+
      '<span style="color:var(--gray)">SO: '+os+' · '+data.total+' CVEs totales según el escaneo del activo.</span></div>';
    document.getElementById('MCO').classList.add('on');return;
  }}
  const nCrit=cves.filter(c=>parseFloat(c.score)>=80).length;
  const nHigh=cves.filter(c=>{{const s=parseFloat(c.score);return s>=60&&s<80;}}).length;
  const meta=
    '<div class="cve-meta">'+
      '<span class="cve-meta-pill" style="color:var(--crit)">'+ICO.dot+' Críticos (≥80): <b>'+nCrit+'</b></span>'+
      '<span class="cve-meta-pill" style="color:var(--high)">'+ICO.dot+' Altos (60-79): <b>'+nHigh+'</b></span>'+
      '<span class="cve-meta-pill">'+cves.length+' CVEs del SO</span>'+
      '<span class="cve-meta-pill">Score activo: <b>'+esc(data.score)+'</b></span>'+
      '<span class="cve-meta-pill">'+esc(data.total)+' CVEs totales</span>'+
    '</div>'+
    '<div class="cve-note">'+ICO.info+' CVEs que afectan al sistema operativo del activo (<b>'+os+'</b>). '+
    'El escaneo del activo reporta '+esc(data.total)+' CVEs en total.</div>'+
    '<input type="text" class="cve-filter" placeholder="Filtrar CVEs…" oninput="cveModalFilter(this)">';
  const rows=cves.map(c=>{{
    const sc=parseFloat(c.score)||0;
    const cls=sc>=80?'s-crit':sc>=60?'s-high':sc>=40?'s-med':'s-low';
    const url='https://www.cve.org/CVERecord?id='+encodeURIComponent(c.id);
    return`<tr>
      <td><a href="${{url}}" target="_blank" class="ext">${{esc(c.id)}} ${{ICO.ext}}</a></td>
      <td><span class="score-pill ${{cls}}">${{esc(c.score)}}</span></td>
      <td>${{esc(c.exploit)||'—'}}</td>
      <td>${{esc(c.os)||'—'}}</td>
      <td>${{cveAction(c.exploit,sc)}}</td>
    </tr>`;
  }}).join('');
  document.getElementById('MCOB').innerHTML=meta+
    '<div class="tbl-wrap"><table class="dtbl" id="cve-modal-tbl"><thead><tr>'+
    '<th>CVE ID</th><th>Score</th><th>Explotabilidad</th><th>SO / Aplicación</th><th>Acción recomendada</th>'+
    '</tr></thead><tbody>'+rows+'</tbody></table></div>';
  document.getElementById('MCO').classList.add('on');
}}
function cveModalFilter(inp){{
  const q=inp.value.toLowerCase().trim();
  const tb=document.getElementById('cve-modal-tbl');
  if(!tb)return;
  tb.querySelectorAll('tbody tr').forEach(tr=>{{
    tr.style.display=(!q||tr.textContent.toLowerCase().includes(q))?'':'none';
  }});
}}
function cveAction(exploit,score){{
  if((exploit||'').toLowerCase().includes('actively'))return'⚠️ Parchear urgente — explotabilidad activa';
  if(score>=80)return'Parchear lo antes posible (score crítico)';
  if(score>=60)return'Parchear en el próximo ciclo de mantenimiento';
  return'Monitorizar y parchear cuando sea posible';
}}
function closeCVEModal(){{document.getElementById('MCO').classList.remove('on');}}
function prio2pill(p){{
  const m={{'🔴 CRÍTICO':'p-crit','🟠 ALTO':'p-high','🟡 MEDIO':'p-med','🟢 BAJO':'p-low'}};
  return`<span class="prio-pill ${{m[p]||'p-low'}}">${{sevLabel(p)}}</span>`;
}}
const SEV_DOT={{'🔴':'#c62828','🟠':'#e65100','🟡':'#f57f17','🟢':'#2e7d32','💀':'#8B0000','⬜':'#9ca3af'}};
function sevLabel(t){{
  const s=String(t),i=s.indexOf(' ');
  if(i>0&&SEV_DOT[s.slice(0,i)])
    return'<span class="sev-dot" style="color:'+SEV_DOT[s.slice(0,i)]+'">'+ICO.dot+'</span>'+esc(s.slice(i+1));
  return esc(s);
}}
function esc(s){{return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');}}
document.addEventListener('keydown',e=>{{if(e.key==='Escape')closeModal();}});

// ── SCROLL TOP ────────────────────────────────────────────────────────────────
const mc=document.getElementById('MC');
mc.addEventListener('scroll',()=>{{document.getElementById('STB').classList.toggle('on',mc.scrollTop>400);}});
function stbClick(){{mc.scrollTo({{top:0,behavior:'smooth'}});}}

document.addEventListener('click',e=>{{
  if(!e.target.closest('.sa')&&!e.target.closest('#SRP'))
    document.getElementById('SRP').style.display='none';
}});

// Finalizar pantalla de carga
window._ldStopPulse&&window._ldStopPulse();
_storeOrigKpis();
try{{buildInvSidebar();}}catch(e){{console.warn('inv sidebar',e);}}
setTimeout(()=>{{try{{applyInvRowColors();}}catch(e){{}}}},400);
window._ldDone&&window._ldDone();
"""


def paso_generar_html(mes_es: str, datos: dict, diff_cves: dict, empresa: str = "", cfg: dict = None, cambios: list = None, reincidentes: list = None, tendencia_hist: list = None, riesgo_crem_manual: Optional[float] = None, cves_prioritarios: list = None, enrich_map: dict = None) -> Path:
    seccion("Generando HTML interactivo", 3, 4)

    # Initialized early — filled later, referenced in tbl() closure
    cve_por_activo: dict = {}

    # Inventario de activos desde config.json
    _inv = _get_inventario(cfg or {})
    inventario_js = json.dumps({
        k: {
            "desc":  v.get("descripcion", ""),
            "crit":  v.get("criticidad", "").upper(),
            "emoji": _CRIT_META.get(v.get("criticidad","").upper(), _CRIT_META[""])[0],
            "color": _CRIT_META.get(v.get("criticidad","").upper(), _CRIT_META[""])[1],
            "bg":    _CRIT_META.get(v.get("criticidad","").upper(), _CRIT_META[""])[2],
        }
        for k, v in _inv.items()
    }, ensure_ascii=False)

    def esc(s): return (str(s).replace("&","&amp;").replace("<","&lt;")
                        .replace(">","&gt;").replace('"',"&quot;"))
    def badge(nivel):
        n = (nivel or "").lower().strip()
        cfg = {"critical":("var(--crit)","var(--crit-bg)"),
               "high":("var(--high)","var(--high-bg)"),
               "medium":("var(--med)","var(--med-bg)"),
               "low":("var(--low)","var(--low-bg)")}
        fg, bg = cfg.get(n, ("var(--gray)","var(--lgray)"))
        return f'<span class="badge" style="color:{fg};background:{bg}">{esc(nivel or "-")}</span>'
    def score_pill(s):
        try: v = int(float(s))
        except Exception: return f"<span>{esc(s)}</span>"
        cls = "s-crit" if v>=80 else "s-high" if v>=60 else "s-med" if v>=40 else "s-low"
        return f'<span class="score-pill {cls}">{v}</span>'
    def prio_pill(p):
        cls = {"🔴 CRÍTICO":"p-crit","🟠 ALTO":"p-high","🟡 MEDIO":"p-med","🟢 BAJO":"p-low"}.get(p,"p-low")
        return f'<span class="prio-pill {cls}">{_sev_label(p)}</span>'
    def cve_link(cid):
        if str(cid).startswith("CVE-"):
            btn = f'<button class="sol-btn" onclick="return openSecHL(\'sec-cvesol\',\'{esc(cid)}\',this)" title="Ver solución en este informe">{_ico("wrench")}</button>'
            return f'<a href="{CVE_BASE_URL}{esc(cid)}" target="_blank" class="ext">{esc(cid)} {_ico("ext")}</a>{btn}'
        return esc(cid)

    nuevos_cve = diff_cves.get("nuevos", set())
    resueltos_cve = diff_cves.get("resueltos", set())

    # ── Nuevos datos para KPIs ejecutivos ─────────────────────────────────────
    _rein_list = reincidentes or []
    risk_data      = calcular_risk_score(datos, diff_cves, _rein_list, tendencia_hist)
    if riesgo_crem_manual is not None:
        risk_data = _aplicar_riesgo_externo(risk_data, riesgo_crem_manual)
    top3_incidentes = construir_top3_incidentes(datos)
    crem_panels    = construir_crem_panels(datos, _inv)

    # Asset cross-ref index
    asset_idx: dict[str, list] = {}
    def _aidx(asset, sid, lbl):
        a = str(asset).strip()
        if not a or a == "nan": return
        if a not in asset_idx: asset_idx[a] = []
        if sid not in [x[0] for x in asset_idx[a]]: asset_idx[a].append((sid, lbl))
    for key, sid, lbl in [
        ("cve_assets","sec-ca","CVE"),("sys_conf","sec-sys","Sys"),
        ("sec_conf","sec-sec","Sec"),("threats","sec-thr","Amenazas"),
        ("anomaly","sec-ano","Anomalías"),("cloud_app","sec-cld","Cloud"),
        ("accounts","sec-acc","Cuentas"),
    ]:
        df = datos.get(key, pd.DataFrame())
        if df.empty: continue
        col = "Device name" if key == "cve_assets" else "Asset"
        if "Activos" in df.columns:  # dedup
            for v in df["Activos"].dropna():
                for a in v.split(","):
                    _aidx(a.strip(), sid, lbl)
        elif col in df.columns:
            for v in df[col].dropna().unique(): _aidx(v, sid, lbl)

    # ── Coincidencia parcial bidireccional para inventario ────────────────────
    def _inv_lookup(name_s: str):
        """Match exacto o parcial: 'dsqlsrv01' coincide con 'dsqlsrv01.dominio.local' y viceversa."""
        name_l = name_s.lower()
        # Exacto
        if name_s in _inv: return name_s, _inv[name_s]
        # Parcial: clave del inventario contenida en el nombre del CSV o viceversa
        for k, v in _inv.items():
            kl = k.lower()
            if kl and (kl in name_l or name_l in kl):
                return k, v
        return None, None

    def abadges(name):
        name_s = str(name).strip()
        links = asset_idx.get(name_s, [])
        inv_key, inv_data = _inv_lookup(name_s)
        inv_html = ""
        if inv_data:
            crit = inv_data.get("criticidad", "").upper()
            emoji, color, bg, css = _CRIT_META.get(crit, _CRIT_META[""])
            desc_txt = esc(inv_data.get("descripcion", ""))
            inv_html = (
                f' <span class="inv-badge inv-{css}"' +
                f' data-crit="{esc(crit)}" data-key="{esc(inv_key or "")}"' +
                f' style="color:{color};background:{bg};border-color:{color}"' +
                f' title="{esc(inv_key or name_s)}: {desc_txt}">{emoji} {esc(crit) if crit else "INV"}</span>'
            )
        if not links:
            return inv_html
        name_j = esc(json.dumps(name_s))  # JSON HTML-safe para atributo onclick="…"
        refs = " ".join(
            f'<a href="#{s}" class="aref" title="Ver «{esc(name_s)}» en {esc(l)}" '
            f'onclick="return openSecHL(\'{s}\',{name_j},this)">{esc(l)}</a>'
            for s, l in links)
        return inv_html + " " + refs

    # Almacenar datos de activos para modal
    activos_json_data = {}
    df_plan_all = construir_plan(datos, diff_cves, enrich_map=enrich_map)
    if not df_plan_all.empty:
        for _, r in df_plan_all.iterrows():
            eq = str(r.get("Activo / Equipo",""))
            try:
                detalles = json.loads(r.get("_detalles_json","[]"))
            except Exception:
                detalles = []
            activos_json_data[eq] = {
                "prioridad": str(r.get("Prioridad","")),
                "modulos": str(r.get("Módulos afectados","")),
                "detalles": detalles,
            }

    # Umbral a partir del cual una tabla usa renderizado virtual (JS) en lugar de HTML estático
    VIRTUAL_THRESHOLD = 2000

    def tbl(tid, headers, rows_data, col_nivel=None, col_score=None, col_cve=None,
            col_asset=None, col_prio=None, ml_cols=None,
            cve_nuevos_set=None, exportable=True, row_assets=None):
        # row_assets: lista paralela a rows_data con el/los nombre(s) de activo COMPLETOS
        # (sin truncar) de cada fila. Se emiten como data-assets="…" para que el blink
        # de referencias cruzadas encuentre la fila aunque la celda visible esté truncada.
        ml = ml_cols or set()

        # CSV export (siempre, independiente del tamaño)
        csv_data_b64 = ""
        if exportable and rows_data:
            import base64, io
            buf = io.StringIO()
            buf.write(",".join(f'"{h}"' for h in headers) + "\n")
            for row in rows_data:
                buf.write(",".join(f'"{str(c).replace(chr(10),"|").replace(chr(13),"")}"'
                                   for c in row) + "\n")
            csv_data_b64 = base64.b64encode(buf.getvalue().encode("utf-8-sig")).decode()

        export_btn = (f'<button class="btn-exp" onclick="exportCSV(\'{tid}\',\'{csv_data_b64}\')" '
                      f'title="Exportar CSV">{_ico("download")} CSV</button>') if csv_data_b64 else ""

        # ── Tabla grande: renderizado virtual vía JS (SIN límite de filas) ──────
        if len(rows_data) > VIRTUAL_THRESHOLD:
            # Serializar TODAS las filas — el JS renderiza bajo demanda (scroll)
            raw_data = []
            for _ri, row in enumerate(rows_data):
                nv = (row[col_nivel] if col_nivel is not None and col_nivel < len(row) else "").lower()
                rc = {"critical":"r-crit","high":"r-high","medium":"r-med"}.get(nv, "")
                item = {"c": [str(c) for c in row], "r": rc}
                if row_assets and _ri < len(row_assets) and row_assets[_ri]:
                    item["a"] = str(row_assets[_ri]).lower()
                raw_data.append(item)
            data_json = json.dumps(raw_data, ensure_ascii=False)
            heads_json = json.dumps(headers, ensure_ascii=False)
            col_nivel_js  = col_nivel  if col_nivel  is not None else -1
            col_score_js  = col_score  if col_score  is not None else -1
            col_cve_js    = col_cve    if col_cve    is not None else -1
            col_prio_js   = col_prio   if col_prio   is not None else -1
            col_asset_js  = col_asset  if col_asset  is not None else -1
            total_label = f'{len(rows_data):,} filas — mostrando <span id="{tid}-showing">200</span> (scroll para más)'

            return f"""<div class="tbl-hdr">
  {export_btn}
  <span style="font-size:11px;color:var(--gray)">{total_label}</span>
  <input type="text" id="{tid}-filter" placeholder="Filtrar en esta tabla…"
    style="margin-left:8px;padding:3px 8px;border:1px solid var(--mgray);border-radius:4px;font-size:12px;width:200px"
    oninput="vtFilter('{tid}')">
</div>
<div class="tbl-wrap vt"><table id="{tid}" class="dtbl">
<thead><tr>{"".join(f"<th>{esc(h)}</th>" for h in headers)}</tr></thead>
<tbody id="{tid}-body"></tbody>
</table></div>
<script>
(function(){{
  var DATA={data_json};
  var filtered=DATA;
  var pageSize=200;
  var colNivel={col_nivel_js};
  var colScore={col_score_js};
  var colCve={col_cve_js};
  var colPrio={col_prio_js};
  var colAsset={col_asset_js};
  function jsAttr(v){{return JSON.stringify(v).replace(/"/g,'&quot;');}}
  function renderCell(v,ci){{
    if(ci===colCve&&v.startsWith('CVE-')) {{
      var btn='<button class="sol-btn" onclick="return openSecHL(\\'sec-cvesol\\',\\''+v+'\\',this)" title="Ver solución en este informe">'+ICO.wrench+'</button>';
      return'<a href="https://www.cve.org/CVERecord?id='+v+'" target="_blank" class="ext">'+v+' '+ICO.ext+'</a>'+btn;
    }}
    if(ci===colScore){{
      var s=parseInt(v)||0;
      var cls=s>=80?'s-crit':s>=60?'s-high':s>=40?'s-med':'s-low';
      return'<span class="score-pill '+cls+'">'+s+'</span>';
    }}
    if(ci===colNivel){{
      var n=v.toLowerCase();
      var cfg={{'critical':['var(--crit)','var(--crit-bg)'],'high':['var(--high)','var(--high-bg)'],
               'medium':['var(--med)','var(--med-bg)'],'low':['var(--low)','var(--low-bg)']}};
      var clr=cfg[n]||['var(--gray)','var(--lgray)'];
      return'<span class="badge" style="color:'+clr[0]+';background:'+clr[1]+'">'+v+'</span>';
    }}
    if(ci===colPrio){{
      var pm={{'🔴 CRÍTICO':'p-crit','🟠 ALTO':'p-high','🟡 MEDIO':'p-med','🟢 BAJO':'p-low'}};
      var sd={{'🔴':'#c62828','🟠':'#e65100','🟡':'#f57f17','🟢':'#2e7d32'}};
      var sp=v.indexOf(' '),em=sp>0?v.slice(0,sp):'',tx=sp>0?v.slice(sp+1):v;
      var dot=(sd[em]&&window.ICO)?'<span class="sev-dot" style="color:'+sd[em]+'">'+ICO.dot+'</span>':'';
      return'<span class="prio-pill '+(pm[v]||'p-low')+'">'+dot+tx.replace(/&/g,'&amp;').replace(/</g,'&lt;')+'</span>';
    }}
    var escaped = v.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
    if(ci===colAsset&&v){{
      // Mismos botones que la tabla estatica: detalle del activo y sus CVEs.
      var b='';
      if(typeof ACTIVOS!=='undefined'&&ACTIVOS[v]&&window.ICO)
        b+='<button class="modal-btn" onclick="showModal('+jsAttr(v)+')" title="Ver detalle del activo">'+ICO.search+'</button>';
      if(typeof CVE_ACTIVO!=='undefined'&&CVE_ACTIVO[v]&&window.ICO)
        b+='<button class="modal-btn cve-btn" onclick="showCVEModal('+jsAttr(v)+')" title="Ver todos los CVEs de este activo">'+ICO.cve+'</button>';
      return escaped+b;
    }}
    return escaped.replace(/\\n/g, '<br>');
  }}
  function renderRows(data, limit){{
    var tb=document.getElementById('{tid}-body');
    if(!tb)return;
    var n=Math.min(limit||pageSize, data.length);
    var html='';
    for(var i=0;i<n;i++){{
      var row=data[i];
      html+='<tr class="'+row.r+'"'+(row.a?' data-assets="'+row.a.replace(/"/g,'&quot;')+'"':'')+'>';
      for(var ci=0;ci<row.c.length;ci++)
        html+='<td>'+renderCell(row.c[ci],ci)+'</td>';
      html+='</tr>';
    }}
    tb.innerHTML=html;
    var sp=document.getElementById('{tid}-showing');
    if(sp)sp.textContent=n+(n<data.length?' de '+data.length:'');
  }}
  // Registro por tabla: cada tabla virtual filtra SUS propios datos
  window.__vt=window.__vt||{{}};
  window.__vt['{tid}']=function(){{
    var inp=document.getElementById('{tid}-filter');
    if(!inp)return;
    var q=inp.value.toLowerCase().trim();
    filtered=q?DATA.filter(function(r){{return r.c.some(function(c){{return c.toLowerCase().includes(q);}})||(r.a&&r.a.indexOf(q)>-1);}})
               :DATA;
    renderRows(filtered, pageSize);
  }};
  window.vtFilter=window.vtFilter||function(id){{
    var fn=window.__vt&&window.__vt[id];
    if(fn)fn();
  }};
  // Load more on scroll
  var wrap=document.getElementById('{tid}');
  if(wrap){{
    var par=wrap.closest('.tbl-wrap')||wrap.parentElement;
    if(par)par.addEventListener('scroll',function(){{
      if(par.scrollTop+par.clientHeight>=par.scrollHeight-50){{
        pageSize=Math.min(pageSize+200, filtered.length);
        renderRows(filtered, pageSize);
      }}
    }});
  }}
  // Initial render
  if(document.readyState==='loading')
    document.addEventListener('DOMContentLoaded',function(){{renderRows(filtered,pageSize);}});
  else renderRows(filtered, pageSize);
  // Re-render when section opens
  var sec=wrap&&wrap.closest('details.sec');
  if(sec)sec.addEventListener('toggle',function(){{if(sec.open)renderRows(filtered,pageSize);}});
}})();
</script>"""

        # ── Tabla normal: HTML estático ───────────────────────────────────────
        out = f'<div class="tbl-hdr">{export_btn}</div>'
        out += f'<div class="tbl-wrap"><table id="{tid}" class="dtbl">'
        out += '<thead><tr>' + "".join(f'<th>{esc(h)}</th>' for h in headers) + '</tr></thead><tbody>'
        for _ri, row in enumerate(rows_data):
            nv = (row[col_nivel] if col_nivel is not None and col_nivel<len(row) else "").lower()
            rc = {"critical":"r-crit","high":"r-high","medium":"r-med"}.get(nv,"")
            is_new = cve_nuevos_set and row[0] in cve_nuevos_set if cve_nuevos_set else False
            if is_new: rc += " r-new"
            _da = ""
            if row_assets and _ri < len(row_assets) and row_assets[_ri]:
                _da = f' data-assets="{esc(str(row_assets[_ri]).lower())}"'
            out += f'<tr class="{rc.strip()}"{_da}>'
            for ci, cell in enumerate(row):
                if col_cve is not None and ci == col_cve:
                    tag = cve_link(cell)
                    if is_new: tag = f'<span class="new-badge">NUEVO</span> {tag}'
                    out += f'<td>{tag}</td>'
                elif col_score is not None and ci == col_score:
                    out += f'<td>{score_pill(cell)}</td>'
                elif col_nivel is not None and ci == col_nivel:
                    out += f'<td>{badge(cell)}</td>'
                elif col_asset is not None and ci == col_asset:
                    modal_btn = (f'<button class="modal-btn" onclick="showModal({esc(json.dumps(cell))})" '
                                 f'title="Ver detalle del activo">{_ico("search")}</button>') if cell in activos_json_data else ""
                    cve_btn = (f'<button class="modal-btn cve-btn" onclick="showCVEModal({esc(json.dumps(cell))})" '
                               f'title="Ver todos los CVEs de este activo">{_ico("cve")}</button>') if cell in cve_por_activo else ""
                    out += f'<td>{esc(cell)}{abadges(cell)}{modal_btn}{cve_btn}</td>'
                elif col_prio is not None and ci == col_prio:
                    out += f'<td>{prio_pill(cell)}</td>'
                elif ci in ml:
                    out += f'<td class="ml-cell">{esc(cell).replace(chr(10),"<br>")}</td>'
                else:
                    out += f'<td>{esc(cell)}</td>'
            out += '</tr>'
        out += '</tbody></table></div>'
        return out

    # Visibilidad por módulo (de .api_meta.json). secs_ocultas se rellena dentro
    # de sec() y se consulta luego al construir la navegación.
    _modules    = datos.get("_modules", {}) or {}
    secs_ocultas: set = set()

    def sec(sid, icon, title, count, content, open_=False):
        # Extraer número para dinámica de contadores
        import re as _re
        num = _re.search(r'\d+', str(count))
        n_val = num.group(0) if num else "0"
        # Secciones ligadas a un módulo de datos: ocultar si no está contratado,
        # o mostrar "sin incidencias" si está contratado pero sin filas.
        _mods = SECCION_MODULOS.get(sid)
        if _mods and n_val == "0":
            if _seccion_aplica(_modules, sid):
                content = ('<div class="sec-empty-ok">✓ Sin incidencias detectadas '
                           'este mes para este módulo.</div>')
            else:
                secs_ocultas.add(sid)
                return ""   # módulo no contratado → no se renderiza la sección
        return (f'<details id="{sid}" {"open" if open_ else ""} class="sec" data-sid="{sid}" data-total="{n_val}">'
                f'<summary onclick="secClick(\'{sid}\')">'
                f'<span class="sl"><span class="sarr">▶</span><span class="sico">{_icoify(icon)}</span>'
                f'<span class="stit">{esc(title)}</span></span>'
                f'<span class="scnt" id="cnt-{sid}">{count}</span></summary>'
                f'<div class="sbody">{content}</div></details>')

    # ── Build sections ────────────────────────────────────────────────────────
    res_filas = construir_resumen(datos)
    total = sum(f["Total"] for f in res_filas)
    alto  = sum(f["Alto / Crítico"] for f in res_filas)
    medio = sum(f["Medio"] for f in res_filas)
    bajo  = sum(f["Bajo"] for f in res_filas)

    # Chart data
    chart_bar_labels = json.dumps([f["Módulo de Seguridad"] for f in res_filas])
    chart_bar_alto   = json.dumps([f["Alto / Crítico"] for f in res_filas])
    chart_bar_medio  = json.dumps([f["Medio"] for f in res_filas])
    chart_bar_bajo   = json.dumps([f["Bajo"] for f in res_filas])
    chart_donut_data = json.dumps([alto, medio, bajo])

    secs = []

    # Secciones a construir — para barra de progreso en consola
    _SEC_NAMES = ["CREM Panels","Resumen","CVE Eventos","CVE Activos","Config. Sistema",
                  "Config. Seguridad","Amenazas","Anomalías","Cloud","Cuentas",
                  "Analítica Predictiva","Cambios","Reincidentes","Tendencia","Plan Actuación"]
    _sec_prog = [0]  # contador mutable
    def _sec_done(name):
        _sec_prog[0] += 1
        ok(f"  HTML [{_sec_prog[0]:2d}/{len(_SEC_NAMES)}] {name}")

    # Diff banner
    diff_banner = ""
    if nuevos_cve or resueltos_cve:
        diff_banner = (f'<div class="diff-banner">'
                       f'<span class="diff-new">{_ico("new")} {len(nuevos_cve)} CVEs nuevos este mes</span>'
                       f'<span class="diff-res">{_ico("check")} {len(resueltos_cve)} CVEs resueltos</span>'
                       f'<span class="diff-src">vs. {esc(diff_cves.get("dir_anterior") or "mes anterior")}</span>'
                       f'</div>')

    # ── TOP 3 INCIDENTES (ACCIONES PRIORITARIAS) ─────────────────────────────
    def _inc_nivel_cls(n):
        return {"Critical":"r-crit","High":"r-high"}.get(n,"r-med")
    def _inc_badge(n):
        cfg2 = {"Critical":("var(--crit)","var(--crit-bg)"),
                "High":("var(--high)","var(--high-bg)"),
                "Medium":("var(--med)","var(--med-bg)")}
        fg2, bg2 = cfg2.get(n, ("var(--gray)","var(--lgray)"))
        return f'<span class="badge" style="color:{fg2};background:{bg2}">{esc(n)}</span>'

    if top3_incidentes:
        t3_items = ""
        for inc in top3_incidentes:
            wb_html = (f'<a href="{esc(inc["wb_link"])}" target="_blank" class="ext" '
                       f'title="{esc(inc["wb_id"])}">{_ico("ext")} WB</a>') if inc.get("wb_link") else (
                       f'<span class="t3-id">{esc(inc["wb_id"])}</span>' if inc.get("wb_id") else "")
            t3_items += f"""<div class="t3-card {_inc_nivel_cls(inc["nivel"])}">
  <div class="t3-hdr">
    <span class="t3-ico">{_icoify(inc["ico"])}</span>
    <span class="t3-cat">{esc(inc["cat"])}</span>
    {_inc_badge(inc["nivel"])}
    <span class="t3-date">{esc(inc["date"])}</span>
    <span class="t3-wb">{wb_html}</span>
  </div>
  <div class="t3-event">{esc(inc["event"])}</div>
  <div class="t3-asset">{_ico("device")} {esc(inc["asset"]) if inc["asset"] else "—"}</div>
</div>"""
        top3_html = f'<div class="t3-label">{_ico("target")} Acciones Prioritarias</div><div class="t3-row">{t3_items}</div>'
    else:
        # Este bloque solo mira amenazas, anomalías y cuentas: decir "sin
        # incidentes críticos/altos" contradecía el KPI de cabecera, que también
        # cuenta los hallazgos de configuración (sys-conf / sec-conf).
        top3_html = (f'<div class="sec-empty-ok">{_ico("check")} Sin amenazas, anomalías ni '
                     f'compromisos de cuenta de nivel crítico/alto este mes.</div>')

    # ── CREM 5 PANELS ────────────────────────────────────────────────────────
    def _panel_badge(n):
        m = {"Critical":("#c62828","#fdecea"),"High":("#e65100","#fff3e0"),
             "Medium":("#f57f17","#fffde7"),"Low":("#2e7d32","#f1f8e9")}
        fg3, bg3 = m.get(n, ("#6b7280","#f8f9fb"))
        return f'<span style="font-size:9.5px;font-weight:700;padding:1px 6px;border-radius:3px;color:{fg3};background:{bg3}">{esc(n)}</span>'

    def _crit_chip(crit):
        meta2 = _CRIT_META.get((crit or "").upper(), _CRIT_META[""])
        label2 = crit.title() if crit else "Sin catalogar"
        return f'<span style="font-size:9px;font-weight:700;padding:1px 5px;border-radius:3px;color:{meta2[1]};background:{meta2[2]}">{meta2[0]} {esc(label2)}</span>'

    def _score_mini(s):
        try: v = int(float(s))
        except Exception: return esc(str(s))
        cls2 = "s-crit" if v>=80 else "s-high" if v>=60 else "s-med" if v>=40 else "s-low"
        return f'<span class="score-pill {cls2}" style="font-size:11px">{v}</span>'

    def _build_panel(icon, title, sid_link, items_html):
        return (f'<div class="cp-panel">'
                f'<div class="cp-hdr"><span class="cp-ico">{_icoify(icon)}</span>'
                f'<span class="cp-title">{esc(title)}</span>'
                f'<a class="cp-link" href="#{sid_link}" onclick="openSec(\'{sid_link}\')" title="Ver sección">→</a></div>'
                f'<div class="cp-body">{items_html or "<div class=cp-empty>Sin datos este mes</div>"}</div>'
                f'</div>')

    def _cp_name(text, sid, extra=""):
        """Nombre del activo como enlace a su fila en la sección de detalle (con blink)."""
        t = str(text)
        tj = esc(json.dumps(t))  # JSON HTML-safe para atributo onclick="…"
        return (f'<div class="cp-name"><a class="cp-namelink" href="#{sid}" '
                f'title="Ver «{esc(t)}» en el detalle" '
                f'onclick="return openSecHL(\'{sid}\',{tj},this)">{esc(t)}</a>'
                f'{(" " + extra) if extra else ""}</div>')

    # Devices
    dev_html = ""
    for it in crem_panels.get("devices",[]):
        dev_html += (f'<div class="cp-item">'
                     f'{_cp_name(it["name"], "sec-ca", _crit_chip(it["crit"]))}'
                     f'<div class="cp-meta">{_score_mini(it["score"])} &nbsp; '
                     f'<span class="cp-tag">{esc(it["cves"])} CVEs</span>'
                     f'{"  <span class=cp-os>" + esc(it["os"][:22]) + "</span>" if it["os"] else ""}</div>'
                     f'</div>')

    # Internet
    inet_html = ""
    for it in crem_panels.get("internet",[]):
        inet_html += (f'<div class="cp-item">'
                      f'{_cp_name(it["name"], "sec-ca", _crit_chip(it["crit"]))}'
                      f'<div class="cp-meta">{_score_mini(it["score"])} &nbsp; '
                      f'<span class="cp-tag">{esc(it["ip"])}</span> &nbsp; '
                      f'<span class="cp-tag">{esc(it["cves"])} CVEs</span></div>'
                      f'</div>')

    # Accounts
    acc_html = ""
    for it in crem_panels.get("accounts",[]):
        acc_html += (f'<div class="cp-item">'
                     f'{_cp_name(it["account"], "sec-acc")}'
                     f'<div class="cp-meta">{_panel_badge(it["nivel"])} &nbsp; '
                     f'<span class="cp-os">{esc(it["event"][:40])}</span></div>'
                     f'</div>')

    # Applications
    app_html = ""
    for it in crem_panels.get("applications",[]):
        app_html += (f'<div class="cp-item">'
                     f'{_cp_name(it["app"], "sec-cve")}'
                     f'<div class="cp-meta">{_score_mini(it["max_score"])} &nbsp; '
                     f'<span class="cp-tag">{esc(it["cves"])} CVEs</span></div>'
                     f'</div>')

    # Cloud
    cld_html = ""
    for it in crem_panels.get("cloud",[]):
        cld_html += (f'<div class="cp-item">'
                     f'{_cp_name(it["app"], "sec-cld")}'
                     f'<div class="cp-meta">{_panel_badge(it["nivel"])} &nbsp; '
                     f'<span class="cp-os">{esc(it["asset"][:30])}</span></div>'
                     f'</div>')

    # Resumen de exposición por criticidad de inventario (activos con CVE Alto/Crítico, score>=60)
    def _crit_of_asset(name_s: str) -> str:
        n = str(name_s).lower()
        for k, v in _inv.items():
            kl = str(k).lower()
            if kl and (kl in n or n in kl):
                return v.get("criticidad","").upper()
        return ""

    crit_counts = {"MUY CRITICO":0, "CRITICO":0, "NO CRITICO":0, "":0}
    df_ca_sum = datos.get("cve_assets", pd.DataFrame())
    if not df_ca_sum.empty:
        _scores = pd.to_numeric(_serie(df_ca_sum, "CVE event risk score", "0"), errors="coerce").fillna(0)
        _names  = _serie(df_ca_sum, "Device name").astype(str)
        for nm, sc in zip(_names[_scores>=60], _scores[_scores>=60]):
            crit_counts[_crit_of_asset(nm)] = crit_counts.get(_crit_of_asset(nm),0) + 1

    crem_summary_html = '<div class="crem-summary">' + "".join(
        f'<div class="cs-item"><span class="cs-ico">{_CRIT_META[k][0]}</span>'
        f'<span class="cs-lbl">{esc(k.title() if k else "Sin catalogar")}</span>'
        f'<span class="cs-val" style="color:{_CRIT_META[k][1]}">{crit_counts[k]}</span></div>'
        for k in ("MUY CRITICO","CRITICO","NO CRITICO","")
    ) + '</div>'

    crem_html = (crem_summary_html +
                 f'<div class="crem-row">'
                 f'{_build_panel("💻","Dispositivos","sec-ca",dev_html)}'
                 f'{_build_panel("🌐","Internet / Expuesto","sec-ca",inet_html)}'
                 f'{_build_panel("👤","Cuentas","sec-acc",acc_html)}'
                 f'{_build_panel("📱","Aplicaciones","sec-cve",app_html)}'
                 f'{_build_panel("☁️","Cloud Assets","sec-cld",cld_html)}'
                 f'</div>')

    secs.append(sec("sec-crem","🔎","Vista CREM — TOP 3 por Dimensión","5 dimensiones",
                    crem_html, open_=True))
    _sec_done("CREM Panels")

    # ── Resumen + gráficos ────────────────────────────────────────────────────
    rows_res = [[f["Módulo de Seguridad"],str(f["Total"]),str(f["Alto / Crítico"]),
                 str(f["Medio"]),str(f["Bajo"])] for f in res_filas]
    charts_html = f"""
<div class="charts-row">
  <div class="chart-box">
    <div class="chart-title">Distribución de severidad</div>
    {_svg_donut(alto, medio, bajo)}
  </div>
  <div class="chart-box chart-box-wide">
    <div class="chart-title">Eventos por módulo</div>
    <div class="chart-svg">{_svg_stacked_bars(res_filas)}</div>
  </div>
</div>"""
    secs.append(sec("sec-res","📊","Resumen Ejecutivo",f"{total} eventos",
        top3_html + diff_banner + charts_html + tbl("t-res",["Módulo","Total","Alto/Crít.","Medio","Bajo"],
            rows_res, exportable=False), open_=True))
    _sec_done("Resumen")

    # CVE events
    df = datos.get("cve_events", pd.DataFrame())
    rows_cve = [[str(r.get("Vulnerability ID","")),str(r.get("CVE impact score","")),
                 str(r.get("Global exploit potential","")),str(r.get("OS/Application",""))[:45],
                 str(r.get("Impact scope","")),str(r.get("First seen time",""))[:10]]
                for _, r in df.iterrows()] if not df.empty else []
    cve_diff_note = ""
    if nuevos_cve or resueltos_cve:
        cve_diff_note = (f'<div class="note-info">'
                         f'<span class="diff-new">{_ico("new")} {len(nuevos_cve)} nuevos este mes (fondo lila)</span>'
                         f'{"  ·  " if resueltos_cve else ""}'
                         f'{"<span class=diff-res>" + _ico("check") + " " + str(len(resueltos_cve)) + " resueltos vs mes anterior</span>" if resueltos_cve else ""}'
                         f'</div>')
    secs.append(sec("sec-cve","🔓","Vulnerabilidades CVE",str(len(rows_cve)),
        cve_diff_note + tbl("t-cve",["CVE ID","Score","Explotabilidad","SO / Aplicación","Equipos","1ª Detección"],
            rows_cve, col_cve=0, col_score=1, cve_nuevos_set=nuevos_cve)))
    _sec_done("CVE Eventos")

    # CVE solutions — enriquecimiento NVD + CISA KEV + EPSS (versión que corrige)
    _cvesol = cves_prioritarios or []
    if _cvesol:
        _sol_rows = ""
        for it in _cvesol:
            _badges = ""
            if it.get("kev"):
                _badges += '<span class="badge b-err">Explotado</span> '
            _ep = it.get("epss", 0) or 0
            if _ep >= 0.5:
                _badges += f'<span class="badge b-warn">EPSS {_ep*100:.0f}%</span>'
            elif _ep > 0:
                _badges += f'<span class="badge b-gray">EPSS {_ep*100:.0f}%</span>'
            _cvss = it.get("cvss")
            _cvss_txt = f"{_cvss}" if _cvss is not None else str(it.get("score", "—"))
            raw_sol = str(it.get("solucion", "")).strip()
            if raw_sol:
                grouped_sols = agrupar_soluciones([raw_sol])
                _sol = "<br>".join(f"• {esc(s)}" for s in grouped_sols)
            else:
                _sol = "<span style='color:var(--gray)'>Ver aviso del fabricante</span>"
            _sol_rows += (f'<tr><td style="white-space:nowrap">{cve_link(it.get("id",""))} {_badges}</td>'
                          f'<td>{score_pill(_cvss_txt)}</td>'
                          f'<td class="num">{it.get("activos",0)}</td>'
                          f'<td>{esc(str(it.get("descripcion","")))}</td>'
                          f'<td style="font-weight:600;color:var(--dark)">{_sol}</td></tr>')
        _cvesol_html = (
            '<div class="note-info">Fuentes: NVD (NIST) · CISA KEV · EPSS (FIRST.org). '
            'La solución indica la versión que corrige la vulnerabilidad. '
            '<b>Explotado</b> = explotación activa conocida.</div>'
            '<div class="tbl-wrap" style="max-height: 480px; overflow-y: auto;"><table class="dtbl">'
            '<thead><tr><th>CVE</th><th>CVSS</th><th>Activos</th>'
            '<th>Descripción</th><th>Solución recomendada</th></tr></thead>'
            f'<tbody>{_sol_rows}</tbody></table></div>')
        secs.append(sec("sec-cvesol", "🔧", "CVEs Prioritarios · Solución",
                        str(len(_cvesol)), _cvesol_html))

    # CVE assets
    df = datos.get("cve_assets", pd.DataFrame())
    rows_ca = [[str(r.get("Device name","")),str(r.get("Operating system",""))[:35],
                str(r.get("IP address","")),str(r.get("CVE event risk score","")),
                str(r.get("Total CVEs","")),str(r.get("Average Unpatched Time (AUT)",""))[:6]]
               for _, r in df.iterrows()] if not df.empty else []
    assets_ca = [str(r.get("Device name","")) for _, r in df.iterrows()] if not df.empty else []
    secs.append(sec("sec-ca","💻","Activos con Mayor Exposición CVE",str(len(rows_ca)),
        tbl("t-ca",["Dispositivo","SO","IP","Score","Total CVEs","AUT (días)"],
            rows_ca, col_score=3, col_asset=0, row_assets=assets_ca)))
    _sec_done("CVE Activos")

    # Sys-conf
    df = datos.get("sys_conf", pd.DataFrame())
    rows_sys = []
    col_e2 = "Risk event"; col_n3 = "Nivel" if "Nivel" in df.columns else "Event risk level"
    col_a3 = "Activos" if "Activos" in df.columns else "Asset"
    n3_col = "n_activos" if "n_activos" in df.columns else None
    assets_sys = []
    for _, r in df.iterrows():
        det = _parse_detail(str(r.get("Detail",r.get("Detail info",""))))
        extra = det.get("osName", det.get("modulesNotOptimized",""))[:50]
        na = str(r.get(n3_col,"")) if n3_col else ""
        rows_sys.append([str(r.get(col_e2,""))[:65],str(r.get(col_a3,""))[:50],
                         str(r.get(col_n3,"")),str(r.get("Detectado","") or r.get("Detected",""))[:16],
                         extra or (f"{na} activos" if na else "")])
        assets_sys.append(str(r.get(col_a3,"")))  # lista de activos COMPLETA (sin truncar)
    secs.append(sec("sec-sys","⚙️","Configuración del Sistema",str(len(rows_sys)),
        tbl("t-sys",["Evento","Activos afectados","Nivel","Detectado","Detalle"],
            rows_sys, col_nivel=2, row_assets=assets_sys)))
    _sec_done("Config. Sistema")

    # Sec-conf
    df = datos.get("sec_conf", pd.DataFrame())
    rows_sec = []
    col_n4 = "Nivel" if "Nivel" in df.columns else "Event risk level"
    col_a4 = "Activos" if "Activos" in df.columns else "Asset"
    n4_col = "n_activos" if "n_activos" in df.columns else None
    assets_sec = []
    for _, r in df.iterrows():
        na = str(r.get(n4_col,"")) if n4_col else ""
        rows_sec.append([str(r.get(col_e2,""))[:65],str(r.get(col_a4,""))[:50],
                         str(r.get(col_n4,"")),str(r.get("Detectado","") or r.get("Detected",""))[:16],
                         f"{na} activos" if na else ""])
        assets_sec.append(str(r.get(col_a4,"")))  # lista de activos COMPLETA (sin truncar)
    secs.append(sec("sec-sec","🛡️","Configuración de Seguridad",str(len(rows_sec)),
        tbl("t-sec",["Evento","Activos afectados","Nivel","Detectado","Nº activos"],
            rows_sec, col_nivel=2, row_assets=assets_sec)))
    _sec_done("Config. Seguridad")

    # Threats
    df = datos.get("threats", pd.DataFrame())
    rows_thr = [[str(r.get("Risk event",""))[:65],str(r.get("Asset","")),
                 str(r.get("Event risk level","")),str(r.get("Detected",""))[:16],
                 _parse_detail(str(r.get("Detail info",""))).get("ruleName","")[:45]]
                for _, r in df.iterrows()] if not df.empty else []
    assets_thr = [str(r.get("Asset","")) for _, r in df.iterrows()] if not df.empty else []
    secs.append(sec("sec-thr","🚨","Detecciones de Amenazas",str(len(rows_thr)),
        tbl("t-thr",["Evento","Activo","Nivel","Detectado","Regla"],
            rows_thr, col_nivel=2, col_asset=1, row_assets=assets_thr)))
    _sec_done("Amenazas")

    # Anomaly
    df = datos.get("anomaly", pd.DataFrame())
    rows_ano = [[str(r.get("Risk event",""))[:65],
                 f"{r.get('Asset','')} ({r.get('Identity type','')})",
                 str(r.get("Event risk level","")),str(r.get("Detected",""))[:16],
                 str(r.get("Detail info",""))[:45]]
                for _, r in df.iterrows()] if not df.empty else []
    assets_ano = [str(r.get("Asset","")) for _, r in df.iterrows()] if not df.empty else []
    secs.append(sec("sec-ano","📡","Anomalías Detectadas",str(len(rows_ano)),
        tbl("t-ano",["Evento","Activo / Identidad","Nivel","Detectado","Detalle"],
            rows_ano, col_nivel=2, row_assets=assets_ano)))
    _sec_done("Anomalías")

    # Cloud
    df = datos.get("cloud_app", pd.DataFrame())
    rows_cld = [[str(r.get("_app",""))[:35],str(r.get("Asset","")),
                 _parse_detail(str(r.get("Detail info",""))).get("appCategory","")[:40],
                 str(r.get("Event risk level","")),str(r.get("Detected",""))[:16]]
                for _, r in df.iterrows()] if not df.empty else []
    assets_cld = [str(r.get("Asset","")) for _, r in df.iterrows()] if not df.empty else []
    secs.append(sec("sec-cld","☁️","Cloud Apps de Riesgo",str(len(rows_cld)),
        tbl("t-cld",["Aplicación","Equipo","Categoría","Nivel","Detectado"],
            rows_cld, col_nivel=3, col_asset=1, row_assets=assets_cld)))
    _sec_done("Cloud Apps")

    # Accounts
    df = datos.get("accounts", pd.DataFrame())
    rows_acc = [[str(r.get("Risk event",""))[:50],str(r.get("Impact scope",""))[:40],
                 str(r.get("Event risk level","")),str(r.get("Detected",""))[:16],
                 _parse_detail(str(r.get("Detail info",""))).get("ips","")[:40]]
                for _, r in df.iterrows()] if not df.empty else []
    assets_acc = [f'{r.get("Asset","")} {r.get("Impact scope","")}'
                  for _, r in df.iterrows()] if not df.empty else []
    secs.append(sec("sec-acc","👤","Compromiso de Cuentas",str(len(rows_acc)),
        tbl("t-acc",["Evento","Cuenta / Activo","Nivel","Detectado","IPs origen"],
            rows_acc, col_nivel=2, row_assets=assets_acc)))
    _sec_done("Cuentas")

    # Analítica Predictiva — rutas de ataque simuladas (ASM Attack Paths)
    df = datos.get("attack_paths", pd.DataFrame())
    rows_pred = []
    if not df.empty:
        for _, r in df.iterrows():
            # La exportación manual del portal fusiona ambos extremos en
            # «Entry / target assets» y detalla entryPoint/targetPoint en
            # «Detail info»; la extracción por API trae columnas separadas.
            det   = _parse_detail(str(r.get("Detail info","")))
            ambos = str(r.get("Entry / target assets","")).strip()
            entry = (str(r.get("Entry assets","")).strip()
                     or det.get("entryPoint","") or ambos)
            target = (str(r.get("Target assets","")).strip()
                      or det.get("targetPoint","") or ambos)
            rows_pred.append([
                str(r.get("Risk event",""))[:80],
                entry[:60],
                target[:60],
                str(r.get("Attack path risk score","")),
                str(r.get("Detected",""))[:16],
            ])
    pred_note = '<div class="note">ℹ️ Rutas de ataque simuladas por ASM (predictivo). Indica posibles vectores que un atacante podría aprovechar.</div>'
    secs.append(sec("sec-pred","🔮","Analítica Predictiva (Rutas de Ataque)",str(len(rows_pred)),
        pred_note + (tbl("t-pred",["Ruta de Ataque","Activos Entrada","Activos Objetivo","Score","Detectado"],
            rows_pred, col_score=3) if rows_pred else
        '<div class="note">Sin rutas de ataque detectadas. Requiere módulo ASM Attack Paths.</div>')))
    _sec_done("Analítica Predictiva")

    # Cambios mes a mes — vista comparativa enriquecida (Δ con %, críticos, nuevos/resueltos)
    if cambios:
        _dir_ant_lbl = esc(diff_cves.get("dir_anterior") or "mes anterior")
        def _chg_delta(c):
            if c["Mes anterior"] == "—":
                return '<span class="chg-pill chg-na">sin histórico</span>'
            v = c.get("_var_num", 0); pct = c.get("_pct", "")
            if v == 0:
                return '<span class="chg-pill chg-flat">= 0</span>'
            worse = v > 0  # más incidencias = peor
            cls = "chg-up" if worse else "chg-down"
            ic  = _ico("trend-up") if worse else _ico("trend-down")
            return f'<span class="chg-pill {cls}">{ic}{v:+d}{(" · " + esc(pct)) if pct else ""}</span>'
        def _chg_crit(c):
            if c["Mes anterior"] == "—":
                return '—'
            ca, cc = c.get("_crit_ant", 0), c.get("_crit_act", 0)
            cd = cc - ca
            cdcls = "chg-up" if cd > 0 else "chg-down" if cd < 0 else "chg-flat"
            return (f'<span class="chg-crit"><b>{ca}</b> → <b>{cc}</b> '
                    f'<span class="chg-cd {cdcls}">{cd:+d}</span></span>')
        _chg_body = "".join(
            f'<tr><td class="chg-mod">{esc(c["Módulo"])}</td>'
            f'<td class="chg-num">{c["Mes anterior"]}</td>'
            f'<td class="chg-num chg-actn">{c["Mes actual"]}</td>'
            f'<td>{_chg_delta(c)}</td>'
            f'<td>{_chg_crit(c)}</td>'
            f'<td class="chg-new">{("+" + c["Activos nuevos"]) if c["Activos nuevos"] != "—" else "—"}</td>'
            f'<td class="chg-res">{("−" + c["Activos resueltos"]) if c["Activos resueltos"] != "—" else "—"}</td></tr>'
            for c in cambios)
        _chg_html = (
            f'<div class="chg-note">{_ico("info")} Comparación con <b>{_dir_ant_lbl}</b>. '
            f'En incidencias, <span class="chg-down-t">menos = mejora</span> y '
            f'<span class="chg-up-t">más = empeora</span>. «Críticos» = eventos Critical/High (o CVE score ≥ 80).</div>'
            f'<div class="tbl-wrap"><table class="dtbl chg-tbl">'
            f'<thead><tr><th>Módulo</th><th>Anterior</th><th>Actual</th><th>Variación</th>'
            f'<th>Críticos (ant → act)</th><th>Nuevos</th><th>Resueltos</th></tr></thead>'
            f'<tbody>{_chg_body}</tbody></table></div>')
        secs.append(sec("sec-cambios","📈","Cambios vs Mes Anterior",f"{len(cambios)} módulos", _chg_html))
    _sec_done("Cambios")

    # Activos reincidentes
    rein = reincidentes or []
    rows_rein_h = [[r["activo"][:45],r["cve_id"],r["score"],r["meses"]+" mes(es)",r["accion"]] for r in rein]
    secs.append(sec("sec-rein","♻️","Activos Reincidentes",f"{len(rein)} CVEs sin resolver",
        '<div class="note">⚠️ CVEs presentes en 2 o más meses consecutivos sin resolver.</div>' +
        (tbl("t-rein",["Activo","CVE ID","Score","Meses sin resolver","Acción urgente"],
            rows_rein_h, col_score=2) if rein else '<div class="note">Sin reincidentes detectados este mes.</div>')))
    _sec_done("Reincidentes")

    # Tendencia mensual. El histórico solo llega hasta el mes anterior (los CSV
    # de este mes aún no están archivados), así que se añade el mes en curso para
    # que la franja de Riesgo CREM y el gráfico terminen en el dato de hoy.
    _sc_tend = pd.to_numeric(datos.get("cve_events", pd.DataFrame()).get("CVE impact score", pd.Series()),
                             errors="coerce").fillna(0)
    tend = list(tendencia_hist or []) + [{
        "mes":        mes_es.replace("_", " "),
        "cve":        len(datos.get("cve_events", pd.DataFrame())),
        "cve_crit":   int((_sc_tend >= 80).sum()),
        "threats":    len(datos.get("threats", pd.DataFrame())),
        "sys_issues": len(datos.get("sys_conf", pd.DataFrame())),
        "risk_score": risk_data.get("score"),
    }]
    if tend:
        tend_canvas = _build_tendencia_html(tend, mes_es)
    else:
        tend_canvas = '<div class="note">Sin datos históricos. Se generará al tener 2+ meses de informes.</div>'
    secs.append(sec("sec-tend","📉","Tendencia Mensual",f"{len(tend)} meses",tend_canvas))
    _sec_done("Tendencia")

    # Plan
    df_plan = construir_plan(datos, diff_cves, enrich_map=enrich_map)
    rows_plan = [[str(r.get("Activo / Equipo","")),str(r.get("Prioridad","")),
                  str(r.get("Problemas detectados","")),str(r.get("Acciones a realizar","")),
                  str(r.get("Módulos afectados",""))]
                 for _, r in df_plan.iterrows()] if not df_plan.empty else []
    plan_leg = (f'<div class="plan-leg">'
                f'<span class="prio-pill p-crit">{_sev_label("🔴 CRÍTICO")}</span> Inmediata &nbsp;'
                f'<span class="prio-pill p-high">{_sev_label("🟠 ALTO")}</span> 72h &nbsp;'
                f'<span class="prio-pill p-med">{_sev_label("🟡 MEDIO")}</span> 7 días &nbsp;'
                f'<span class="prio-pill p-low">{_sev_label("🟢 BAJO")}</span> Mantenimiento'
                f'</div>')
    secs.append(sec("sec-plan","🔧","Plan de Actuación por Equipo",f"{len(rows_plan)} activos",
        plan_leg + tbl("t-plan",["Activo / Equipo","Prioridad","Problemas detectados",
                                  "Acciones a realizar","Módulos afectados"],
            rows_plan, col_prio=1, col_asset=0, ml_cols={2,3}), open_=True))
    _sec_done("Plan Actuación")

    # CVE por activo — para popup "ver todos los CVEs".
    # El export CVE no trae vínculo directo equipo→CVE ("Impact scope" es un número,
    # el nº de activos afectados). Vinculamos por coincidencia de SISTEMA OPERATIVO:
    # el SO del activo (cve-assets) contra la lista "OS/Application" de cada CVE.
    df_cve_ev = datos.get("cve_events", pd.DataFrame())
    df_cve_as = datos.get("cve_assets", pd.DataFrame())

    if not df_cve_ev.empty and not df_cve_as.empty and "OS/Application" in df_cve_ev.columns:
        # Pre-normalizar los CVE una sola vez
        cve_ev_rows = [{
            "id":      str(rc.get("Vulnerability ID","")),
            "score":   str(rc.get("CVE impact score","")),
            "exploit": str(rc.get("Global exploit potential","")),
            "os":      str(rc.get("OS/Application","")),
            "os_l":    str(rc.get("OS/Application","")).lower(),
        } for _, rc in df_cve_ev.iterrows()]
        _sig_cache: dict = {}   # firma SO → lista de CVEs (muchos activos comparten SO)
        def _cves_for_sig(sig: str):
            if sig in _sig_cache:
                return _sig_cache[sig]
            lst = []
            if sig:
                seen = set()
                for e in cve_ev_rows:
                    if sig in e["os_l"] and e["id"] not in seen:
                        seen.add(e["id"])
                        lst.append({"id":e["id"],"score":e["score"],
                                    "exploit":e["exploit"],"os":e["os"]})
                lst.sort(key=lambda x: float(x["score"]) if x["score"].replace(".","").isdigit() else 0,
                         reverse=True)
            _sig_cache[sig] = lst
            return lst
        for _, ra in df_cve_as.iterrows():
            dev = str(ra.get("Device name","")).strip()
            if not dev: continue
            cve_por_activo[dev] = {"total": str(ra.get("Total CVEs","?")),
                                   "score": str(ra.get("CVE event risk score","?")),
                                   "os": str(ra.get("Operating system","")),
                                   "cves": _cves_for_sig(_os_sig(ra.get("Operating system","")))}
    cve_por_activo_json = json.dumps(cve_por_activo, ensure_ascii=False)

    # Search index — COMPLETO, sin límite de filas
    sidx = []
    rows_cambios_sidx = [[c["Módulo"],c["Mes anterior"],c["Mes actual"]] for c in (cambios or [])]
    _idx_sections = [("sec-cve",rows_cve),("sec-ca",rows_ca),("sec-sys",rows_sys),
                     ("sec-sec",rows_sec),("sec-thr",rows_thr),("sec-ano",rows_ano),
                     ("sec-cld",rows_cld),("sec-acc",rows_acc),("sec-plan",rows_plan),
                     ("sec-pred",rows_pred),
                     ("sec-cambios",rows_cambios_sidx),
                     ("sec-rein",[list(r.values()) for r in (reincidentes or [])])]
    _total_idx = sum(len(r) for _,r in _idx_sections)
    _done_idx = 0
    with _progress() as prog:
        task = prog.add_task("Construyendo índice de búsqueda…", total=_total_idx)
        for sid, rows_d in _idx_sections:
            for ri, row in enumerate(rows_d):
                sidx.append({"s":sid,"r":ri,"t":" ".join(str(c) for c in row).lower(),
                             "d":[_strip_sev(str(c))[:60] for c in row[:3]]})
            prog.advance(task, len(rows_d))
    sidx_json = json.dumps(sidx, ensure_ascii=False)
    ok(f"Índice de búsqueda: {len(sidx):,} entradas")
    activos_json = json.dumps(activos_json_data, ensure_ascii=False)

    # Referencias cruzadas por activo (para los hipervínculos del modal 🔍)
    def _refs_for(name: str):
        refs = asset_idx.get(name)
        if refs:
            return refs
        nl = name.lower()
        for k, v in asset_idx.items():
            kl = k.lower()
            if kl and (kl in nl or nl in kl):
                return v
        return []
    axref_json = json.dumps({a: _refs_for(a) for a in activos_json_data}, ensure_ascii=False)

    SLABELS = json.dumps({
        "sec-crem":"Vista CREM","sec-res":"Resumen","sec-cve":"CVE Eventos","sec-ca":"CVE Activos",
        "sec-sys":"Config. Sistema","sec-sec":"Config. Seg.","sec-thr":"Amenazas",
        "sec-ano":"Anomalías","sec-cld":"Cloud","sec-acc":"Cuentas",
        "sec-pred":"Analítica Predictiva","sec-plan":"Plan Actuación","sec-cambios":"Cambios","sec-rein":"Reincidentes","sec-tend":"Tendencia"
    })
    # risk gauge data for header
    _rs = risk_data.get("score", 0)
    _rs_disp = f"{_rs:.1f}" if float(_rs) != int(_rs) else str(int(_rs))
    _rn = risk_data.get("nivel", "—")
    _rd = risk_data.get("delta", 0)
    _rt = risk_data.get("trend", "—")
    _rc = risk_data.get("color", "#6b7280")
    _rg_arrow = "▲" if _rd > 0 else ("▼" if _rd < 0 else "●")
    _rg_trend_cls = "rg-worse" if _rd > 0 else ("rg-better" if _rd < 0 else "rg-same")
    _rg_trend_str = f"{_rg_arrow} {abs(int(_rd))} pts" if _rd != 0 else "Sin cambio"
    gen_date = datetime.now().strftime("%Y-%m-%d %H:%M")

    # Item de navegación: se omite si la sección está oculta (módulo no contratado)
    def _ni(sid, icon, label, n, on=False):
        if sid in secs_ocultas:
            return ""
        cls = "ni on" if on else "ni"
        return (f'<div class="{cls}" onclick="navTo(\'{sid}\',this)" data-t="{sid}">'
                f'<span class="nico">{_icoify(icon)}</span>{label}'
                f'<span class="ncnt">{n}</span></div>')

    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>CREM — Informe Seguridad {esc(mes_es)}</title>
<style>
{_css_tecnico()}
</style>
</head>
<body>

<!-- LOADING SCREEN -->
<div id="loading-screen">
  <div class="ld-logo"><span class="ld-mark">{_ico("shield")}</span><span class="ld-word">empresa</span></div>
  <div class="ld-pct" id="ld-pct">0%</div>
  <div class="ld-bar-wrap"><div class="ld-bar" id="ld-bar"></div></div>
  <div class="ld-title">Cargando informe de seguridad…</div>
  <div class="ld-step" id="ld-step">Iniciando…</div>
</div>

<script>
// Progress tracking — runs before rest of page renders
(function(){{
  var steps=[
    'Cargando estructura…',
    'Procesando vulnerabilidades CVE…',
    'Cargando configuración del sistema…',
    'Procesando detecciones de amenazas…',
    'Cargando anomalías y cloud…',
    'Construyendo plan de actuación…',
    'Renderizando tablas…',
    'Aplicando filtros y búsqueda…',
    'Listo',
  ];
  var step=0, pct=0;
  var bar=document.getElementById('ld-bar');
  var pctEl=document.getElementById('ld-pct');
  var stepEl=document.getElementById('ld-step');
  function setProgress(p,s){{
    pct=Math.min(p,99);
    if(bar){{bar.style.width=pct+'%';}}
    if(pctEl)pctEl.textContent=pct+'%';
    if(stepEl&&s)stepEl.textContent=s;
  }}
  // Advance progress as page parses (each script block fires this)
  window._ldStep=function(n){{
    var p=Math.round((n/steps.length)*95);
    setProgress(p, steps[n]||'');
  }};
  window._ldDone=function(){{
    setProgress(100,'¡Listo!');
    setTimeout(function(){{
      var ls=document.getElementById('loading-screen');
      if(ls)ls.classList.add('hidden');
      setTimeout(function(){{if(ls)ls.remove();}},500);
    }},300);
  }};
  // Pulse animation while waiting
  var pulse=setInterval(function(){{
    if(pct<95)setProgress(pct+0.3);
  }},120);
  window._ldStopPulse=function(){{clearInterval(pulse);}};
}})();
</script>

<div class="hdr"><div class="hdr-in">
  <button class="nav-toggle-btn" id="navToggleBtn" onclick="toggleNav()" title="Abrir navegación" aria-label="Abrir navegación">{_ico("list")}</button>
  <div class="logo">
    <span class="logo-mark">{_ico("shield")}</span>
    <span class="logo-word">empresa</span>
  </div>
  <div class="hdivider"></div>
  <div class="hinfo">
    <div class="htit">Informe de Seguridad{"&nbsp;— " + esc(empresa) if empresa else ""}</div>
    <div class="hsub">TrendAI CREM  ·  {esc(mes_es)}  ·  {gen_date}</div>
  </div>
  <div class="rg-pill" onclick="navTo('sec-crem',document.querySelector('[data-t=sec-crem]'))" title="Índice de Riesgo CREM — ver detalle">
    <div class="rg-score" style="color:{_rc}">{_rs_disp}</div>
    <div class="rg-right">
      <div class="rg-label">Riesgo CREM</div>
      <div class="rg-nivel" style="color:{_rc}">{esc(_rn)}</div>
      <div class="rg-trend {_rg_trend_cls}">{_rg_trend_str}</div>
    </div>
  </div>
  <div class="hdivider"></div>
  <div class="kpis">
    <div class="kpi k-tot"><div class="kn" id="kpi-tot">{total}</div><div class="kl">Total</div></div>
    <div class="kpi k-crit"><div class="kn" id="kpi-crit">{alto}</div><div class="kl">Alto/Crít.</div></div>
    <div class="kpi k-med"><div class="kn" id="kpi-med">{medio}</div><div class="kl">Medio</div></div>
    <div class="kpi k-low"><div class="kn" id="kpi-low">{bajo}</div><div class="kl">Bajo</div></div>
    <div class="kpi k-vis"><div class="kn" id="kpi-vis">{total}</div><div class="kl">Visibles</div></div>
  </div>
</div></div>

<div class="tb">
  <div class="tbg">
    <button class="btn" onclick="expAll()">▼ Todo</button>
    <button class="btn" onclick="colAll()">▲ Colapsar</button>
  </div>
  <div class="tbsep"></div>
  <span class="flbl">Severidad:</span>
  <div class="tbg">
    <span class="chip" onclick="tglF(this,'critical')"><span class="cdot" style="color:var(--crit)"></span>Crítico</span>
    <span class="chip" onclick="tglF(this,'high')"><span class="cdot" style="color:var(--high)"></span>Alto</span>
    <span class="chip" onclick="tglF(this,'medium')"><span class="cdot" style="color:var(--med)"></span>Medio</span>
    <span class="chip" onclick="tglF(this,'low')"><span class="cdot" style="color:var(--low)"></span>Bajo</span>
  </div>
  <div class="inv-sep"></div>
  <span class="flbl">Activo:</span>
  <div class="tbg" id="inv-chips">
    <span class="inv-chip" onclick="tglInv(this,'MUY CRITICO')">{_dot("#8B0000")} Muy Crítico</span>
    <span class="inv-chip" onclick="tglInv(this,'CRITICO')">{_dot("#c62828")} Crítico</span>
    <span class="inv-chip" onclick="tglInv(this,'NORMAL')">{_dot("#f57f17")} Normal</span>
    <span class="inv-chip" onclick="tglInv(this,'NO CRITICO')">{_dot("#2e7d32")} No Crítico</span>
    <span class="inv-chip" onclick="tglInv(this,'SIN_INV')">{_dot("#9ca3af")} Sin catalogar</span>
  </div>
  <span class="flt-info" id="flt-info">
    <span id="flt-info-txt">Filtro activo</span>
    <button class="flt-clear" onclick="clearAllFilters()" title="Limpiar filtros">✕</button>
  </span>
  <div class="sa">
    <div class="sw">
      <span class="si">{_ico("search")}</span>
      <input type="text" id="SI" placeholder="Buscar equipo, CVE, evento…"
             oninput="doSearch(this.value)" onkeydown="skd(event)">
    </div>
    <span class="sc" onclick="clrSearch()">✕</span>
  </div>
</div>

<div class="srp" id="SRP">
  <div class="srp-h"><span id="SRC">0 resultados</span><span style="color:var(--gray)">↵ navegar</span></div>
  <div id="SRL"></div>
</div>

<div class="lay" id="lay">
  <div class="nav-backdrop" id="navBackdrop" onclick="closeNav()"></div>
  <nav class="nav">
    <div class="navt">Secciones</div>
    <div class="ni on" onclick="navTo('sec-crem',this)" data-t="sec-crem"><span class="nico">{_ico("search")}</span>Vista CREM<span class="ncnt">5</span></div>
    <div class="ni" onclick="navTo('sec-res',this)" data-t="sec-res"><span class="nico">{_ico("chart")}</span>Resumen<span class="ncnt">{len(res_filas)}</span></div>
    {_ni("sec-cve","🔓","CVE Eventos",len(rows_cve))}
    {_ni("sec-cvesol","🔧","CVEs · Solución",len(cves_prioritarios or [])) if cves_prioritarios else ""}
    {_ni("sec-ca","💻","CVE Activos",len(rows_ca))}
    {_ni("sec-sys","⚙️","Config. Sistema",len(rows_sys))}
    {_ni("sec-sec","🛡️","Config. Seg.",len(rows_sec))}
    {_ni("sec-thr","🚨","Amenazas",len(rows_thr))}
    {_ni("sec-ano","📡","Anomalías",len(rows_ano))}
    {_ni("sec-cld","☁️","Cloud Apps",len(rows_cld))}
    {_ni("sec-acc","👤","Cuentas",len(rows_acc))}
    <div class="ni" onclick="navTo('sec-pred',this)" data-t="sec-pred"><span class="nico">{_ico("target")}</span>Analítica Predictiva<span class="ncnt">{len(rows_pred)}</span></div>
    {"" if not cambios else '<div class="ni" onclick="navTo(\'sec-cambios\',this)" data-t="sec-cambios"><span class="nico">'+_ico("trend-up")+'</span>Cambios<span class="ncnt">'+str(len(cambios or []))+'</span></div>'}
    <div class="ni" onclick="navTo('sec-rein',this)" data-t="sec-rein"><span class="nico">{_ico("recycle")}</span>Reincidentes<span class="ncnt">{len(reincidentes or [])}</span></div>
    <div class="ni" onclick="navTo('sec-tend',this)" data-t="sec-tend"><span class="nico">{_ico("trend-down")}</span>Tendencia<span class="ncnt">{len(tendencia_hist or [])}</span></div>
    <div class="ni" onclick="navTo('sec-plan',this)" data-t="sec-plan"><span class="nico">{_ico("wrench")}</span>Plan Actuación<span class="ncnt">{len(rows_plan)}</span></div>
  </nav>
  <main class="main" id="MC">
    {_html_degradaciones()}
    {''.join(secs)}
  </main>
</div>

<!-- MODAL CVE ACTIVO -->
<div class="modal-overlay" id="MCO" onclick="if(event.target===this)closeCVEModal()">
  <div class="modal" style="max-width:860px">
    <div class="modal-hdr" style="background:#1a3a5c">
      <span class="modal-title" id="MCOT">CVEs del activo</span>
      <button class="modal-close" onclick="closeCVEModal()">✕</button>
    </div>
    <div class="modal-body" id="MCOB"></div>
  </div>
</div>

<!-- MODAL DETALLE ACTIVO -->
<div class="modal-overlay" id="MO" onclick="if(event.target===this)closeModal()">
  <div class="modal">
    <div class="modal-hdr">
      <span class="modal-title" id="MT">Detalle del activo</span>
      <button class="modal-close" onclick="closeModal()">✕</button>
    </div>
    <div class="modal-body" id="MB"></div>
  </div>
</div>

<button class="stb" id="STB" onclick="stbClick()" title="Ir arriba">↑</button>

<script>
{_js_tecnico(activos_json=activos_json, axref_json=axref_json, cve_por_activo_json=cve_por_activo_json, inventario_js=inventario_js, sidx_json=sidx_json, slabels_json=SLABELS, chart_bar_labels=chart_bar_labels, chart_bar_alto=chart_bar_alto, chart_bar_medio=chart_bar_medio, chart_bar_bajo=chart_bar_bajo, chart_donut_data=chart_donut_data)}
</script>
</body></html>"""

    mes_safe = mes_es.replace("/","-").replace(" ","_")
    ruta = CTX.dir_informe / f"Revisión_CREM_{mes_safe}.html"
    ruta.write_text(html, encoding="utf-8")
    ok(f"HTML guardado → [bold]{ruta}[/]")
    return ruta

# ==============================================================================
# 14. EXCELS
# ==============================================================================
def paso_generar_excels(datos: dict):
    seccion("Generando Excel de revisión", 0, 0)
    CTX.dir_excels.mkdir(exist_ok=True)
    for clave, df in datos.items():
        if isinstance(df, pd.DataFrame) and not df.empty:
            ruta = CTX.dir_excels / f"crem-{clave}.xlsx"
            pub = df[[c for c in df.columns if not c.startswith("_")]]
            pub.to_excel(ruta, index=False, engine="openpyxl")
            wb = openpyxl.load_workbook(ruta); ws = wb.active
            for col in ws.columns:
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(
                    max((len(str(c.value or "")) for c in col), default=8) + 4, 60)
            wb.save(ruta)
            ok(f"{ruta.name:<30} {len(pub)} filas")

# ==============================================================================
# 14b. HTML EJECUTIVO
# ==============================================================================
def paso_generar_html_ejecutivo(
    mes_es: str, datos: dict, diff_cves: dict,
    empresa: str = "", cfg: dict = None,
    cambios: list = None, reincidentes: list = None,
    tendencia_hist: list = None,
    riesgo_crem_manual: Optional[float] = None,
    cves_prioritarios: list = None,
    enrich_map: dict = None,
) -> Path:
    seccion("Generando HTML ejecutivo", 3, 4)
    cfg = cfg or {}
    cambios = cambios or []
    reincidentes = reincidentes or []
    tendencia_hist = tendencia_hist or []
    cves_prioritarios = cves_prioritarios or []

    res_filas = construir_resumen(datos)
    plan_df   = construir_plan(datos, diff_cves, enrich_map=enrich_map)

    total_eventos     = sum(f["Total"] for f in res_filas)
    total_crit_alto   = sum(f["Alto / Crítico"] for f in res_filas)
    total_medio       = sum(f["Medio"] for f in res_filas)
    total_bajo        = sum(f["Bajo"] for f in res_filas)
    nuevos_cves       = len(diff_cves.get("nuevos", set()))
    resueltos_cves    = len(diff_cves.get("resueltos", set()))
    persistentes_cves = len(diff_cves.get("persistentes", set()))
    activos_riesgo    = len(datos.get("cve_assets", pd.DataFrame()))
    balance_cve       = resueltos_cves - nuevos_cves   # >0 = mejora neta

    # ── Riesgo CREM consolidado (0–100): el KPI ejecutivo principal ────────────
    risk_data = calcular_risk_score(datos, diff_cves, reincidentes, tendencia_hist)
    if riesgo_crem_manual is not None:
        risk_data = _aplicar_riesgo_externo(risk_data, riesgo_crem_manual)
    risk_score = risk_data.get("score", 0)
    risk_nivel = risk_data.get("nivel", "—")
    risk_color = risk_data.get("color", "#6b7280")
    risk_delta = risk_data.get("delta", 0)

    # ── Top-3 incidentes críticos (Workbench): el "qué ha pasado" concreto ─────
    top3_incidentes = construir_top3_incidentes(datos)

    # Top acciones críticas/altas del plan
    if not plan_df.empty and "Prioridad" in plan_df.columns:
        top_plan = plan_df[plan_df["Prioridad"].str.contains("CRÍTICO|ALTO", regex=True, na=False)].head(8)
        n_acc_crit = int(plan_df["Prioridad"].str.contains("CRÍTICO", na=False).sum())
        n_acc_alto = int(plan_df["Prioridad"].str.contains("ALTO", na=False).sum())
    else:
        top_plan = pd.DataFrame()
        n_acc_crit = n_acc_alto = 0

    # ── Narrativa ejecutiva auto-generada (situación + foco recomendado) ───────
    def _delta_frase():
        if not tendencia_hist or risk_data.get("prev_score") is None:
            return ""
        if risk_delta > 5:
            return f" El nivel de riesgo ha <b>empeorado {abs(int(risk_delta))} puntos</b> respecto al mes anterior."
        if risk_delta < -5:
            return f" El nivel de riesgo ha <b>mejorado {abs(int(risk_delta))} puntos</b> respecto al mes anterior."
        return " El nivel de riesgo se mantiene <b>estable</b> respecto al mes anterior."

    if risk_score >= 75:
        _cab = "La exposición al riesgo es <b>crítica</b> y requiere actuación inmediata."
    elif risk_score >= 50:
        _cab = "La exposición al riesgo es <b>alta</b> y exige atención prioritaria."
    elif risk_score >= 25:
        _cab = "La exposición al riesgo es <b>moderada</b> y bajo control con las acciones previstas."
    elif n_acc_crit > 0:
        _cab = "El nivel de riesgo global es <b>bajo</b>, si bien persisten acciones críticas puntuales que conviene cerrar."
    else:
        _cab = "La postura de seguridad es <b>saludable</b>, sin exposición crítica relevante."

    _cuerpo = (f" En el período de <b>{mes_es}</b> se registraron <b>{total_eventos}</b> eventos, "
               f"de los cuales <b>{total_crit_alto}</b> son de severidad crítica/alta.")
    if nuevos_cves or resueltos_cves:
        if balance_cve > 0:
            _cuerpo += f" Se resolvieron <b>{resueltos_cves}</b> vulnerabilidades frente a <b>{nuevos_cves}</b> nuevas (balance neto <b>favorable</b>: {balance_cve:+d})."
        elif balance_cve < 0:
            _cuerpo += f" Aparecieron <b>{nuevos_cves}</b> vulnerabilidades nuevas frente a <b>{resueltos_cves}</b> resueltas (balance neto <b>desfavorable</b>: {balance_cve:+d})."
        else:
            _cuerpo += f" El balance de vulnerabilidades es neutro (<b>{nuevos_cves}</b> nuevas y <b>{resueltos_cves}</b> resueltas)."
    if reincidentes:
        _cuerpo += f" Persisten <b>{len(reincidentes)}</b> CVEs reincidentes sin resolver."
    if n_acc_crit or n_acc_alto:
        _cuerpo += f" Hay <b>{n_acc_crit + n_acc_alto}</b> acciones prioritarias pendientes ({n_acc_crit} críticas, {n_acc_alto} altas)."
    narrativa_html = _cab + _cuerpo + _delta_frase()

    notas = cfg.get("notas_adicionales", "")
    contacto = cfg.get("contacto_tecnico", "")
    gen_ts = datetime.now().strftime("%d/%m/%Y %H:%M")

    def esc(s): return (str(s).replace("&","&amp;").replace("<","&lt;")
                        .replace(">","&gt;").replace('"',"&quot;"))

    # ── Gráfico de distribución de riesgo por módulo: SVG autocontenido ──
    risk_svg = _svg_stacked_bars(res_filas)

    # ── Gauge Riesgo CREM + sparkline tendencia del score ─────────────────────
    gauge_svg = _svg_risk_gauge(risk_score, risk_color, risk_nivel)
    # El histórico solo contiene meses ya archivados: el último punto debe ser
    # el score de este mes, que es el que se rotula como «Actual» al lado.
    _spark_vals = [t.get("risk_score") for t in tendencia_hist] + [risk_score]
    spark_svg = _svg_sparkline(_spark_vals, color=risk_color)
    if risk_delta > 5:
        risk_trend_html = f'<span class="rk-worse">▲ +{abs(int(risk_delta))} pts vs. mes anterior</span>'
    elif risk_delta < -5:
        risk_trend_html = f'<span class="rk-better">▼ {abs(int(risk_delta))} pts vs. mes anterior</span>'
    elif risk_data.get("prev_score") is not None:
        risk_trend_html = '<span class="rk-same">● Estable vs. mes anterior</span>'
    else:
        risk_trend_html = '<span class="rk-same">Sin histórico previo</span>'

    # ── Top-3 incidentes críticos (tarjetas) ──────────────────────────────────
    def _t3_cls(nivel):
        n = (nivel or "").lower()
        return "r-crit" if n == "critical" else ("r-high" if n == "high" else "r-med")
    top3_html = ""
    if top3_incidentes:
        for it in top3_incidentes:
            wb = ""
            if it.get("wb_link"):
                wb = f'<a href="{esc(it["wb_link"])}" target="_blank" class="t3-wb">↗ {esc(it.get("wb_id","Workbench"))}</a>'
            elif it.get("wb_id"):
                wb = f'<span class="t3-wb">{esc(it["wb_id"])}</span>'
            top3_html += f"""
        <div class="t3-card {_t3_cls(it.get('nivel'))}">
          <div class="t3-hdr">
            <span class="t3-cat">{esc(it.get('cat',''))}</span>
            <span class="prio-pill {'p-crit' if (it.get('nivel','').lower()=='critical') else 'p-alto'}">{esc(it.get('nivel','—'))}</span>
            <span class="t3-date">{esc(it.get('date',''))}</span>
          </div>
          <div class="t3-event">{esc(it.get('event','—'))}</div>
          <div class="t3-foot"><span class="t3-asset">{esc(it.get('asset','—'))}</span>{wb}</div>
        </div>"""
    else:
        # Ídem: el bloque no cubre los hallazgos de configuración, que sí suman
        # en el KPI «Alto/Crít.» de la cabecera.
        top3_html = ('<div class="t3-empty">Sin amenazas, anomalías ni compromisos de cuenta '
                     'de nivel crítico/alto este mes ✓</div>')

    # ── CVEs prioritarios con solución detallada (NVD + KEV + EPSS) ────────────
    cve_prio_rows = ""
    for it in cves_prioritarios[:15]:
        cid = esc(it.get("id", ""))
        cve_a = (f'<a href="https://www.cve.org/CVERecord?id={cid}" target="_blank">{cid}</a>')
        badges = ""
        if it.get("kev"):
            badges += '<span class="cve-badge kev">★ Explotado</span>'
        ep = it.get("epss", 0) or 0
        if ep >= 0.5:
            badges += f'<span class="cve-badge epss-hi">EPSS {ep*100:.0f}%</span>'
        elif ep > 0:
            badges += f'<span class="cve-badge epss">EPSS {ep*100:.0f}%</span>'
        cvss = it.get("cvss")
        cvss_txt = f'{cvss}' if cvss is not None else str(it.get("score", "—"))
        sev = it.get("severidad", "")
        sev_cls = {"CRITICAL": "p-crit", "HIGH": "p-alto"}.get(sev, "")
        raw_sol = str(it.get("solucion", "")).strip()
        if raw_sol:
            grouped_sols = agrupar_soluciones([raw_sol])
            sol = "<br>".join(f"• {esc(s)}" for s in grouped_sols)
        else:
            sol = "<span style='color:var(--t3)'>Ver aviso del fabricante</span>"
        act = it.get("activos", 0)
        cve_prio_rows += f"""
        <tr>
          <td class="bold" style="white-space:nowrap">{cve_a}{(' ' + badges) if badges else ''}</td>
          <td class="num"><span class="cve-score {sev_cls}">{esc(cvss_txt)}</span></td>
          <td class="num">{act}</td>
          <td style="color:var(--t2)">{esc(it.get('descripcion',''))}</td>
          <td class="cve-sol">{sol}</td>
        </tr>"""
    cves_prio_html = f"""
  <div class="section-title">{_ico("cve")} CVEs prioritarios · solución recomendada</div>
  <div class="card">
    <div class="card-body" style="padding:0">
      <table>
        <thead><tr><th>CVE</th><th>CVSS</th><th>Activos</th><th>Descripción</th><th>Solución recomendada</th></tr></thead>
        <tbody>{cve_prio_rows}</tbody>
      </table>
    </div>
  </div>
  <div class="cve-src">Fuentes: NVD (NIST) · CISA KEV · EPSS (FIRST.org). ★ = explotación activa conocida.</div>
""" if cve_prio_rows else ""

    # ── Filas tabla módulos ───────────────────────────────────────────────────
    # El resumen y la comparativa mensual nombran distinto el mismo módulo
    # ("Detecciones Amenaza" vs "Amenazas"), así que se emparejan por id de
    # módulo. Con la búsqueda por subcadena que había antes, esa fila se quedaba
    # sin tendencia («—») aunque sí hubiera variación.
    _cambios_por_id = {}
    for _c in cambios:
        _mc = MODULO_POR_ETIQUETA.get(_c.get("Módulo","").lower())
        if _mc:
            _cambios_por_id[_mc.id] = _c
    def _tendencia_icon(mod):
        m = MODULO_POR_ETIQUETA.get(mod.lower())
        c = _cambios_por_id.get(m.id) if m else None
        if c is None:
            return '<span class="tend-eq">—</span>'
        v = c.get("_var_num", 0)
        if   v > 0:  return f'<span class="tend-up">▲ +{v}</span>'
        elif v < 0:  return f'<span class="tend-dn">▼ {v}</span>'
        return '<span class="tend-eq">● =</span>'

    mod_rows = ""
    for f in res_filas:
        mod = esc(f["Módulo de Seguridad"])
        total = f["Total"]
        ca    = f["Alto / Crítico"]
        med   = f["Medio"]
        low   = f["Bajo"]
        risk_cls = "risk-crit" if ca > 0 else ("risk-med" if med > 0 else "risk-ok")
        tend = _tendencia_icon(f["Módulo de Seguridad"])
        mod_rows += f"""
        <tr>
          <td><span class="mod-name">{mod}</span></td>
          <td class="num">{total}</td>
          <td class="num {risk_cls} bold">{ca}</td>
          <td class="num">{med}</td>
          <td class="num">{low}</td>
          <td>{tend}</td>
        </tr>"""

    # ── Filas plan ejecutivo ──────────────────────────────────────────────────
    plan_rows = ""
    if not top_plan.empty:
        for _, row in top_plan.iterrows():
            p = str(row.get("Prioridad",""))
            p_cls = "p-crit" if "CRÍTICO" in p else "p-alto"
            activo  = esc(str(row.get("Activo / Equipo","—")))
            problem = esc(resumir_lineas(row.get("Problemas detectados","—")))
            accion  = esc(resumir_lineas(row.get("Acciones a realizar","—")))
            modulo  = esc(str(row.get("Módulos afectados","—")))
            if "CRÍTICO" in p:
                sla = f'{cfg.get("sla_critico_dias", 1)}d'
            elif "ALTO" in p:
                sla = f'{cfg.get("sla_alto_dias", 3)}d'
            else:
                sla = f'{cfg.get("sla_medio_dias", 7)}d'
            plan_rows += f"""
        <tr>
          <td><span class="prio-pill {p_cls}">{_sev_label(p)}</span></td>
          <td class="bold">{activo}</td>
          <td style="white-space:pre-wrap">{problem}</td>
          <td style="white-space:pre-wrap">{accion}</td>
          <td class="num">{sla}</td>
          <td><span class="mod-tag">{modulo}</span></td>
        </tr>"""
    else:
        plan_rows = '<tr><td colspan="6" class="empty-row">Sin acciones de alta prioridad este mes ✓</td></tr>'

    # ── Cambios mes a mes ─────────────────────────────────────────────────────
    cambios_html = ""
    if cambios:
        for c in cambios[:7]:
            v = c.get("_var_num", 0); pct = c.get("_pct", "")
            suf = f" · {esc(pct)}" if pct else ""
            if c.get("Mes anterior") == "—":
                badge = '<span class="chg-eq">nuevo</span>'
            elif v > 0:
                badge = f'<span class="chg-up">▲ +{v}{suf}</span>'
            elif v < 0:
                badge = f'<span class="chg-dn">▼ {v}{suf}</span>'
            else:
                badge = '<span class="chg-eq">= 0</span>'
            cambios_html += (f'<div class="chg-row"><span class="chg-mod">{esc(c.get("Módulo",""))}</span>'
                             f'<span class="chg-vals">{esc(str(c.get("Mes anterior","—")))} → '
                             f'<b>{esc(str(c.get("Mes actual","—")))}</b></span>{badge}</div>')
    else:
        cambios_html = '<div class="chg-row" style="color:#888">Sin datos de comparación mensual</div>'

    # ── Reincidentes ─────────────────────────────────────────────────────────
    reinc_html = ""
    if reincidentes:
        for r in reincidentes[:5]:
            if isinstance(r, str):
                cve_id = r
                label = r
            else:
                cve_id = r.get("cve_id", "CVE-?")
                score = r.get("score", "?")
                meses = r.get("meses", "?")
                activo = r.get("activo", "?")
                label = f"<b>{cve_id}</b> (Score: {score} · {activo} activo{'s' if activo != '1' else ''} · {meses} meses sin resolver)"
            reinc_html += f'<div class="reinc-item"><a href="https://www.cve.org/CVERecord?id={esc(cve_id)}" target="_blank">↗ {label}</a></div>'
    else:
        reinc_html = '<div class="reinc-item ok">Sin CVEs reincidentes ✓</div>'

    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Informe Ejecutivo CREM — {esc(empresa)} — {esc(mes_es)}</title>
<style>
{_css_ejecutivo()}
</style>
</head>
<body>

<div class="hdr"><div class="hdr-in">
  <div class="logo">
    <span class="logo-mark">{_ico("shield")}</span>
    <span class="logo-word">empresa</span>
    <span class="eje-badge">Ejecutivo</span>
  </div>
  <div class="hdivider"></div>
  <div class="hinfo">
    <div class="htit">Informe de Seguridad{"&nbsp;— " + esc(empresa) if empresa else ""}</div>
    <div class="hsub">TrendAI CREM · Revisión CREM · {esc(mes_es)} · {gen_ts}</div>
  </div>
  <div class="rg-pill" title="Índice de Riesgo CREM (0–100)">
    <div class="rg-score" style="color:{risk_color}">{(f"{risk_score:.1f}" if risk_score != int(risk_score) else int(risk_score))}</div>
    <div class="rg-right">
      <div class="rg-label">Riesgo CREM</div>
      <div class="rg-nivel" style="color:{risk_color}">{esc(risk_nivel)}</div>
      <div class="rg-trend {'rg-worse' if risk_delta > 5 else ('rg-better' if risk_delta < -5 else 'rg-same')}">{f'▲ +{abs(int(risk_delta))} pts' if risk_delta > 5 else (f'▼ {abs(int(risk_delta))} pts' if risk_delta < -5 else 'Estable')}</div>
    </div>
  </div>
  <div class="hdivider"></div>
  <div class="kpis">
    <div class="kpi k-tot"><div class="kn">{total_eventos}</div><div class="kl">Total</div></div>
    <div class="kpi k-crit"><div class="kn">{total_crit_alto}</div><div class="kl">Alto/Crít.</div></div>
    <div class="kpi k-med"><div class="kn">{total_medio}</div><div class="kl">Medio</div></div>
    <div class="kpi k-low"><div class="kn">{total_bajo}</div><div class="kl">Bajo</div></div>
  </div>
</div></div>

<div class="wrap">
  {_html_degradaciones()}
  <div class="section-title">{_ico("chart")} Situación de seguridad</div>
  <div class="hero">
    <div class="hero-gauge">
      <div class="gauge-cap">Índice de Riesgo CREM</div>
      {gauge_svg}
      <div class="hero-trend">{risk_trend_html}</div>
    </div>
    <div class="hero-body">
      <div class="hero-verdict">{narrativa_html}</div>
      <div class="hero-chips">
        <span class="hero-chip"><b>{total_eventos}</b> eventos</span>
        <span class="hero-chip {'crit' if total_crit_alto else 'ok'}"><b>{total_crit_alto}</b> crítico/alto</span>
        <span class="hero-chip {'crit' if balance_cve < 0 else 'ok'}">Balance CVE <b>{balance_cve:+d}</b></span>
        <span class="hero-chip {'crit' if reincidentes else 'ok'}"><b>{len(reincidentes)}</b> reincidentes</span>
        <span class="hero-chip {'crit' if (n_acc_crit + n_acc_alto) else 'ok'}"><b>{n_acc_crit + n_acc_alto}</b> acciones prioritarias</span>
      </div>
    </div>
  </div>

  <div class="kpi-grid">
    <div class="kpi">
      <div class="kpi-val {'red' if total_crit_alto > 0 else ''}">{total_eventos}</div>
      <div class="kpi-label">Total eventos</div>
      <div class="kpi-sub">{total_crit_alto} crítico/alto · {total_medio} medio</div>
    </div>
    <div class="kpi">
      <div class="kpi-val {'red' if nuevos_cves > 0 else 'green'}">{nuevos_cves}</div>
      <div class="kpi-label">CVEs nuevos</div>
      <div class="kpi-sub">vs. mes anterior</div>
    </div>
    <div class="kpi">
      <div class="kpi-val green">{resueltos_cves}</div>
      <div class="kpi-label">CVEs resueltos</div>
      <div class="kpi-sub">{persistentes_cves} persistentes</div>
    </div>
    <div class="kpi">
      <div class="kpi-val {'amber' if activos_riesgo > 0 else 'green'}">{activos_riesgo}</div>
      <div class="kpi-label">Activos con riesgo</div>
      <div class="kpi-sub">con CVEs activos</div>
    </div>
    <div class="kpi">
      <div class="kpi-val {'amber' if reincidentes else 'green'}">{len(reincidentes)}</div>
      <div class="kpi-label">CVEs reincidentes</div>
      <div class="kpi-sub">{cfg.get('meses_reincidente',2)}+ meses sin resolver</div>
    </div>
  </div>

  {"<div class='cve-diff'>" +
    (f"<span class='cve-pill cve-new'>▲ {nuevos_cves} nuevos</span>" if nuevos_cves else "") +
    (f"<span class='cve-pill cve-res'>▼ {resueltos_cves} resueltos</span>" if resueltos_cves else "") +
    (f"<span class='cve-pill cve-per'>● {persistentes_cves} persistentes</span>" if persistentes_cves else "") +
    "</div>" if (nuevos_cves or resueltos_cves or persistentes_cves) else ""}

  <div class="section-title">{_ico("alert")} Incidentes destacados del mes</div>
  <div class="t3-grid">{top3_html}</div>
{cves_prio_html}

  <div class="section-title">{_ico("list")} Riesgo por módulo</div>
  <div class="two-col">
    <div class="card" style="margin-bottom:0">
      <div class="card-hdr">Tabla de módulos</div>
      <div class="card-body" style="padding:0">
        <table>
          <thead>
            <tr>
              <th>Módulo</th><th>Total</th><th>Crítico/Alto</th><th>Medio</th><th>Bajo</th><th>Tendencia</th>
            </tr>
          </thead>
          <tbody>{mod_rows}</tbody>
        </table>
      </div>
    </div>
    <div class="card" style="margin-bottom:0">
      <div class="card-hdr">Distribución de riesgo por módulo</div>
      <div class="card-body">
        <div class="chart-svg">{risk_svg}</div>
      </div>
    </div>
  </div>

  <div class="section-title">{_ico("target")} Acciones prioritarias (Crítico / Alto)</div>
  <div class="card">
    <div class="card-body" style="padding:0">
      <table>
        <thead>
          <tr><th>Prioridad</th><th>Activo</th><th>Problema</th><th>Acción recomendada</th><th>SLA</th><th>Módulo</th></tr>
        </thead>
        <tbody>{plan_rows}</tbody>
      </table>
    </div>
  </div>

  <div class="section-title">{_ico("trend-up")} Comparativa mensual</div>
  {f'''<div class="card">
    <div class="card-hdr">Evolución del Riesgo CREM ({len([v for v in _spark_vals if v is not None])} meses)</div>
    <div class="card-body">
      <div class="spark-box">
        <div class="spark-svg">{spark_svg}</div>
        <div class="spark-meta">Actual<b style="color:{risk_color}">{(f"{risk_score:.1f}" if risk_score != int(risk_score) else int(risk_score))}</b></div>
      </div>
    </div>
  </div>''' if spark_svg else ""}
  <div class="two-col">
    <div class="card" style="margin-bottom:0">
      <div class="card-hdr">Variación vs. mes anterior</div>
      <div class="card-body">{cambios_html}</div>
    </div>
    <div class="card" style="margin-bottom:0">
      <div class="card-hdr">CVEs reincidentes ({cfg.get('meses_reincidente',2)}+ meses)</div>
      <div class="card-body">{reinc_html}</div>
    </div>
  </div>

  {f'''<div class="section-title">{_ico("note")} Notas adicionales</div>
  <div class="card">
    <div class="card-body">
      <div class="notes-box">{esc(notas)}</div>
    </div>
  </div>''' if notas.strip() else ""}

  {f'''<div class="section-title">{_ico("user")} Contacto técnico</div>
  <div class="card">
    <div class="card-body" style="font-size:13px;color:var(--t2)">{esc(contacto)}</div>
  </div>''' if contacto.strip() else ""}

</div><!-- /wrap -->

<div class="footer">
  EMPRESA · Informe ejecutivo CREM · {esc(empresa)} · {esc(mes_es)} · Generado {gen_ts}
</div>
</body></html>"""

    mes_safe = mes_es.replace("/","-").replace(" ","_")
    ruta = CTX.dir_informe / f"Revisión_CREM_{mes_safe}_ejecutivo.html"
    ruta.write_text(html, encoding="utf-8")
    ok(f"HTML ejecutivo guardado → [bold]{ruta}[/]")
    return ruta

# ==============================================================================
# 15. RIESGO CREM MANUAL
# ==============================================================================
def _obtener_riesgo_crem_api(empresa_dir: Optional[Path]) -> Optional[float]:
    """
    Consulta el Cyber Risk Index a la API de Vision One con el .env de la empresa.

    Distingue dos situaciones que antes se confundían en el mismo aviso:
      · la clave no está puesta  → es configuración, no un fallo. Informativo.
      · la clave está y falla    → eso sí es una degradación del informe.
    Sin esa distinción salía un WARNING todos los meses en clientes que
    deliberadamente no usan la API, y el ruido acaba tapando los avisos reales.
    """
    if not empresa_dir: return None
    env_file = empresa_dir / ".env"
    if not env_file.exists():
        info(f"Sin .env en [dim]{empresa_dir.name}[/] → se usa el Riesgo CREM calculado.")
        return None
    if not _leer_env(env_file).get("TRENDAI_API_KEY", "").strip():
        info("TRENDAI_API_KEY sin configurar → se usa el Riesgo CREM calculado. "
             "[dim](rellénala en el .env de la empresa para leer el score real del portal)[/]")
        return None
    try:
        from trendai_api import TrendAIClient
        client = TrendAIClient.from_env(str(env_file))
        res = client.get_cyber_risk_index()
        if res.get("ok"):
            ok(f"Cyber Risk Index obtenido de Vision One API: [cyan bold]{res['score']}[/] ({res['level']})")
            return res["score"]
        degradado("Riesgo CREM", f"la API de Vision One respondió sin datos ({res.get('message')})",
                  "se usa el score calculado, que puede no coincidir con el del portal")
        return None
    except Exception as e:
        degradado("Riesgo CREM", f"fallo al consultar la API de Vision One ({e})",
                  "se usa el score calculado, que puede no coincidir con el del portal")
        return None


def _pedir_riesgo_crem_manual(score_auto: float, no_input: bool = False, empresa_dir: Optional[Path] = None) -> Optional[float]:
    """
    Permite fijar manualmente el Riesgo CREM cuando la heurística interna
    no coincide con el score real mostrado en el portal Vision One.
    Permite escribir 'api' para consultar automáticamente la API si hay .env.
    Devuelve None si se mantiene el valor automático.
    """
    if no_input: return None
    console.print()
    console.print(Rule("[bold cyan]Riesgo CREM[/]", style="cyan", align="left"))
    console.print(f"  [dim]Score calculated automáticamente: [cyan]{score_auto}[/][/]")
    console.print("  [dim]Si prefieres usar el valor exacto del portal, escribe el número (ej: 36.2) o [cyan]'api'[/] para obtenerlo por API.[/]")
    console.print("  [dim]Enter = mantener el valor automático.[/]")
    resp = input("  → ").strip().replace(",", ".")
    if not resp: return None
    if resp.lower() == "api":
        score_api = _obtener_riesgo_crem_api(empresa_dir)
        if score_api is not None:
            return score_api
        return None
    try:
        valor = float(resp)
    except ValueError:
        warn("Valor no válido, se mantiene el score automático.")
        return None
    if not (0 <= valor <= 100):
        warn("Fuera de rango (0-100), se mantiene el score automático.")
        return None
    ok(f"Riesgo CREM fijado manualmente: [cyan]{valor}[/]")
    return valor


# ==============================================================================
# 16. UTILS
# ==============================================================================
def mes_a_espanol(s: str) -> str:
    sl = s.lower()
    for es in MESES_ES.values():
        if sl.startswith(es.lower()): return s
    for en, es in MESES_ES.items():
        if sl.startswith(en.lower()): return es + s[len(en):]
    return s

def resumen_final(ruta_word, ruta_html, ruta_html_eje, ruta_log, elapsed):
    lines = []
    if ruta_word:     lines.append(f"[green]Word:[/]         [dim]{ruta_word}[/]")
    if ruta_html:     lines.append(f"[green]HTML técnico:[/] [dim]{ruta_html}[/]")
    if ruta_html_eje: lines.append(f"[green]HTML ejecutivo:[/][dim]{ruta_html_eje}[/]")
    lines.append(f"[green]Log:[/]          [dim]{ruta_log}[/]")
    console.print()
    console.print(Panel(
        "\n".join(lines),
        title="[bold green]PROCESO COMPLETADO ✓[/]",
        subtitle=f"[dim]{elapsed:.1f}s[/]",
        border_style="green", expand=False))
    console.print()

# ==============================================================================
# 17. MAIN
# ==============================================================================

# ==============================================================================
# MENÚ DE SELECCIÓN DE EMPRESA
# ==============================================================================
def _seleccionar_empresa() -> str:
    """
    Muestra un menú numerado con las carpetas de empresa detectadas.
    Una carpeta es empresa si contiene una subcarpeta CSV/ o INFORMES/.
    """
    # Detectar carpetas de empresa en el directorio actual
    empresas = []
    _SKIP = {"plantilla","datos","__pycache__",".git","default","debug","info_doc","CLIENTES"}
    # Buscar primero en CLIENTES/ (nueva estructura); fallback al directorio raíz
    _base_cli = BASE_DIR / "CLIENTES"
    _base_search = _base_cli if _base_cli.is_dir() else BASE_DIR
    for p in sorted(_base_search.iterdir()):
        if not p.is_dir(): continue
        if p.name in _SKIP: continue
        if p.name.startswith("."): continue
        # Es empresa si tiene CSV/ o INFORMES/ o csv/ dentro
        if (p / "CSV").is_dir() or (p / "csv").is_dir() or (p / "INFORMES").is_dir():
            empresas.append(p.name)

    console.print()
    console.print(Rule("[bold cyan]Selección de empresa[/]", style="cyan", align="left"))

    if not empresas:
        # No hay carpetas de empresa — pedir nombre libre
        console.print("  [dim]No se encontraron carpetas de empresa existentes.[/]")
        console.print("  [cyan]Nombre de la empresa[/]  (se creará la carpeta)\n")
        return input("  → ").strip()

    # Mostrar tabla de selección
    table = RichTable(box=None, padding=(0, 2), show_header=False)
    table.add_column("Núm", style="bold cyan", width=4)
    table.add_column("Empresa", style="bold")
    table.add_column("CSV", style="dim")
    table.add_column("Histórico", style="dim")

    for i, nombre in enumerate(empresas, 1):
        p = Path(nombre)
        csv_ok   = "✓" if (p / "CSV").is_dir() else "—"
        hist_ok  = "✓" if (p / "INFORMES" / "CSV").is_dir() else "—"
        table.add_row(f"[{i}]", nombre, f"CSV {csv_ok}", f"Histórico {hist_ok}")

    console.print(table)
    console.print()
    console.print(f"  [dim]0 = introducir nombre nuevo[/]\n")

    while True:
        console.print(f"  Selecciona empresa [cyan][1-{len(empresas)}][/] o [cyan][0][/] para nueva:")
        val = input("  → ").strip()
        if val == "0":
            console.print("  [cyan]Nombre de la nueva empresa[/]")
            return input("  → ").strip()
        if val.isdigit() and 1 <= int(val) <= len(empresas):
            elegida = empresas[int(val) - 1]
            ok(f"Empresa seleccionada: [yellow bold]{elegida}[/]")
            return elegida
        warn(f"Opción inválida. Introduce un número entre 0 y {len(empresas)}.")


def main():
    t0 = time.monotonic()
    # Portabilidad: trabajar SIEMPRE desde la carpeta del script, de modo que
    # todas las rutas relativas (CLIENTES/, plantilla/, subprocesos a la API)
    # se resuelvan igual en cualquier equipo y sin importar desde dónde se lance.
    os.chdir(BASE_DIR)
    parser = argparse.ArgumentParser(description=f"{PROYECTO} v{VERSION}")
    parser.add_argument("--mes",       metavar="MES_ANO", help='Ej: "Mayo 2026"')
    parser.add_argument("--no-input",  action="store_true",
                        help="Modo no interactivo (cron). Requiere --mes.")
    parser.add_argument("--riesgo-crem", metavar="SCORE", type=float, default=None,
                        help="Fija manualmente el Riesgo CREM (0-100, admite decimales) en vez del "
                             "calculado automáticamente. Útil cuando difiere del score real del portal.")
    parser.add_argument("--solo-word", action="store_true",
                        help="Solo regenerar Word/HTML usando caché datos/")
    parser.add_argument("--excels",    action="store_true",
                        help="Generar xlsx de revisión en informe/excels/")
    parser.add_argument("--conservar-csv", action="store_true",
                        help="No mover los CSV al histórico; copiarlos y conservarlos en CSV/")
    parser.add_argument("--empresa",   metavar="NOMBRE",
                        help='Nombre de la empresa cliente. Ej: "Ayuntamiento de Barcelona"')
    parser.add_argument("--template",  default="ambos",
                        choices=["tecnico","ejecutivo","ambos"],
                        help="Tipo de informe: tecnico (Word+HTML), ejecutivo (HTML light), ambos")

    parser.add_argument("--enriquecer-cve", action="store_true",
                        help="(Obsoleto) El enriquecimiento de CVEs (NVD + CISA KEV + EPSS) "
                             "ahora se realiza SIEMPRE. Este flag se mantiene por compatibilidad.")
    parser.add_argument("--prueba", action="store_true",
                        help="Modo prueba: genera en [EMPRESA]/PRUEBAS/ sin archivar los CSV, "
                             "sin persistir el Riesgo CREM ni tocar el histórico.")
    parser.add_argument("--api-riesgo", action="store_true",
                        help="Consulta automáticamente la API de Vision One para obtener el Cyber Risk Index real.")
    args = parser.parse_args()


    banner()

    # Período
    _ayer = date.today().replace(day=1) - timedelta(days=1)
    default_mes = mes_a_espanol(_ayer.strftime("%B %Y"))
    if args.mes:
        mes = mes_a_espanol(args.mes)
    elif args.no_input:
        mes = default_mes; info(f"--no-input: usando período por defecto [{mes}]")
    else:
        console.print(f"\n  [cyan]Período del informe[/]  (Enter = [yellow]{default_mes}[/])\n")
        val = input("  → ").strip()
        mes = mes_a_espanol(val) if val else default_mes
    info(f"Período: [yellow bold]{mes}[/]")

    # Empresa — menú de selección rápida por carpetas existentes
    if args.empresa:
        empresa = args.empresa.strip()
    elif args.no_input:
        empresa = ""
    else:
        empresa = _seleccionar_empresa()
    if empresa:
        info(f"Empresa: [yellow bold]{empresa}[/]")

    mes_safe = mes.replace("/","-").replace(" ","_")

    # El contexto se construye ENTERO aquí y se instala una sola vez. Antes eran
    # cuatro bloques `global` sueltos repartidos por la función.
    _cli = BASE_DIR / "CLIENTES"
    _nombre_dir = (empresa.replace("/","_").replace("\\","_").replace(":","_").strip()
                   if empresa else "default")
    empresa_dir = (_cli / _nombre_dir) if _cli.is_dir() else (BASE_DIR / _nombre_dir)

    # Salida: [EMPRESA]/INFORMES/[MES_AÑO]/ · en modo prueba → [EMPRESA]/PRUEBAS/[MES_AÑO]/
    if args.prueba:
        dir_informe = empresa_dir / "PRUEBAS" / mes_safe
        # Caché de datos aislada en PRUEBAS/ para no tocar nada de INFORMES/.
        dir_datos = dir_informe / "datos"
    else:
        dir_informe = empresa_dir / "INFORMES" / mes_safe
        dir_datos = (empresa_dir / "INFORMES" / "CSV"
                     / f"csv-{mes_safe.replace('_', '-').lower()}" / "datos")

    ctx = Contexto(empresa=empresa, empresa_dir=empresa_dir,
                   dir_csv=empresa_dir / "CSV",
                   dir_datos=dir_datos, dir_informe=dir_informe)
    instalar_contexto(ctx)

    for _d in (ctx.dir_csv, ctx.dir_informe, ctx.dir_datos):
        _d.mkdir(parents=True, exist_ok=True)
    if args.prueba:
        warn("MODO PRUEBA: no se archivan CSV, no se persiste el Riesgo CREM ni se toca el histórico.")
    info(f"Carpeta CSV:       [dim]{ctx.dir_csv}[/]")
    info(f"Carpeta histórica: [dim]{ctx.dir_historico}[/]")
    info(f"Carpeta informe:   [dim]{ctx.dir_informe}[/]")

    # Leer config.json de la empresa
    cfg = _leer_config(ctx.empresa_dir)

    ruta_log = _setup_log(mes_safe)

    # Validación
    seccion("Validando CSVs", 0, 4)
    _validar_csvs()

    # Carga
    datos = cargar_todos(usar_cache=args.solo_word)
    if args.excels: paso_generar_excels(datos)

    # Diff CVEs
    seccion("Comparando CVEs con mes anterior", 1, 4)
    dir_anterior = _buscar_dir_anterior(mes)
    if dir_anterior:
        info(f"Directorio anterior encontrado: [dim]{dir_anterior}[/]")
    else:
        info("No se encontró carpeta csv-MES-AÑO anterior. Sin comparación CVE.")
    diff_cves = comparar_cves(datos, dir_anterior)

    # Generar
    # Análisis de cambios y reincidentes
    cambios = analizar_cambios(datos, dir_anterior, diff_cves)
    meses_min = cfg.get("meses_reincidente", 2)
    reincidentes = detectar_reincidentes(CTX.empresa_dir, datos, mes, meses_min)
    if reincidentes:
        ok(f"Reincidentes: [new]{len(reincidentes)}[/] CVEs sin resolver {meses_min}+ meses")
    tendencia_hist = leer_tendencia_historica(CTX.empresa_dir, mes)

    # Enriquecimiento de CVEs (NVD + CISA KEV + EPSS) — SIEMPRE activo.
    # La primera vez descarga de NVD (usa la API key si está configurada) y cachea;
    # las siguientes ejecuciones son casi instantáneas.
    cves_prioritarios = []
    seccion("Enriqueciendo CVEs (NVD · KEV · EPSS)", 1, 4)
    _nvd_key = _resolver_nvd_key(cfg, CTX.empresa_dir)
    if _nvd_key:
        info(f"NVD API key detectada [dim](…{_nvd_key[-4:]})[/] → 50 req/30s")
    else:
        warn("Sin NVD API key → 5 req/30s (más lento). Añade NVD_API_KEY al .env de la raíz del proyecto.")
    _cache_cve = BASE_DIR / "cve_cache"
    _enrich = enriquecer_cves(datos, _cache_cve, api_key=_nvd_key)
    cves_prioritarios = construir_cves_prioritarios(datos, _enrich, top_n=None)
    if cves_prioritarios:
        ok(f"CVEs prioritarios con solución: [new]{len(cves_prioritarios)}[/]")

    # Riesgo CREM (obtiene Cyber Risk Index de API si existe key; de lo contrario usa manual/automático)
    _risk_auto = calcular_risk_score(datos, diff_cves, reincidentes, tendencia_hist)
    riesgo_crem_manual = args.riesgo_crem
    if riesgo_crem_manual is None and CTX.empresa_dir:
        if getattr(args, "api_riesgo", False) or (CTX.empresa_dir / ".env").exists():
            riesgo_crem_manual = _obtener_riesgo_crem_api(CTX.empresa_dir)
    if riesgo_crem_manual is None:
        riesgo_crem_manual = _pedir_riesgo_crem_manual(_risk_auto["score"], no_input=args.no_input, empresa_dir=CTX.empresa_dir)



    if riesgo_crem_manual is not None:
        _final_score = riesgo_crem_manual
        if   _final_score >= 75: _final_nivel = "CRÍTICO"
        elif _final_score >= 50: _final_nivel = "ALTO"
        elif _final_score >= 25: _final_nivel = "MEDIO"
        else:                    _final_nivel = "BAJO"
    else:
        _final_score, _final_nivel = _risk_auto["score"], _risk_auto["nivel"]
    _risk_info_final = {"score": _final_score, "nivel": _final_nivel, "manual": riesgo_crem_manual is not None}

    template = getattr(args, "template", "tecnico")
    ruta_word     = None
    ruta_html     = None
    ruta_html_eje = None

    if template in ("tecnico", "ambos"):
        ruta_word = paso_generar_word(mes, datos, diff_cves, ctx.empresa, cfg, cambios, reincidentes, enrich_map=_enrich)
        ruta_html = paso_generar_html(mes, datos, diff_cves, ctx.empresa, cfg, cambios, reincidentes, tendencia_hist,
                                       riesgo_crem_manual=riesgo_crem_manual,
                                       cves_prioritarios=cves_prioritarios,
                                       enrich_map=_enrich)

    if template in ("ejecutivo", "ambos"):
        ruta_html_eje = paso_generar_html_ejecutivo(mes, datos, diff_cves, ctx.empresa, cfg, cambios, reincidentes, tendencia_hist,
                                                     riesgo_crem_manual=riesgo_crem_manual,
                                                     cves_prioritarios=cves_prioritarios,
                                                     enrich_map=_enrich)
        if ruta_html_eje and ruta_html_eje.exists():
            ruta_pdf = ruta_html_eje.with_suffix(".pdf")
            info(f"Generando PDF ejecutivo en: [dim]{ruta_pdf.name}[/]")
            ok_pdf = convertir_html_a_pdf(ruta_html_eje, ruta_pdf)
            if ok_pdf:
                ok(f"PDF ejecutivo guardado → [bold]{ruta_pdf}[/]")
            else:
                degradado("PDF ejecutivo", "PyQt6 no pudo generar el PDF", "solo queda la versión HTML del informe ejecutivo")

    # Archivar CSVs al histórico del mes + persistir Riesgo CREM final.
    # Por defecto se MUEVEN (CSV/ queda limpia); en modo caché o con --conservar-csv se copian.
    # En modo prueba NO se archiva nada: los CSV quedan intactos y el histórico no se toca.
    if args.prueba:
        info("Modo prueba: se omite el archivado de CSV y la persistencia del Riesgo CREM.")
    else:
        _mover_csv = not (args.conservar_csv or args.solo_word)
        _actualizar_csv_historico(CTX.empresa_dir, mes, risk_info=_risk_info_final, mover=_mover_csv)

    # Abrir HTML automáticamente si config lo indica
    if cfg.get("abrir_html_al_terminar", False):
        import webbrowser
        for f in filter(None, [ruta_html, ruta_html_eje]):
            webbrowser.open(str(f.resolve())); break

    if template in ("tecnico", "ambos") and ruta_word is None:
        err("No se pudo generar el Word."); sys.exit(1)

    resumen_final(ruta_word, ruta_html, ruta_html_eje, ruta_log, time.monotonic() - t0)
    # Lo último que se ve: qué salió incompleto. Un ✓ final con avisos debajo es
    # honesto; un ✓ final a secas sobre un informe recortado, no.
    _resumen_degradaciones()

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        console.print("\n[yellow]  Interrumpido.[/]"); sys.exit(0)
    except Exception:
        console.print("\n[bold red]  ERROR INESPERADO:[/]"); traceback.print_exc(); sys.exit(1)